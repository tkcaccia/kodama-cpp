# SPDX-FileCopyrightText: 2026 Stefano Cacciatore
# SPDX-License-Identifier: MIT

from __future__ import annotations

import csv
from pathlib import Path
from textwrap import dedent

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
MANUSCRIPT = ROOT / "kodama_cpp_jmlr_manuscript.docx"
SELF_REVIEW = ROOT / "kodama_cpp_jmlr_self_review.docx"
TEX = ROOT / "kodama_cpp_jmlr_technical_supplement.tex"
BIB = ROOT / "kodama_cpp_refs.bib"
ARCH_FIGURE = ROOT / "kodama_cpp_architecture.png"
SENSITIVITY_FIGURE = ROOT / "kodama_m_tcycle_sensitivity.png"
ENSEMBLE_CONVERGENCE_FIGURE = ROOT / "jmlr_pilot_20260716" / "ensemble_convergence.png"
NONSPATIAL_PANEL_FIGURE = (
    ROOT
    / "jmlr_nonspatial_panel_20260718"
    / "nonspatial_visualization_validation.png"
)
METREF_BACKEND_FIGURE = ROOT / "MetRef__pls_lda__cpu_cuda_optimized_20260720.png"
HPC_RESULT_DIR = ROOT / "jmlr_hpc_kodama_20260806"
HPC_VISUALIZATION_FIGURE = HPC_RESULT_DIR / "kodama_vs_classic_silhouette.png"
HPC_VISUALIZATION_SUMMARY = HPC_RESULT_DIR / "kodama_vs_classic_dataset_summary.csv"
HPC_INFERENCE_SUMMARY = HPC_RESULT_DIR / "kodama_vs_classic_inference.csv"
HPC_DATASET_INVENTORY = HPC_RESULT_DIR / "dataset_inventory.csv"
HPC_RUNTIME_SUMMARY = HPC_RESULT_DIR / "statistical_analysis_chiamaka" / "runtime_by_dataset.csv"
IMAGENET_COMPARISON_FIGURE = (
    ROOT / "jmlr_imagenet_comparison_20260807" / "imagenet_classic_vs_kodama.png"
)
PENDIGITS_ADVERSE_FIGURE = (
    ROOT / "jmlr_cycle20_pendigits_m100_t100_20260808" / "pendigits_cpu_adverse.png"
)
COIL20_KNN_FIGURE = (
    ROOT / "jmlr_cycle21_coil20_m100_t100_20260808" / "coil20_knn_cpu_metal.png"
)
SATIMAGE_FIGURE = (
    ROOT / "jmlr_cycle24_satimage_metal_corrected_m100_t100_20260808" / "satimage_cpu_metal_corrected.png"
)


TOKENS = {
    "font": "Calibri",
    "body_size": 11,
    "h1_size": 16,
    "h2_size": 13,
    "h3_size": 12,
    "heading_blue": RGBColor(0x2E, 0x74, 0xB5),
    "heading_dark": RGBColor(0x1F, 0x4D, 0x78),
    "title": RGBColor(0x0B, 0x25, 0x45),
    "muted": RGBColor(0x55, 0x55, 0x55),
    "table_fill": "E8EEF5",
    "callout_fill": "F4F6F9",
    "border": "B7C7D9",
}


AUTHORS = [
    ("Moussa Kassim", "1,2,*", "moussa.kassim@icgeb.org"),
    ("Martin Ocharo", "1,2,*", "martin.ocharo@icgeb.org"),
    ("Dalia Ahmed", "1", "dalia.ahmed@icgeb.org"),
    ("Dupe Ojo", "1", "dupe.ojo@icgeb.org"),
    ("Alessia Vignoli", "3,4", "vignoli@cerm.unifi.it"),
    ("Leonardo Tenori", "3,4,†", "tenori@cerm.unifi.it"),
    ("Stefano Cacciatore", "1,2,†", "stefano.cacciatore@icgeb.org"),
]

AUTHOR_NOTES = [
    ("*", "Moussa Kassim and Martin Ocharo contributed equally."),
    ("†", "Leonardo Tenori and Stefano Cacciatore are co-corresponding authors."),
]

AFFILIATIONS = [
    (
        "1",
        "Bioinformatics Unit, International Centre for Genetic Engineering and Biotechnology "
        "(ICGEB), Cape Town 7925, South Africa",
    ),
    (
        "2",
        "Department of Integrative Biomedical Sciences, Institute of Infectious Disease & Molecular "
        "Medicine (IDM), University of Cape Town, Cape Town 7925, South Africa",
    ),
    (
        "3",
        'Department of Chemistry "Ugo Schiff", University of Florence, Sesto Fiorentino, Italy',
    ),
    (
        "4",
        "Magnetic Resonance Center (CERM), University of Florence, Sesto Fiorentino, Italy",
    ),
]

ABSTRACT = (
    "KODAMA searches for latent structure by maximizing the cross-validated predictability of an "
    "evolving label vector. We present kodama-cpp and make five contributions. First, KODAMA is "
    "implemented as a standalone C++17 library with a typed, wrapper-independent API, allowing thin "
    "R and Python interfaces to share one numerical implementation. Second, native multicore CPU, "
    "NVIDIA CUDA, and Apple Metal backends implement float32 nearest-neighbor search, k-means, "
    "PCA, graph operations, and reusable workspaces without silent CPU fallback. Third, kodama-cpp "
    "introduces label-aware SIMPLS followed by LDA in the requested-component latent space as a new "
    "KODAMA evaluator, distinct from the historical PLS-DA decoder. Fourth, the current search uses a common prediction-guided proposal and "
    "evaluation scaffold, while PLS-LDA adds a disclosed transition-coarsening move and fragmentation "
    "term to its state score. Proposal size decreases smoothly, an error-scaled temperature permits "
    "early exploration, and each cycle performs exactly one new cross-validation pass. Fifth, the landmark-projection strategy "
    "is formalized for general matrix input through exact-quota "
    "stratified sampling, optimization on landmarks, and supervised projection to the remaining "
    "samples. These policies retain cross-validated predictability as the internal signal but need "
    "not reproduce historical stochastic trajectories. We validate backend identity, numerical "
    "behavior, runtime, memory, landmark-selection contracts, and external-label diagnostics "
    "separately so that systems acceleration is not conflated with the optimization objective. In "
    "an explicitly exploratory 10-dataset matched analysis, PLS-LDA KODAMA followed by UMAP improved "
    "truth-label silhouette on six datasets (median change +0.025, bootstrap 95% interval "
    "[-0.067, +0.084]); KNN and openTSNE were less favorable, precluding a universal claim."
)


SECTIONS = [
    (
        "1. Introduction",
        [
            (
                "High-dimensional data analysis often begins with clustering or dimensionality "
                "reduction, but neither operation directly asks whether the discovered groups are "
                "predictable from the data. KODAMA takes a different view: a useful labeling of "
                "samples should be reproducible by a classifier under cross-validation. The labels "
                "are not treated as ground truth; they are optimization variables whose quality is "
                "measured by held-out prediction."
            ),
            (
                "The original KODAMA paper introduced this idea as knowledge discovery by accuracy "
                "maximization, and the subsequent R package made it available as an unsupervised "
                "and semi-supervised feature-extraction tool for noisy high-dimensional data. The "
                "Bioinformatics package paper describes the algorithm as repeated random label "
                "initialization, iterative maximization of cross-validated accuracy by label "
                "swapping, and construction of a dissimilarity matrix from the resulting label "
                "vectors. The R documentation exposes the same logic through M independent "
                "iterative processes, Tcycle optimization steps, KNN or PLS-DA classifiers, and "
                "separate KODAMA.matrix and KODAMA.visualization functions."
            ),
            (
                "The repeated classifier fits inside KODAMA create a systems problem that is not "
                "resolved by translating individual R expressions into C++. Neighbor indices, fold "
                "matrices, class encodings, latent projections, and evolving labels must be reused "
                "across many cross-validation calls without changing the accepted label trajectory. "
                "A useful accelerator implementation must therefore preserve the mathematical state "
                "machine while changing where that state lives and how it moves."
            ),
            (
                "kodama-cpp addresses this problem with a standalone C++17 core and explicit CPU, "
                "CUDA, and Apple Metal backends. It currently supports two classifier families in "
                "the KODAMA optimization layer: KNN for local neighborhood consistency and PLS-LDA "
                "for low-rank linear discrimination. R and Python wrappers call the same typed API, "
                "so fold construction, proposals, acceptance decisions, and result metadata are not "
                "reimplemented by each language interface. Their public functions, defaults, accepted "
                "choices, one-based best-run convention, and result fields are matched; Python only "
                "transliterates dotted R argument names to snake_case."
            ),
            (
                "The manuscript is organized around five novelties. The first is software ownership: "
                "KODAMA becomes a standalone C++ library from which R and Python wrappers can be built "
                "without reimplementing folds, proposals, classifiers, or graph construction. The "
                "second is heterogeneous execution through native CUDA and Apple Metal backends under "
                "the same float32 result contract. The third is a label-aware SIMPLS evaluator followed "
                "by LDA in the requested-component latent space, a new KODAMA classifier route rather than "
                "a reimplementation of the historical PLS-DA decoder. The fourth is a guided label-evolution strategy that "
                "uses out-of-fold predictions, class-transition evidence, adaptive proposal sizes, "
                "degeneracy guards, and error-scaled cooling while preserving one new CV evaluation per "
                "cycle. The fifth is a general landmark-projection procedure that makes the expensive "
                "optimization depend on a representative "
                "subset and then projects each optimized solution to all samples."
            ),
            (
                "We deliberately frame this contribution as a software-methods paper, not as a "
                "claim that KODAMA replaces clustering, semi-supervised learning, or manifold "
                "learning. Cross-validated accuracy is the internal optimization criterion used "
                "to construct candidate label vectors. External labels, when available, are used "
                "only for diagnostics such as adjusted Rand index, silhouette, local purity, and "
                "runtime-quality tradeoffs."
            ),
        ],
    ),
    (
        "2. KODAMA objective",
        [
            (
                "Let X in R^(n x p) denote the input matrix, c in {1,...,K}^n a candidate label "
                "vector, F a classifier family, and Pi a fold assignment. For each validation "
                "sample i, the classifier is trained on samples not in Pi(i) and predicts c_i from "
                "x_i. The empirical objective is the held-out accuracy A(c; X, F, Pi). KODAMA "
                "searches for a labeling c* that maximizes this quantity, turning unsupervised "
                "structure discovery into a self-guided supervised prediction problem."
            ),
            (
                "One KODAMA run starts from an initial partition. In the historical R interface, "
                "this can be one label per sample or a clustering-derived vector W; a fraction of "
                "samples can be selected for each run, with 75% as the documented default. The "
                "current C++ interface keeps the same idea but exposes starting labels, splitting, "
                "landmarks, and fixed or constrained groups as typed inputs. During each cycle, "
                "misclassified samples or constrained groups generate candidate label moves. A "
                "move is accepted if it improves the objective, or under a temperature rule that "
                "allows exploration early in the run. Landmark projection is treated as a general "
                "matrix algorithm: label "
                "optimization is performed on selected landmarks and the chosen classifier restores "
                "a complete label vector before ensemble construction."
            ),
            (
                "The M runs are independent. This independence is important both statistically, "
                "because it averages over random initialization and proposal order, and "
                "computationally, because backend schedulers can process runs without "
                "changing the objective. The final KODAMA representation is obtained from the "
                "ensemble of optimized label vectors: pairs of samples that repeatedly receive "
                "compatible labels are pulled closer in the derived graph or dissimilarity, while "
                "unstable pairs are weakened."
            ),
        ],
    ),
    (
        "3. Implementation architecture",
        [
            (
                "The core library exposes C++ functions rather than wrapper-specific entry points. "
                "A MatrixView accepts double or float inputs, but all analysis paths convert to and "
                "retain float32 matrices and workspaces. The core owns fold assignments, compact class "
                "encodings, current and best label vectors, classifier buffers, and typed result "
                "metadata. KODAMAGraphResult owns only graph arrays, two PCA starts, and metadata; "
                "it never owns or aliases the source MatrixView. The public API is grouped into "
                "cross-validation kernels, core label optimization, graph preparation, KODAMA matrix "
                "construction, PCA, and embedding utilities."
            ),
            (
                "Backend selection is strict and observable. CPU, CUDA, and Metal entry points report "
                "the backend that actually executed, and an unavailable accelerator raises an error. "
                "This rule is essential for reproducible benchmarking because a nominal GPU call that "
                "silently falls back to CPU would preserve outputs while invalidating performance claims."
            ),
            (
                "KNN kernels have backend-specific implementations under one voting contract. The "
                "dependency-light CPU path uses a package-owned float32 HNSW implementation with "
                "a deterministic connected base seed before per-node-locked parallel insertion, "
                "batched candidate distance evaluation, contiguous graph storage, and parallel "
                "batched queries. CUDA "
                "provides package-owned exact and recall-tuned IVF-Flat search plus GPU k-means using "
                "CUDA Toolkit libraries. Apple Metal provides native exact and recall-tuned IVF-Flat "
                "search plus GPU k-means using Metal compute kernels. Full-data accelerator graph "
                "preparation uses exact search when n <= 5,000 or n^2 p <= 2e8 and otherwise builds "
                "a resident IVF-Flat index tuned against a deterministic exact pilot to at least 0.99 "
                "recall. The selected index, nlist, nprobe, and pilot recall are returned. Neighbor "
                "structures are reused where the mathematics permits, avoiding repeated wrapper-level "
                "index construction."
            ),
            (
                "PLS-LDA uses a SIMPLS strategy following the fastPLS implementation rather than "
                "a simplified SVD approximation. CPU, CUDA, and Metal paths operate on float32 "
                "matrices and compute class-label cross-products directly instead of constructing "
                "dense one-hot response matrices where possible. The "
                "requested component count is treated as the evaluated component count whenever "
                "mathematically feasible, rather than being replaced by an internal model-selection "
                "shortcut. The library exposes no PLS-cKNN classifier; KODAMA optimization is limited "
                "to KNN and PLS-LDA."
            ),
            (
                "CUDA and Metal KODAMA calls keep one full-data graph resident through all M runs "
                "and allocate one persistent scratch lane per concurrent worker. Each lane retains "
                "landmark and fold layouts, compact labels, projections, KNN votes, cross-products, "
                "PLS weights, and LDA workspaces. Landmark indices and labels are scattered behind "
                "per-lane epoch tags directly into a resident run-major M-by-n result matrix. Exact "
                "constrained-majority relabeling and agreement-graph correction consume that matrix "
                "on-device, removing per-run projected-label downloads and the final whole-matrix "
                "re-upload. A new M run changes landmark membership and fold "
                "contents, so those values are refreshed once in the existing buffers; Tcycle "
                "proposals overwrite only label-dependent state. This is buffer and data-lifecycle "
                "residency, while stochastic proposal decisions remain host orchestrated."
            ),
            (
                "Latent-space LDA is fitted from sufficient statistics. For training scores T, "
                "class counts n_c, class means mu_c, q retained components, and C active classes, "
                "the pooled within-class covariance is [T' T - sum_c n_c mu_c mu_c'] / "
                "max(1, n - C). Each discriminant direction solves Sigma w_c = mu_c by Cholesky "
                "factorization and triangular substitution; Sigma is never explicitly inverted. "
                "The diagonal regularizer is lambda = rho s, where s = trace(Sigma) / q when "
                "finite and positive and s = 1 otherwise. The deterministic sequence rho in "
                "{1e-8, 1e-6, 1e-5, 1e-4, 1e-3, 1e-2} is advanced only when Cholesky fails. "
                "Prediction retains the ordinary LDA discriminant t' w_c - 0.5 mu_c' w_c + "
                "log(n_c / n); the sequence is a numerical safeguard, not a tuned parameter."
            ),
            (
                "The library also exposes float32 randomized PCA as a preprocessing and embedding-"
                "initialization primitive. A Gaussian range sketch, backend-specific automatic "
                "oversampling and power-iteration policy, orthonormal basis, and compact eigensolve "
                "produce scores, loadings, singular values, and explained variance. CPU, CUDA, and "
                "Metal entry points share this contract. The implementation follows the current "
                "fastEmbedR/fastPLS strategy but is standalone and does not link either package, "
                "Armadillo, Eigen, or RAFT."
            ),
            (
                "UMAP and openTSNE are direct standalone ports of pinned fastEmbedR revision "
                "814350a5. UMAP builds fuzzy graphs by default, with binary weighting "
                "available explicitly; both are constructed directly in float32 CSR form, "
                "including the smooth-kNN bandwidth calculation and epochs-per-sample schedule; "
                "openTSNE uses the matching sparse affinities and exact or FFT-grid repulsion. "
                "The current openTSNE port also retains persistent CPU workers, cached FFT plans, "
                "parallel FFT preparation, and chunked CUDA graph launches. The CUDA path reuses "
                "pooled workspaces. Metal provides a native UMAP optimizer using the same graph, "
                "bandwidth, and epoch schedule with fixed-point atomic updates, plus the native "
                "fastEmbedR FFT-grid openTSNE schedule with portable float-bit CAS accumulation. "
                "None of these paths "
                "links fastEmbedR or R."
            ),
            (
                "Visualization initialization is part of the typed backend contract. One "
                "backend-native float32 PCA produces both starts from the raw matrix: openTSNE "
                "centers and scales scores to maximum component standard deviation 1e-4, while "
                "UMAP centers, max-absolute scales to 10, adds seeded float32 jitter, and recenters. "
                "Explicit coordinates take precedence; otherwise wrappers use explicit raw data, "
                "then a stored start only when its CPU/CUDA/Metal backend matches, and finally the "
                "reported graph-spectral UMAP or random openTSNE fallback. In a controlled CPU "
                "openTSNE comparison with matched neighbors and settings, fastEmbedR and kodama-cpp "
                "coordinates were identical. CPU UMAP is algorithmically matched but not asserted "
                "bitwise because kodama-cpp retains float32 iterative state whereas the compared R "
                "entry point stores its embedding in a double matrix."
            ),
            (
                "KODAMA.graph prepares one full-data neighbor graph and backend-matched UMAP and "
                "openTSNE PCA starts without retaining the raw matrix. Four input forms share one "
                "dispatch contract: raw data; a prepared graph result; that result plus a separately "
                "supplied raw matrix; or a bare neighbor graph. Prepared and bare graph inputs avoid "
                "rebuilding neighbor search. With explicit raw features, landmark geometry and "
                "PLS-LDA use the original float32 variables; without them, the graph is converted into "
                "a PLS-compatible float32 representation with the standard self-tuning normalized "
                "Laplacian transform. Fixed-graph KNN reuses supplied neighborhoods and is not asserted "
                "to reproduce the stochastic trajectory of data-native KNN."
            ),
            (
                "The graph-only PLS-LDA transform is deterministic conditional on the supplied seed. "
                "For row i, sigma_i is the median finite neighbor distance, falling back to the row "
                "mean and then to 1e-6 when needed. Each directed edge has weight "
                "w_ij = exp(-d_ij^2 / max(1e-12, sigma_i sigma_j)); reciprocal edges are added by "
                "taking the larger weight, and the operator is symmetrically normalized as "
                "D^{-1/2} W D^{-1/2}. Random Gaussian feature columns are repeatedly multiplied by "
                "this sparse operator, orthonormalized for at least eight iterations, and finally "
                "standardized before entering the ordinary PLS-LDA core."
            ),
            (
                "The implementation deliberately keeps classifier paths clean. KODAMA optimization "
                "is exposed as KNN or PLS-LDA; auxiliary cross-validation kernels exist for testing "
                "and benchmarking, but the KODAMA optimizer is not a collection of benchmark-specific "
                "branches. The graph, embedding, and clustering utilities are similarly separated "
                "from the label-evolution loop."
            ),
        ],
    ),
    (
        "4. Relationship to public KODAMA literature",
        [
            (
                "The contribution of kodama-cpp is not a new definition of KODAMA. It is a "
                "reorganization of the published method into a reusable numerical system. The "
                "2014 PNAS paper is the methodological reference for accuracy maximization. The "
                "2017 Bioinformatics paper is the software reference for the R package, including "
                "KNN and PLS-DA classifiers, repeated suboptimal runs, and construction of a "
                "dissimilarity matrix from the optimized label vectors. The R documentation is the "
                "public reference for defaults such as M = 100, Tcycle = 20, 75% sample selection, "
                "and the output fields dissimilarity, accuracy, proximity, classification vectors, "
                "scores, entropy, and landmarks."
            ),
            (
                "The main differences are architectural and computational. The R package owns a "
                "complete R workflow and uses R-package dependencies for data handling, nearest "
                "neighbors, and visualization. kodama-cpp owns the numerical kernels and exposes "
                "them through C++ types so that R and Python wrappers can share one backend. This "
                "makes comparison with the KODAMA R 2.4.1/2.4 predecessor a compatibility question: held-out "
                "label predictability remains the central signal and the output roles can be mapped, "
                "while exact trajectories need not match because the current standard adds grouped "
                "adaptive proposals, transition-driven coarsening, and a disclosed degeneracy guard. "
                "Float32 storage, package-owned CPU/CUDA/Metal search, accelerator kernels, and "
                "wrapper-independent graph objects are additional implementation changes."
            ),
            (
                "The PLS routes must not be conflated. KODAMA R 2.4.1 exposes its historical "
                "PLS-DA decoder. kodama-cpp introduces label-aware SIMPLS followed by LDA in the "
                "requested-component latent space, with fixed component-count semantics and native "
                "float32 CPU, CUDA, and Metal execution. The MetRef evaluation therefore reports "
                "historical PLS-DA quantitatively, but treats it as a contextual predecessor rather "
                "than a classifier-matched implementation baseline."
            ),
            (
                "All predecessor performance comparisons in this article use KODAMA R "
                "2.4.1/2.4; later KODAMA releases are excluded from timing, quality, and "
                "methodological comparisons. Optional R wrapper utilities are documented "
                "separately and are not part of the predecessor performance or quality comparison."
            ),
        ],
    ),
    (
        "5. Evaluation",
        [
            (
                "We evaluate kodama-cpp at three levels. First, kernel benchmarks isolate KNNCV "
                "and PLSLDACV from the full KODAMA matrix pipeline so that nearest-neighbor search "
                "and PLS-LDA fitting can be interpreted separately. Second, core optimizer "
                "benchmarks measure repeated label-vector evolution over multiple seeds. Third, "
                "matrix-level validation measures end-to-end KODAMA.matrix calls and reports both "
                "cross-validated accuracy and external-label diagnostics where reference labels are "
                "available."
            ),
            (
                "The evaluation is intentionally separated into these layers because KODAMA has "
                "several scaling regimes. KNN acceleration is dominated by neighbor search and "
                "voting. PLS-LDA acceleration is dominated by repeated SIMPLS and LDA work across "
                "folds and cycles. The final graph and visualization stages depend primarily on "
                "the sparse graph size. Reporting the layers separately makes it possible to "
                "attribute a speedup to the part of the implementation that produced it."
            ),
            (
                "The evaluation also separates optimization from validation. Held-out CV accuracy is "
                "the internal classifier signal, while the current standard ranks proposals with the "
                "disclosed label-only acceptance score. Neither quantity is interpreted as external "
                "biological, chemical, or semantic truth. When reference labels are "
                "available, they are held outside the optimizer and reported after the fact. This "
                "distinction follows the broader cautionary literature on model selection, feature "
                "selection bias, and circular analysis."
            ),
            (
                "All timings in this section are wall-clock seconds from the benchmark drivers, "
                "and all accuracy values are computed on held-out cross-validation folds. CPU "
                "kernel measurements use the configured single-thread CPU path unless explicitly "
                "labeled otherwise; CUDA measurements use the same CUDA Toolkit runtime used by the "
                "test suite. Tables are intentionally separated by scope: kernel-level rows do not "
                "claim end-to-end speedups, and graph-input rows are treated as API validation unless "
                "paired with data-input KODAMA results under the same M, Tcycle, seed, and graph settings."
            ),
        ],
    ),
    (
        "6. Availability and reproducibility",
        [
            (
                "kodama-cpp is intended for release in the public tkcaccia/kodama-cpp GitHub "
                "repository. Original project code is licensed under MIT, while identified "
                "compatible third-party portions retain their upstream terms. In particular, the "
                "Metal backend is marked MIT AND Apache-2.0 because its fastEmbedR lineage retains "
                "the Faiss-mlx fused list-scan/top-k organization under Apache-2.0. The reviewed "
                "release architecture is a standalone C++17 core repository, a separate R wrapper "
                "repository, and Python binding source versioned with the core."
            ),
            (
                "Reproducibility is supported through CMake builds, CPU, CUDA, and Metal tests, wrapper "
                "validation tests, benchmark scripts, and explicit recording of backend parameters. "
                "GitHub Actions builds and tests the CPU core on Linux and macOS, builds the native "
                "Metal core on macOS, and installs and tests both wrapper packages. A separate LLVM "
                "workflow measures CPU source coverage, while Doxygen publishes the typed API and a "
                "compact C++/R/Python walkthrough at https://tkcaccia.github.io/kodama-cpp/. The C++ "
                "core does not depend on R data readers; wrappers translate host-language objects into "
                "contiguous matrices before calling the library."
            ),
            (
                "As of 6 August 2026, the public main branches identify abbreviated core commit "
                "0b8b873 and R-wrapper commit 7090592. Neither repository yet has an "
                "immutable release tag, the Python wrapper is still distributed within the core source "
                "tree, and the uploaded HPC manifests record git_commit as NA. The multi-dataset analysis "
                "is therefore reported as exploratory. A confirmatory release benchmark must be rerun from "
                "a clean tagged core and wrapper pair and record source checksums and the container digest."
            ),
            (
                "A versioned provenance audit maps every shipped source-like file to an SPDX license "
                "identifier and records pinned KODAMA, fastPLS, fastEmbedR, faissR, FAISS, Faiss-mlx, "
                "and cuVS snapshots. Stefano Cacciatore confirmed that tkcaccia is his historical Git "
                "identity and separately authorized MIT distribution of the KODAMA and "
                "fastPLS-derived portions whose copyright he owns. The audit preserves publication "
                "coauthor credit, third-party notices, and full license texts, and is enforced by a "
                "release-time source-header and retained-license checksum test."
            ),
        ],
    ),
    (
        "7. Limitations",
        [
            (
                "The current implementation prioritizes KNN and PLS-LDA because they are central "
                "to the KODAMA optimization principle and can be implemented cleanly on CPU, CUDA, "
                "and Apple Metal. Additional classifiers are outside this release unless they preserve the "
                "cross-validated accuracy objective and can be tested without changing the method "
                "for a specific benchmark."
            ),
            (
                "The main engineering limitation is that KODAMA performs many cross-validation "
                "fits. This makes memory locality, buffer reuse, label compaction, and GPU "
                "scheduling as important as asymptotic complexity. The present release includes "
                "float32 buffers, fold reuse, label compaction, a resident full-data graph, persistent "
                "per-worker CUDA/Metal classifier workspaces, and native accelerator KNN and PLS-LDA "
                "kernels. Landmark selection and proposal control remain host orchestrated; a fully "
                "device-side evolutionary state machine is an engineering extension rather than a "
                "claim of the current manuscript."
            ),
            (
                "The native CUDA neighbor-graph builder currently supports at most 256 retained "
                "neighbors per sample, whereas the CPU path supports larger rows. The implementation "
                "raises an error rather than silently truncating k. Benchmark protocols therefore "
                "record unsupported graph sizes as missing and never manufacture an accelerator row "
                "by changing a requested parameter."
            ),
        ],
    ),
]


CONTRIBUTION_ROWS = [
    (
        "Standalone C++ core",
        "Moves folds, classifiers, label evolution, projection, and graph construction into one C++17 library so thin R and Python wrappers share the same typed implementation.",
    ),
    (
        "CUDA and Metal backends",
        "Provides native float32 KNN, k-means, SIMPLS-LDA, PCA, and reusable accelerator state with strict backend identity and no silent CPU fallback.",
    ),
    (
        "Guided CV maximization",
        "Uses a common prediction-guided proposal/evaluation scaffold; PLS-LDA additionally applies transition coarsening and a fragmentation-aware state-score term.",
    ),
    (
        "Landmark projection",
        "Uses exact-quota stratified sampling, optimization on a representative subset, supervised projection, and ensemble graph correction.",
    ),
]


LANDMARKING_NOVELTY_PARAGRAPHS = [
    (
        "Landmark projection reduces the cost of repeated accuracy maximization by optimizing "
        "labels only on a selected subset, fitting a supervised model to each optimized landmark "
        "vector, and predicting labels for the remaining samples before constructing the ensemble "
        "dissimilarity. kodama-cpp implements this as a general algorithmic component for matrix "
        "inputs that does not depend on a wrapper language."
    ),
    (
        "The current selector replaces k-means with one center per requested landmark by a coarse "
        "partition followed by exact population-proportional allocation. For effective landmark "
        "count L and coarse stratum c containing n_c of n samples, q_c = L n_c / n is the expected "
        "allocation. The algorithm samples floor(q_c) rows without replacement, then assigns the "
        "remaining slots by randomized systematic rounding of the fractional quotas. It therefore "
        "returns exactly L distinct landmarks and satisfies E[L_c] = L n_c / n without requiring "
        "L k-means centers."
    ),
    (
        "Sampling strata and labels being optimized are intentionally different objects. After "
        "landmarks are selected, a separate expression-space partition initializes the candidate "
        "labels. KODAMA evolves those labels for Tcycle steps, stores the selected vector from each "
        "of M independent runs, and projects it to non-landmarks with the selected KNN or PLS-LDA "
        "classifier. The final graph is then corrected from agreement across complete projected "
        "vectors, so non-landmarks contribute to the returned representation even though the "
        "iterative CV search is performed on the landmark subset."
    ),
]


LANDMARK_VALIDATION_ROWS = [
    (
        "Selection contract",
        "Unit tests require the exact effective landmark count, distinct sampled rows, fixed-seed repeatability, and valid represented/occupied-stratum diagnostics.",
    ),
    (
        "Backend contract",
        "CPU and CUDA matrix runs expose matching landmark counts and diagnostics; selection remains host-orchestrated and independent of classifier backend.",
    ),
    (
        "Scalability isolation",
        "On flow18, selecting 750,016 landmarks from 300 coarse strata required 0.321 seconds. Recall-tuned resident CUDA IVF-Flat subsequently reduced the 100-neighbor graph from 296.482 to 25.938 seconds.",
    ),
    (
        "Quality evaluation",
        "Landmark tests report downstream CV accuracy, ensemble stability, and external diagnostics separately from selection time so representativeness is not inferred from speed alone.",
    ),
]


IMPLEMENTATION_NOVELTY_PARAGRAPHS = [
    (
        "kodama-cpp preserves cross-validated predictability as the KODAMA signal but contributes "
        "both a systems redesign and an explicit current search policy. The original R implementation "
        "remains the historical reference for accuracy maximization and ensemble construction. In the "
        "new library, numerical state previously recreated across wrapper calls is represented once in "
        "typed C++ objects and reused across folds, proposal cycles, and independent runs."
    ),
    (
        "This distinction matters because the runtime of KODAMA is the product of several nested "
        "operations: M independent searches, Tcycle proposal steps, cross-validation folds, and a "
        "classifier fit or vote inside each fold. Accelerating only one matrix multiplication leaves "
        "allocation, format conversion, index construction, and host-device traffic in the dominant "
        "loop. The implementation therefore optimizes the lifecycle of the complete CV state while "
        "retaining one objective evaluation per proposal."
    ),
    (
        "The backends share mathematical invariants but not source-level execution strategies. KNN "
        "uses the same metric and voting rules with different search engines. PLS-LDA uses the same "
        "SIMPLS deflation and latent LDA model while each accelerator maps the dominant operations to "
        "its own memory and matrix-compute model. This separation avoids benchmark-specific branches "
        "inside the classifier and makes CPU/accelerator parity an explicit test target."
    ),
    (
        "The CPU PLS-LDA fold path projects training rows directly into class latent sums and the "
        "latent Gram matrix, then projects and scores validation rows without retaining complete "
        "score matrices. The randomized SIMPLS refresh remains float32. Its norm reduction preserves "
        "the original float result when representable and uses double accumulation only when the "
        "squared norm exceeds float32 range; the same representation guard is used by Metal. This "
        "restores mathematically feasible fixed-component fits without adding a dataset threshold or "
        "changing the two-step power iteration."
    ),
    (
        "Guided evolution is deliberately implemented above the classifier layer. Previous held-out "
        "predictions determine sample/group proposals and the class-transition graph; the adaptive "
        "schedule determines how many groups may move; and each proposal receives one new CV pass. "
        "This is a common scaffold, not a classifier-independent state policy: KNN uses the diversity-"
        "guarded score directly, whereas PLS-LDA may add transition coarsening and a fragmentation "
        "subtraction before the shared cooling decision. The policy can change stochastic trajectories "
        "relative to historical KODAMA without changing the definition of a CV prediction."
    ),
]


CUDA_BACKEND_PARAGRAPHS = [
    (
        "The CUDA path is package-owned and uses only CUDA Toolkit libraries. Input matrices are "
        "materialized as contiguous float32 buffers. Exact KNN evaluates the requested metric "
        "directly, whereas IVF-Flat builds GPU k-means centroids and inverted lists and selects "
        "nprobe from an exact pilot query when recall tuning is enabled. The selected nlist, nprobe, "
        "pilot recall, and device id are returned with the result. Exact Lloyd assignment uses a "
        "bounded 256 MiB score workspace per independent M lane and batches rows when necessary; "
        "every row still evaluates every centroid with the same distance and tie rules. Assignments, "
        "centroid sums and counts, empty-cluster repair, list counts, prefix offsets, and row-ID "
        "scatter remain on CUDA. A move-only ResidentIVFIndex can retain the input, projection, "
        "centroids, offsets, and IDs for subsequent self or external-query calls without a hidden "
        "process-global cache."
    ),
    (
        "For PLS-LDA, class labels are encoded compactly and the centered X'Y cross-product is formed "
        "from class sums on the GPU, avoiding a dense n-by-K one-hot response. Float32 SIMPLS uses "
        "reusable thread-local device workspaces that grow to the required capacity; the unreachable "
        "legacy double-precision CPU/CUDA path was removed after the August 2026 audit. Fixed-order "
        "class accumulation is followed by separate column-sum and centering kernels, avoiding a "
        "read/write race in X'Y centering. Each stream owns a persistent cuBLAS workspace. CUDA computes "
        "class latent sums as (X'Y)'W and the latent second moment as W'(X'X)W, so training scores need "
        "not be materialized. Validation projection, pooled-covariance correction, Cholesky solve, and "
        "LDA scoring remain float32 operations, while fold caches retain train/validation layouts and "
        "training Gram matrices."
    ),
    (
        "The default rank-one randomized SIMPLS refresh uses two power iterations. CUDA dispatches "
        "the resulting rank-one products as matrix-vector operations and omits the algebraically "
        "redundant 1-by-1 block eigensolve. After the first component, the preceding weight is the "
        "defined warm start, so the implementation also omits generation of a random vector that "
        "would be overwritten before use; this removes launches without changing any SIMPLS value. "
        "The multi-vector refresh remains available under the same deflation equations. Latent LDA "
        "forms pooled covariance from score and class-sum "
        "statistics, retries the trace-normalized ridge sequence only after a failed Cholesky "
        "factorization, and copies back predicted codes rather than score matrices."
    ),
    (
        "The CUDA backend uploads the full graph once and keeps per-worker landmark/fold, label, "
        "projection, voting, SIMPLS, and LDA allocations alive across independent runs. Values that "
        "mathematically change are refreshed in place, and independent runs retain their seeds and "
        "proposal trajectories. Landmark selection, proposal decisions, and compact result collection "
        "remain host orchestrated, so the manuscript does not claim a fully device-side evolutionary "
        "state machine."
    ),
]


METAL_BACKEND_PARAGRAPHS = [
    (
        "The Apple backend is implemented in Objective-C++ on Metal rather than emulating CUDA. A "
        "persistent Metal state owns the process-cached device, command queues, and compute pipelines. "
        "Device discovery falls back from the system default to enumerated Metal devices, and independent "
        "M workers retain separate command queues and float32 workspaces, paralleling CUDA stream lanes. Native kernels "
        "implement exact KNN, projected IVF-Flat search, Lloyd k-means assignment and centroid "
        "updates, empty-cluster repair, and inverted-list count, prefix, and scatter operations. "
        "The same move-only ResidentIVFIndex contract retains Metal input and index buffers across "
        "search calls; self-search therefore does not upload the training matrix again."
    ),
    (
        "The Metal PLS-LDA path uses float32 class-label cross-products and the same SIMPLS component "
        "limit as CPU and CUDA. Metal Performance Shaders execute the dominant Xw and X't matrix "
        "products through MPSMatrixMultiplication. The same two-iteration incremental refresh and "
        "SIMPLS loading deflation are used on all three backends; the comparatively small response-"
        "space products and deflation updates remain host operations. This division is stated "
        "explicitly because a native Metal backend does not imply that every scalar operation is a "
        "GPU kernel. Metal materializes projected scores only in a worker-local device buffer, "
        "uses a label-aware kernel for class sums and MPS matrix multiplication for T'T, uploads the small fitted LDA "
        "model, and returns validation labels after on-device scoring."
    ),
    (
        "Metal is a first-class backend contract: it has named KNNCV, PLSLDACV, CoreKNN, CorePLSLDA, "
        "and KODAMAMatrix entry points, reports Metal in result metadata, and is covered by dedicated "
        "parity tests. Unsupported Metal operations fail explicitly instead of being relabeled CPU "
        "execution. The R integration layer accepts backend = 'metal' explicitly and links the Metal, "
        "MetalPerformanceShaders, and Foundation system frameworks without FAISS or cuVS link flags. "
        "This makes Apple GPU results auditable in the same manner as CUDA results."
    ),
    (
        "Metal also exposes native fuzzy/binary UMAP and FFT-grid openTSNE while retaining the "
        "package-owned graph and backend-native PCA initialization. Following the audited fastEmbedR "
        "source, CPU uses a CSR UMAP epoch schedule, CUDA uses an atomic COO/CSR epoch schedule, and "
        "Metal uses the clean row sampler with fixed-point atomic coordinate updates. openTSNE retains sparse "
        "attractive forces, FFT-grid repulsion, gains, momentum, and centering on Metal. These "
        "native paths are exposed consistently in the C++, R, and Python APIs."
    ),
]


BACKEND_CONTRACT_ROWS = [
    (
        "Internal precision",
        "Contiguous float32 host/device matrices and reusable CUDA buffers.",
        "Float32 shared MTLBuffers and MPSMatrix objects.",
    ),
    (
        "KNN",
        "Exact or recall-tuned IVF-Flat with device-built lists, returned metadata, and an explicit resident-index owner.",
        "Exact or recall-tuned IVF-Flat with device-built lists and the same resident-index lifetime.",
    ),
    (
        "PLS-LDA",
        "GPU label cross-products, SIMPLS workspace, projection, class summaries, and LDA scoring.",
        "Label-aware cross-product with MPS-accelerated X projections and shared SIMPLS/LDA semantics.",
    ),
    (
        "KODAMA evolution",
        "For a fixed classifier: same folds, constrained proposals, classifier-specific score, requested feasible components, and independent M runs.",
        "For a fixed classifier: same folds, constrained proposals, classifier-specific score, requested feasible components, and independent M runs.",
    ),
    (
        "Failure semantics",
        "Unavailable CUDA entry points throw; no CPU fallback is labeled CUDA.",
        "Unavailable Metal entry points throw; no CPU fallback is labeled Metal.",
    ),
]


API_ROWS = [
    (
        "Cross-validation",
        "KNNCV, PLSLDACV",
        "Fold assignment, prediction, fold accuracies, confusion matrices, timing, memory, and backend metadata.",
    ),
    (
        "Core optimization",
        "CoreKNN, CorePLSLDA",
        "Independent label-evolution runs that maximize KNN or PLS-LDA cross-validated accuracy.",
    ),
    (
        "Graph preparation",
        "KODAMA.graph / KODAMAGraph",
        "One reusable KNN graph plus backend-matched UMAP/openTSNE PCA starts, with matrix or external-handle R storage and without retaining the raw matrix.",
    ),
    (
        "KODAMA matrix",
        "KODAMA.matrix / KODAMAMatrix",
        "Landmark selection, splitting, M runs, optimized label vectors, KODAMA dissimilarity, and graph output.",
    ),
    (
        "Graph-input KODAMA",
        "KODAMA.matrix.graph / KODAMAMatrixFromGraph",
        "Runs the same KNN or PLS-LDA KODAMA optimizer from supplied neighbor indices/distances; KNN reuses the graph directly and PLS-LDA uses self-tuning Laplacian features when the data matrix is unavailable.",
    ),
    (
        "Principal components",
        "PCA / KODAMA.pca",
        "Float32 randomized PCA scores, loadings, singular values, explained variance, preprocessing vectors, backend identity, and runtime on CPU, CUDA, or Metal.",
    ),
    (
        "Visualization",
        "KODAMA.visualization",
        "Pinned fastEmbedR UMAP and openTSNE ports from standard or KODAMA-corrected float32 graphs; backend-matched raw-data PCA initialization is the default, fuzzy UMAP weighting is the default, and initialization provenance is returned.",
    ),
    (
        "Graph and clustering",
        "makeSNNGraph equivalent and random-walk clustering",
        "CPU, CUDA, or Metal neighbor-graph construction followed by an explicit CPU random-walk clustering kernel.",
    ),
    (
        "R compatibility utilities",
        "normalization, scaling, dinisurface, helicoid, spirals, swissroll",
        "Dependency-free R implementations preserve the established preprocessing signatures, returned fields, and seeded synthetic-manifold constructions; they are wrapper conveniences rather than C++ kernels.",
    ),
]


COMPLEXITY_ROWS = [
    (
        "KODAMA runs",
        "M independent runs, each with Tcycle proposal cycles plus one initial CV pass.",
        "Embarrassingly parallel over M; each cycle performs exactly one CV evaluation.",
    ),
    (
        "KNN core",
        "Neighbor search/build plus O(M Tcycle n_L k) voting on landmarks, where n_L is the landmark count and k is the KNN predictor size.",
        "Graph-input KNN removes the search/build term and reuses supplied indices and distances.",
    ),
    (
        "PLS-LDA core",
        "O(M Tcycle folds C_PLS(n_train, p, h)) for h requested components, plus latent-space LDA scoring.",
        "Float32 label-aware SIMPLS and fold buffers reduce constants but do not change the repeated-CV structure.",
    ),
    (
        "Graph correction",
        "O(M E) agreement counting for E retained graph edges, followed by row-wise neighbor sorting.",
        "Sparse by default; dense n by n dissimilarities are avoided for scalability.",
    ),
    (
        "Visualization",
        "Runs on the base or KODAMA-corrected sparse graph using UMAP/openTSNE helpers.",
        "Visualization time is reported separately because it is not part of the CV optimization objective.",
    ),
]


GRAPH_INPUT_VALIDATION_ROWS = [
    (
        "KNN graph input",
        "Uses supplied neighbor indices/distances directly in CoreKNNGraph_CPU/CUDA/METAL and KODAMAMatrixFromGraph; no neighbor index is rebuilt.",
        "Preserves graph-based KNN voting mathematics. CUDA and Metal upload the supplied graph once and perform fold masking, voting, and prediction with resident graph buffers.",
    ),
    (
        "PLS-LDA graph input with X",
        "When the original data matrix is supplied with the graph, PLS-LDA uses float32 X and the graph supplies initialization/projection support.",
        "Equivalent classifier path to data-input PLS-LDA; graph handling is an input-format extension.",
    ),
    (
        "PLS-LDA graph-only input",
        "When X is unavailable, a host kernel transforms the graph into self-tuning normalized Laplacian features before backend-selected CorePLSLDA.",
        "Exposed as an optional spectral surrogate, not as evidence that graph-only PLS-LDA matches data-input PLS-LDA in all datasets.",
    ),
    (
        "Validation boundary",
        "Wrapper tests exercise both KNN and PLS-LDA graph-input calls; benchmark comparisons must use the same M, Tcycle, seed, and supplied graph.",
        "The manuscript does not use graph-only PLS-LDA results to support data-input KODAMA performance claims.",
    ),
]


GRAPH_INPUT_CONTRACT_ROWS = [
    ("Raw X", "Raw float32 variables", "Data-native KNN", "Raw variables", "One"),
    ("Prepared P", "Self-tuning graph features", "Fixed supplied graph", "Laplacian surrogate", "Zero"),
    ("Prepared P + X", "Raw float32 variables", "Fixed supplied graph", "Raw variables", "Zero"),
    ("Bare G", "Self-tuning graph features", "Fixed supplied graph", "Laplacian surrogate", "Zero"),
]


GRAPH_API_SMOKE_ROWS = [
    ("KNN", "X", "1", "0.323", "0.954", "0.907"),
    ("KNN", "P + X", "0", "0.037", "0.965", "0.907"),
    ("KNN", "P", "0", "0.023", "0.983", "1.000"),
    ("KNN", "G", "0", "0.023", "0.983", "1.000"),
    ("PLS-LDA", "X", "1", "2.217", "0.936", "1.000"),
    ("PLS-LDA", "P + X", "0", "2.088", "0.936", "1.000"),
    ("PLS-LDA", "P", "0", "1.817", "0.947", "1.000"),
    ("PLS-LDA", "G", "0", "1.801", "0.947", "1.000"),
]


GRAPH_API_SMOKE_NOTE = (
    "MetRef CUDA smoke validation used n = 873, p = 375, k = 100, and M = T = 10. "
    "The prepared R object occupied 1.08 MB and contained no raw-data field, compared "
    "with 2.70 MB for the input matrix; one-time graph/PCA preparation required 0.342 "
    "wall seconds. Exact PLS-LDA agreement confirms that prepared P does not alter "
    "raw-feature classification when X is also supplied. The KNN difference is expected: "
    "P + X uses fixed-graph voting, whereas X uses data-native KNN within each landmark "
    "sample. P and bare G agree exactly because stored PCA starts affect visualization "
    "initialization, not graph-only label evolution. These figures validate API lifecycle "
    "semantics and are not presented as a performance benchmark."
)


GRAPH_PREPARATION_PSEUDOCODE_LINES = [
    "Algorithm: graph preparation and KODAMA.matrix dispatch",
    "",
    "KODAMA.Graph(X, k, backend):",
    "  X32 <- Float32(X)",
    "  G <- BackendKNN(X32, k)",
    "  (Z_U, Z_T) <- BackendPCAStarts(X32)",
    "  return P = {G, Z_U, Z_T, dimensions, backend, timings}",
    "  # X32 and X are not retained by P.",
    "",
    "KODAMA.Matrix(input, optional X):",
    "  raw X        -> prepare P; use raw geometry",
    "  prepared P+X -> reuse P; use caller-owned raw geometry",
    "  prepared P   -> reuse P; use self-tuning graph geometry",
    "  bare G       -> use G with no stored PCA starts and graph geometry",
]


GRAPH_INPUT_PROCEDURE_PARAGRAPHS = [
    (
        "The graph-input API represents a KNN matrix as two rectangular arrays with identical "
        "shape n by k. Entry I[i,j] is the sample identifier of the jth neighbor of sample i, "
        "and D[i,j] is its non-negative dissimilarity. The boundary accepts either zero-based "
        "or one-based identifiers, converts them to the internal convention, discards invalid "
        "identifiers, self-neighbors, and non-finite distances, clamps negative distances to "
        "zero, and retains at most graph.neighbors entries per row. The row order must match "
        "the sample order used by every optional label, constraint, and feature vector."
    ),
    (
        "For graph-only KNN KODAMA, the supplied graph replaces construction of the global "
        "neighbor index and remains the source of neighbors for CV voting, landmark-label "
        "projection, and final agreement correction. A deterministic self-tuning normalized-"
        "Laplacian representation is still computed once to provide geometry for landmark "
        "selection and initial splitting; it does not replace graph-based KNN voting. If the "
        "original data matrix X is available, KODAMAMatrixFromGraphData uses float32 X for "
        "that geometry while continuing to reuse the supplied graph."
    ),
    (
        "PLS-LDA requires a rectangular feature representation. With X supplied, graph-input "
        "PLS-LDA uses the ordinary float32 SIMPLS plus latent-space LDA classifier on X. With "
        "only I and D supplied, it uses the deterministic self-tuning normalized-Laplacian "
        "features described above. This graph-only PLS-LDA route is an explicit spectral "
        "surrogate and is not claimed to be numerically equivalent to data-input PLS-LDA. "
        "Graph normalization and spectral feature extraction are host operations in the "
        "current release; CPU, CUDA, or Metal is then selected for the PLS-LDA evolution."
    ),
    (
        "The remaining KODAMA procedure is unchanged: each of M independent runs selects "
        "landmarks, initializes at most splitting classes, performs Tcycle proposal and CV "
        "updates, and projects its optimized labels to all samples. The output contains the "
        "M by n label matrix, one best raw CV accuracy per run, one graph in either base or "
        "KODAMA-corrected state, timing fields including graph-feature time, and backend "
        "metadata. UMAP or openTSNE can then consume the corrected graph directly; "
        "when X is unavailable, visualization reports its graph-only initialization fallback."
    ),
    (
        "The R and Python graph-input functions expose the same contract. Python uses the "
        "direct snake_case transliteration of dotted R names. Defaults, accepted values, the "
        "one-based best-run convention, and returned fields are otherwise aligned."
    ),
]


GRAPH_INPUT_USAGE_LINES = [
    "C++17",
    "kodama::NeighborGraph g{I, D, k};",
    "kodama::KODAMAMatrixOptions opt;",
    "opt.runs = 100; opt.cycles = 100; opt.classifier = kodama::CoreClassifier::KNN;",
    "auto fit = kodama::KODAMAMatrixFromGraph(g, n, {}, {}, {}, opt);",
    "",
    "R",
    "p <- KODAMA.graph(X, k = 100, backend = \"cuda\")",
    "fit_graph <- KODAMA.matrix(graph = p, classifier = \"knn\")",
    "fit_graph_x <- KODAMA.matrix(data = X, graph = p, classifier = \"pls_lda\")",
    "fit_raw <- KODAMA.matrix(data = X, classifier = \"pls_lda\")",
    "fit_bare <- KODAMA.matrix(graph = list(indices = I, distances = D), classifier = \"knn\")",
    "",
    "Python",
    "p = kodama.graph(X, k=100, backend=\"cuda\")",
    "fit_graph = kodama.matrix(graph=p, classifier=\"knn\")",
    "fit_graph_x = kodama.matrix(data=X, graph=p, classifier=\"pls_lda\")",
    "fit_raw = kodama.matrix(data=X, classifier=\"pls_lda\")",
    "fit_bare = kodama.matrix(graph={\"indices\": I, \"distances\": D}, classifier=\"knn\")",
]


COMPATIBILITY_ROWS = [
    (
        "Objective",
        "KODAMA maximizes cross-validated accuracy of an evolving label vector.",
        "Held-out label predictability remains the classifier signal and raw accuracy is reported; the standard search ranks proposals with a disclosed label-only degeneracy guard.",
    ),
    (
        "Search dynamics",
        "The historical core evolves labels through its published stochastic accuracy-maximization procedure.",
        "Grouped adaptive moves and transition-driven coarsening are explicit search extensions; they use CV predictions but never reference labels.",
    ),
    (
        "Independent runs",
        "R documentation reports M iterative processes and Tcycle optimization cycles.",
        "KODAMAMatrixOptions exposes runs and cycles; runs remain independent across CPU, CUDA, and Metal backends.",
    ),
    (
        "Sample selection",
        "The documented FUN_SAM default selects 75% of samples per iterative process.",
        "Landmark and splitting controls preserve the same sampling role while exposing it in typed options.",
    ),
    (
        "Classifiers",
        "The R package documents KNN and PLS-DA classifiers.",
        "The KODAMA optimizer exposes KNN and SIMPLS + PLS-LDA; component count is explicit.",
    ),
    (
        "Output contract",
        "The R interface returns dissimilarity, accuracy, proximity, label vectors, scores, entropy, and landmarks.",
        "The C++ result returns optimized labels, accuracies, graphs, timings, memory, and backend metadata for wrappers.",
    ),
    (
        "Architecture",
        "The R package is built around R workflows and R dependencies.",
        "The core is R/Python independent, CMake-installable, and accelerated where available.",
    ),
]


RELATED_WORK_PARAGRAPHS = [
    (
        "KODAMA is related to, but distinct from, semi-supervised learning and weak supervision. "
        "Classical semi-supervised graph methods assume that at least some labels are observed and "
        "then propagate them under smoothness assumptions, as in Gaussian-field label propagation "
        "and local-global consistency. Pseudo-labeling and weak-supervision systems also create or "
        "aggregate imperfect labels, but they typically use those labels to train a downstream "
        "predictor. KODAMA reverses the emphasis: the label vector itself is the optimization "
        "object, and its quality is the ability of a chosen classifier to reproduce it under "
        "held-out folds."
    ),
    (
        "The closest evaluation literature is cluster validation by stability or prediction. "
        "Methods such as stability-based clustering and prediction strength ask whether a grouping "
        "is reproducible under perturbation or held-out prediction. KODAMA uses the same general "
        "principle of reproducibility, but internalizes it as the search objective over label "
        "vectors. This is why the benchmark reports the internal CV score together with adjusted "
        "Rand index, silhouette, local purity, and active class count rather than presenting CV "
        "accuracy alone as proof of external correctness."
    ),
    (
        "The manuscript also follows warnings from the model-selection and circular-analysis "
        "literature. Cross-validation can be biased when the same data are used both to choose a "
        "model and to estimate final performance; feature selection outside the resampling loop is "
        "a classic example. In kodama-cpp, the internal CV score is intentionally the search "
        "criterion, so external claims must come from separate diagnostics, transparent parameter "
        "sensitivity, and comparisons with standard UMAP and t-SNE on the same datasets."
    ),
]


RELATED_WORK_ROWS = [
    (
        "Semi-supervised graph learning",
        "Label propagation and local-global consistency start from observed labels and smooth them on a graph.",
        "KODAMA starts from synthetic or supplied candidate labels and optimizes their cross-validated predictability.",
    ),
    (
        "Pseudo-labeling and weak supervision",
        "Pseudo-labeling and data programming create imperfect labels to train predictive models.",
        "KODAMA treats the label vector as the object being discovered; the classifier is the measuring instrument.",
    ),
    (
        "Cluster stability and prediction strength",
        "Stability methods evaluate whether clusters survive perturbation or prediction on held-out samples.",
        "KODAMA turns held-out predictability into an optimization objective, then reports external diagnostics separately.",
    ),
    (
        "Manifold visualization",
        "UMAP and t-SNE optimize low-dimensional layouts from neighbor or affinity structures.",
        "KODAMA supplies a corrected graph to the same visualization step; it does not replace those embedding objectives.",
    ),
    (
        "Open-source ML software",
        "JMLR software papers emphasize clear APIs, reproducibility, tests, licensing, and benchmark transparency.",
        "kodama-cpp separates the numerical core from R/Python wrappers and records backend parameters and validation evidence.",
    ),
]


VALIDATION_ROWS = [
    (
        "Correctness",
        "CTest passed on dependency-free CPU, Apple Metal, and CUDA builds; tests cover constrained folds, predictions, confusion matrices, requested component counts, and backend parity.",
    ),
    (
        "Numerics",
        "Float32 MatrixView validation tests exercise CPU, CUDA, and Metal KNN/PLS-LDA paths with typed backend metadata; a fresh CUDA configure/build verifies that SIMPLS has no Armadillo or double-fallback dependency.",
    ),
    (
        "Performance",
        "Kernel-level KNNCV/PLSLDACV timings are separated from core optimizer and KODAMA.matrix timings in the benchmark outputs.",
    ),
    (
        "Continuous integration",
        "GitHub Actions commit a32f718 passed the CI, CPU-coverage, and API-documentation workflows, including Linux/macOS CPU, native Metal, and wrapper jobs.",
    ),
    (
        "Coverage and documentation",
        "LLVM 22.1.1 CPU instrumentation measured 80.48% line, 69.33% branch, 77.50% function, and 79.98% region coverage. Dedicated graph-construction and clustering contract tests raised graph_cluster.cpp to 94.65% line and 73.94% branch coverage. Cycles 19--22 added invalid-input, visualization-option, binary-UMAP compaction/replay, empty-resident-index, unavailable-backend, fixed-seed replay, monotone best-trace, guarded one-class, anti-collapse, shake-recovery, constant-predictor PLS-LDA fallback, and generic raw/graph matrix-orchestration contracts. Cycle 26 added resident-IVF null-view, non-positive-k, unavailable-backend, and move-ownership contracts. CUDA and Metal are excluded from that CPU number and retain hardware tests. Doxygen generation and GitHub Pages deployment passed, and the public API site is reachable at https://tkcaccia.github.io/kodama-cpp/.",
    ),
    (
        "Release",
        "CMake install targets, wrapper build scripts, SPDX source headers, a pinned provenance matrix, retained third-party license texts, and a checksum-backed license audit are present. On 2026-08-06 the public main branches were core 0b8b8736e4ef688e0fccb4c86057ef455de88760 and R wrapper 709059237266c741898d720e650bf3e38e92e7ec. Neither is a release tag, and the uploaded HPC manifests contain git_commit = NA; the final benchmark must therefore be rerun from a clean tagged archive rather than attributed retrospectively.",
    ),
]


INSTALLATION_ROWS = [
    (
        "C++ CPU",
        "Configure with CMake using KODAMA_ENABLE_CUDA=OFF and KODAMA_ENABLE_METAL=OFF; then build, test, and install the exported target.",
    ),
    (
        "C++ Metal",
        "On macOS, configure with KODAMA_ENABLE_METAL=ON while CUDA is off; then build and run CTest. Apple system frameworks are linked automatically.",
    ),
    (
        "C++ CUDA",
        "Activate an environment containing the CUDA toolkit; configure with KODAMA_ENABLE_CUDA=ON; then build and run CTest. FAISS, cuVS, RAFT, RMM, Armadillo, and external graph-clustering libraries are not required.",
    ),
    (
        "R wrapper",
        "Clone https://github.com/tkcaccia/kodama-r, set KODAMA_CPP_ROOT and KODAMA_CPP_BUILD_DIR, then run R CMD INSTALL kodama-r. The wrapper links the installed or selected core build and includes dependency-free KODAMA-compatible normalization, scaling, and synthetic-manifold utilities. Configure records the selected archive path and checksums of the archive and public headers; a changed signature rebuilds both Rcpp bridge objects and relinks the package.",
    ),
    (
        "Python wrapper",
        "Use split-repos/kodama-python from the reviewed core repository and pass KODAMA_CPP_ROOT and KODAMA_CPP_BUILD_DIR as CMake config settings to pip. Publication of the separately versioned Python repository remains a pre-submission action.",
    ),
    (
        "Runtime check",
        "Run ctest for the core library; in R call KODAMA.diagnostics(); in Python import kodama and run the package validation tests.",
    ),
]


LICENSE_DEPENDENCY_ROWS = [
    ("kodama-cpp project code", "MIT, with SPDX copyright assigned to Stefano Cacciatore; historical Git identity tkcaccia."),
    ("KODAMA-derived code", "Public origin GPL (>= 2); local portions whose copyright Stefano Cacciatore owns are separately authorized under MIT. KODAMA coauthors remain credited."),
    ("fastPLS-derived code", "Public origin GPL-3; local SIMPLS/LDA portions whose copyright Stefano Cacciatore owns are separately authorized under MIT. fastPLS coauthors remain credited."),
    ("fastEmbedR and faissR", "MIT snapshots pinned for randomized-PCA policy, UMAP/openTSNE, HNSW, Metal, and nearest-neighbor implementation lineage."),
    ("Native CPU HNSW", "Direct FAISS-derived organization under MIT; Meta and Stefano Cacciatore notices and the complete FAISS license are retained."),
    ("Apple Metal", "MIT AND Apache-2.0 for the combined file; Meta, Faiss-mlx, and modification notices plus both full licenses are retained."),
    ("Native CUDA", "Package-owned exact/IVF KNN, k-means, and SIMPLS/LDA use CUDA Toolkit libraries without FAISS, cuVS, RAFT, RMM, Armadillo, or external graph-clustering links."),
    ("Audit enforcement", "All tracked source-like files carry SPDX identifiers; official license texts are SHA-256 checked by tools/check_license_headers.sh."),
]


WRAPPER_VALIDATION_ROWS = [
    ("C++ core, local CPU", "Dependency-free release-candidate build passed 5/5 configured tests on macOS, including source/license, standalone-namespace, numerical/core, public-API, and float32 smoke validation."),
    ("C++ core, Apple Metal", "Native Metal release-candidate build passed 6/6 configured tests, including the source/license and standalone-namespace audits and frozen public-API snapshot; a clean external CMake consumer linked the installed package and selected Metal."),
    ("C++ core, CUDA", "Commit a32f718 passed 4/4 configured license, numerical/core, public-API, and float32 smoke tests in the CUDA 13.0 compute-capability-12.0 build on chiamaka."),
    ("R wrapper, local CPU", "R CMD build followed by R CMD check --as-cran --no-manual passed installation, documentation, examples, and tests in a UTF-8 locale; one environment-only NOTE reported macOS xcrun temporary detritus. Visualization metadata includes backend, optimizer, initialization provenance, runtime, and UMAP fuzzy-graph diagnostics."),
    ("Python wrapper, local CPU/Metal", "An isolated built wheel passed eight pytest checks, including CPU and physical-Metal UMAP/openTSNE provenance, optimizer, runtime, and fuzzy-graph metadata assertions."),
    ("GitHub Actions", "Commit a32f718 passed CI, CPU-coverage, and API-documentation workflows, including Ubuntu/macOS CPU, native Metal, R/Python, and documentation jobs."),
]


TCYCLE_SENSITIVITY_ROWS = [
    (
        "MetRef",
        "KNN",
        "0.9924 -> 1.0000",
        "0.8847 -> 0.8809",
        "7 -> 9",
        "3.116 -> 3.115",
    ),
    (
        "MetRef",
        "PLS-LDA",
        "0.9786 -> 1.0000",
        "0.9504 -> 0.9924",
        "34 -> 24",
        "119.790 -> 584.089",
    ),
    (
        "USPS",
        "KNN",
        "0.9662 -> 0.9997",
        "0.9414 -> 0.9925",
        "61 -> 32",
        "15.336 -> 15.828",
    ),
    (
        "USPS",
        "PLS-LDA",
        "0.9383 -> 0.9612",
        "0.9220 -> 0.9520",
        "88.5 -> 71",
        "139.135 -> 636.650",
    ),
]


M_SENSITIVITY_ROWS = [
    (
        "MetRef",
        "KNN",
        "1.0000 / 1.0000 / 1.0000",
        "0.9229 / 0.8931 / 0.8809",
        "0.0517 / 0.0517 / 0.0517",
        "0.464 / 1.235 / 3.115",
    ),
    (
        "MetRef",
        "PLS-LDA",
        "0.9985 / 1.0000 / 1.0000",
        "0.9901 / 0.9939 / 0.9924",
        "0.8257 / 0.8439 / 0.7700",
        "113.381 / 277.456 / 584.089",
    ),
    (
        "USPS",
        "KNN",
        "0.9976 / 0.9976 / 0.9997",
        "0.9947 / 0.9925 / 0.9925",
        "0.3370 / 0.3370 / 0.3139",
        "3.431 / 8.227 / 15.828",
    ),
    (
        "USPS",
        "PLS-LDA",
        "0.9639 / 0.9602 / 0.9612",
        "0.9534 / 0.9503 / 0.9520",
        "0.1098 / 0.1080 / 0.1156",
        "123.851 / 319.838 / 636.650",
    ),
]


ENSEMBLE_CONVERGENCE_ROWS = [
    ("MetRef", "KNN", "0.1377 / 0.0783 / 0.0407", "0.9888"),
    ("MetRef", "PLS-LDA", "0.0588 / 0.0453 / 0.0180", "0.9988"),
    ("USPS", "KNN", "0.1011 / 0.0617 / 0.0319", "0.9955"),
    ("USPS", "PLS-LDA", "0.1191 / 0.0803 / 0.0380", "0.9926"),
]


PARAMETER_SENSITIVITY_PARAGRAPHS = [
    (
        "Because KODAMA has two nested stochastic controls, M and Tcycle cannot be justified by visual quality alone. Tcycle controls the depth of label-vector evolution inside one run, whereas M controls how many independent local optima enter the final ensemble. The named release-validation experiment uses MetRef and USPS with both KODAMA classifiers: fixing M = 100 while varying Tcycle in {20, 50, 100}, and fixing Tcycle = 100 while varying M in {20, 50, 100}. It uses landmarks = 100000 subject to the historical 75% rule when n is no larger than landmarks, the default splitting rule, ncomp = 50, knn.k = 30, and a recorded seed."
    ),
    (
        "The Tcycle sweep reports best and median cross-validated accuracy, adjusted Rand index against reference labels used only after optimization, active-class count, and runtime. On USPS, increasing Tcycle from 20 to 100 raised best/median CV accuracy from 0.9662/0.9414 to 0.9997/0.9925 for KNN and from 0.9383/0.9220 to 0.9612/0.9520 for PLS-LDA; median class counts fell from 61 to 32 and from 88.5 to 71. On MetRef, PLS-LDA likewise improved its median CV accuracy from 0.9504 to 0.9924 and reduced median classes from 34 to 24, whereas KNN changed little and remained over-coarsened. Tcycle = 100 is therefore retained as a conservative search-depth setting, not as a claim of universal optimality; the MetRef PLS-LDA selected-run ARI was highest at Tcycle = 50."
    ),
    (
        "The M sweep has a different interpretation. Increasing M raises the number of independent solutions entering the final agreement graph; it is not a search-depth parameter and is not expected to improve the externally scored best run monotonically. Each edge weight is the fraction of runs assigning its two samples to the same label. Its Bernoulli worst-case Monte Carlo standard error is 1/(2 sqrt(M)): 0.1118 at M = 20, 0.0707 at M = 50, and 0.0500 at M = 100. Empirically, the graph-weight RMSE relative to the M = 100 ensemble fell between M = 20 and M = 50 for every dataset/classifier pair, reaching 0.0180-0.0407 at M = 50 with correlations of 0.9888-0.9988. M = 100 is retained to bound ensemble uncertainty, while the nonmonotone selected-run ARI is reported rather than used to tune M."
    ),
]


RELEASE_SENSITIVITY_READY = all(
    not row[0].startswith("Benchmark")
    for row in TCYCLE_SENSITIVITY_ROWS + M_SENSITIVITY_ROWS
)


PILOT_EXPERIMENT_PARAGRAPHS = [
    (
        "The benchmark corpus uses copied float32 RData datasets spanning n = 873 to 5,220,347 samples and p = 11 to 16,384 variables. For the current matched MetRef result, commit 9da48ee was rebuilt with CUDA 13.0 for compute capability 12.0; all three configured CTest targets and the installed R-wrapper testthat suite passed. Broader tables retain explicitly dated release-history snapshots rather than silently mixing them with this current-commit run."
    ),
    (
        "The current package-owned CUDA IVF-Flat path was checked on MNIST70k and MetRef with five-fold cosine KNNCV and k = 10. On MNIST70k it produced accuracy 0.973857 in 4.233 s with nlist = 237 and maximum nprobe = 32, compared with 0.973029 in 4.753 s for the recorded former FAISS/cuVS row. On MetRef it produced 0.816724 in 0.195 s, compared with 0.814433 in 0.297 s. These are implementation spot checks rather than repeated-run performance estimates."
    ),
    (
        "A focused resident-IVF microbenchmark separated construction from repeated search on a deterministic 50,000 by 32 float32 matrix using five timed searches after one warm-up. For all-row self-search, retaining the index reduced mean rebuild-plus-search time from 0.8695 to 0.8091 s on Metal and from 0.1984 to 0.1765 s on CUDA (1.07x and 1.12x). For 1,000-query batches against the same training matrix, time changed from 0.1179 to 0.0306 s on Metal and from 0.0318 to 0.0092 s on CUDA (3.86x and 3.45x). Neighbor overlap was 1.000 in every comparison. These are systems measurements, not dataset-level accuracy estimates."
    ),
    (
        "The current M = 100, Tcycle = 100 MetRef validation shows why cross-validated accuracy is reported together with label-quality diagnostics. Matched CPU and CUDA PLS-LDA runs both reached maximum CV accuracy 1.000 and median CV accuracy 0.992366. Their CV-selected runs contained 24/27 classes and attained ARI 0.7654/0.7732 against 22 withheld reference classes. Backend-native k = 30 fuzzy UMAP produced truth-label silhouette 0.8206/0.7856. Because finite-precision and scheduling differences can select different stochastic runs, these results support matched objective behavior rather than bitwise trajectory identity."
    ),
    (
        "The dependency-distilled CUDA build links CUDA Toolkit libraries but no FAISS, cuVS, RAFT, or RMM soname. Binary wrappers therefore need to locate the CUDA runtime selected at build time, but do not need to provision the removed search libraries."
    ),
    (
        "The exact-quota selector was also isolated on flow18 (1,000,021 rows, 11 variables). With a request exceeding n, the historical rule yielded 750,016 effective landmarks. Matrix-only sampling from 300 coarse strata required 0.321 seconds and represented all 300 strata. An earlier exact-search CUDA KNN run with M = 1 and Tcycle = 10 required 362.348 seconds, including 335.079 seconds for the shared 100-neighbor graph. This diagnosis motivated automatic accelerator graph dispatch. In the accepted resident implementation, graph time fell from the matched 296.482-second exact baseline to 25.938 seconds (11.43x). At M = 100, Tcycle = 100, and 1,000 landmarks, graph-returning KNN and PLS-LDA required 47.358 and 66.135 seconds; labels-only execution required 46.945 and 64.987 seconds because it omitted the final host graph download. Best CV accuracy remained 1.000 for both classifiers; PLS-LDA ARI changed from 0.926970 to 0.924996 between exact and IVF search and was unchanged by residency."
    ),
]


BENCHMARK_PROTOCOL_ROWS = [
    ("Date", "2026-07-27 UTC for the matched current-commit MetRef CPU/CUDA run; earlier validation snapshots are dated explicitly"),
    ("GPU", "NVIDIA GeForce RTX 5060 Ti, 16 GB device memory, driver 595.71.05"),
    ("Build validation", "Fresh CUDA 13.0 compute-capability-12.0 build at 9da48ee succeeded; CTest passed 3/3 and the installed R wrapper passed testthat"),
    ("Runtime", "CUDA Toolkit only for native search/PLS paths; no FAISS, cuVS, RAFT, RMM, or Armadillo link"),
    ("Data format", "RData lists exported to contiguous float32 row-major matrices"),
    ("CPU setting", "Four CPU threads for the matched MetRef end-to-end run; other tables state their worker count"),
    ("Timing boundary", "End-to-end wall time starts immediately before the public call and stops after synchronized return, including allocations and host/device transfers; no warm-up is removed unless a row is explicitly labeled microbenchmark"),
    ("CV kernels", "KNNCV and PLSLDACV measured independently from the full matrix pipeline"),
    ("Core optimizer", "Three seeds per dataset for CoreKNN and CorePLSLDA medians"),
    ("Reported metrics", "Wall time, peak memory where available, CV accuracy, ARI, and active class count"),
]


METAL_VALIDATION_PARAGRAPHS = [
    (
        "The dependency-light backend experiment was run on macOS with CUDA and OpenMP disabled. "
        "The CPU build therefore measures the package-owned HNSW and float32 SIMPLS/LDA implementations, "
        "while the Metal build links only Apple system frameworks. CPU/Metal agreement within the documented "
        "float32 tolerance is the primary acceptance criterion; timings are secondary evidence. A historical "
        "MetRef KNNCV diagnostic reporting 11.145 versus 0.026 seconds lacked a retained warm-up and "
        "synchronization record and has therefore been removed from the quantitative table rather than "
        "defended as a 425-fold speedup."
    ),
    (
        "A focused Apple M3 graph benchmark used 4,000 samples, 32 float32 variables, k = 30, "
        "and three repeated runs. Before concurrent insertion, the accepted implementation builds "
        "a deterministic connected seed containing min(n, max(n_threads + 1, 2 m_HNSW + 1)) "
        "points, where m_HNSW is the internal graph connectivity. Parallel HNSW construction "
        "required a median 6.722 s with one thread and 1.791 s with four threads, a 3.75-fold "
        "scaling gain. The earlier four-thread "
        "serial-insertion implementation required 7.248 s, so the accepted implementation was "
        "4.05-fold faster. Four-thread overlap with the one-thread graph had a three-run median "
        "of 0.99966. Across 100 repeated four-thread builds of a 600 by 24 test problem with "
        "k = 20, minimum and mean brute-force recall were 0.99975 and 0.99987. Both serial and "
        "parallel paths must retain at least 0.99 exact-neighbor recall in CTest, and the complete "
        "CPU suite passes ThreadSanitizer."
    ),
    (
        "The portable float32 distance loop is structured for compiler vectorization; Apple Clang "
        "emits 128-bit NEON arithmetic without a separate architecture path. On x86-64, a one-time "
        "per-index capability check selects a package-owned 256-bit AVX2/FMA kernel when supported "
        "and otherwise retains the portable loop. On an Intel Core i7-13700, five-run medians for "
        "the same 4,000 by 32, k = 30 benchmark decreased from 1.342554 to 1.253022 seconds with one "
        "thread and from 0.362764 to 0.337393 seconds with four threads. Candidate single-thread/four-thread "
        "graph overlap was 1.000 in all five runs. Explicit software prefetching was rejected after "
        "neutral-to-negative timings. HNSW candidate heaps were already reserved per worker, and final "
        "CPU graph correction now similarly reuses one sorting buffer per OpenMP worker; the latter "
        "preserved the exact result hash and had essentially neutral wall time."
    ),
    (
        "Metal IVF-Flat remains an explicit option rather than an automatic replacement for exact search. "
        "On MNIST10k, exact and recall-tuned IVF KNNCV both produced accuracy 0.9490, while elapsed time changed "
        "from 2.054 to 1.662 s. A 0.99 pilot-recall target was rejected because accuracy decreased to 0.9482; "
        "the accepted automatic pilot target is 0.999. End-to-end KODAMA can still favor exact search when "
        "repeated IVF training costs more than the saved query time."
    ),
    (
        "A 7 August 2026 local worktree screen provides preliminary crossover evidence, not release-tagged "
        "confirmatory results. One CPU graph per dataset and seed was supplied unchanged to CPU4 and Metal, "
        "separating optimizer timing from backend-specific graph search. Across seeds 4, 17, and 42 at "
        "M = 20 and Tcycle = 20, median Metal speedups were 1.79x for MetRef KNN, 1.91x for COIL20 KNN, "
        "2.02x for USPS KNN, and 2.33x for USPS PLS-LDA. MetRef PLS-LDA was effectively tied at 0.98x "
        "(25.842 seconds CPU4 versus 26.292 seconds Metal). Median Metal-minus-CPU truth-silhouette changes "
        "were +0.0035, +0.0009, -0.0002, +0.0036, and -0.0044 for those five cells, respectively. "
        "Raw-variable COIL20 PLS-LDA at 50 components was stopped after more than 600 seconds before one "
        "independent run completed; the paired Metal cell was not started. Both outcomes are retained as "
        "censored local feasibility evidence rather than replaced by PCA or a smaller component count. "
        "A separate bounded "
        "M=4/Tcycle=1 scheduler study found identical CV and ARI summaries for one through four Metal "
        "lanes, with optimization times 61.578, 38.685, 47.883, and 40.856 seconds, respectively. Because "
        "back-to-back GPU timing was sensitive to thermal and run order, the accepted automatic planner "
        "uses retained CV-fold bytes and the Metal working-set budget rather than choosing the fastest row. "
        "It selected three lanes for COIL20 and four for MetRef on the tested 8 GB Apple M3."
    ),
    (
        "A full-cycle follow-up used one unchanged CPU HNSW graph and the same CPU UMAP backend "
        "for CPU4 and physical Metal at M = 100 and Tcycle = 100. On USPS, Metal reduced KNN "
        "optimization from 12.451 to 7.888 seconds (1.58x) and pipeline time from 15.480 to "
        "10.972 seconds, while truth-label silhouette was 0.4550/0.4555. PLS-LDA showed the "
        "opposite systems result: CPU4 required 707.057 seconds and Metal 1105.998 seconds, "
        "with truth-label silhouette 0.4227/0.4174. Classic UMAP silhouette was approximately "
        "0.13. Together with full-cycle MetRef, this demonstrates a classifier- and matrix-shape "
        "crossover rather than universal accelerator speedup. The local worktree was dirty, so "
        "these rows are engineering controls awaiting clean tagged CPU4/CUDA replication."
    ),
    (
        "A third full-cycle control used ImageSegmentation (2,310 by 19; seven withheld truth "
        "classes), one unchanged CPU graph, CPU UMAP, seed 4, and M = Tcycle = 100. KNN "
        "optimization required 1.107 seconds on CPU4 and 1.672 seconds on Metal; truth-label "
        "silhouette was 0.5387 and 0.5396, versus 0.2935 and 0.2956 for classic UMAP. PLS-LDA "
        "required 12.557 seconds on CPU4 and 125.108 seconds on Metal; silhouette was 0.6418 "
        "and 0.6503. The requested 50 PLS "
        "components were correctly limited to the 19 mathematically feasible components. This "
        "negative accelerator result reinforces that backend advantage depends on matrix shape; "
        "Metal execution is validated but is not claimed to be universally faster."
    ),
    (
        "A fourth full-cycle control used PageBlocks (5,473 by 10; five strongly imbalanced truth "
        "classes), the same shared-graph design, seed 4, and M = Tcycle = 100. CPU4/Metal KNN "
        "optimization required 2.778/2.722 seconds and median CV accuracy was 0.99683/0.99732. "
        "CPU4/Metal PLS-LDA required 13.655/118.383 seconds while median CV accuracy remained "
        "close at 0.90512/0.90524. Classic and KODAMA truth silhouettes were all negative; the "
        "dataset is therefore retained as a backend-parity and adverse-quality control, not as "
        "evidence of visualization improvement. It also establishes a low-dimensional regime in "
        "which physical Metal PLS-LDA is substantially slower than CPU4."
    ),
    (
        "A fifth full-cycle control used PenDigits (10,992 by 16; ten withheld truth classes), "
        "the same shared-graph design, seed 4, and M = Tcycle = 100. CPU4/Metal KNN "
        "optimization required 4.397/4.925 seconds with median CV accuracy 0.99672/0.99588. "
        "CPU4/Metal PLS-LDA required 39.311/199.897 seconds with median CV accuracy "
        "0.97198/0.97350. In contrast to the high internal scores, classic UMAP truth-label "
        "silhouette was 0.5116 and every KODAMA silhouette was negative (-0.0832 to -0.1009). "
        "The diffuse plots are retained as an adverse ensemble-correction result: CV accuracy "
        "is not an external-label objective, and a fragmented label ensemble can weaken the "
        "corrected graph. No dataset-specific parameter change was accepted."
    ),
    (
        "A sixth local stress control used raw COIL20 (1,440 by 16,384; 20 withheld truth "
        "classes), seed 4, M = Tcycle = 100, and 1,080 effective landmarks under the unchanged "
        "75% rule. CPU4/Metal KNN optimization wall time was 193.377/80.753 seconds, a 2.39x "
        "Metal speedup, and median CV accuracy was 0.99537 on both backends. Complete pipeline "
        "time was 200.950/88.315 seconds. Both KODAMA truth silhouettes (0.4531/0.4278) were "
        "below classic UMAP (0.5266/0.5303), so this is acceleration and adverse-quality evidence, "
        "not a universal improvement claim. Raw CPU4 PLS-LDA exceeded 300 seconds before the "
        "first four independent runs completed and is retained as a censored resource-limit cell."
    ),
    (
        "A seventh local control used SatImage (6,435 by 36; six withheld truth classes), seed 4, "
        "M = Tcycle = 100, and 4,827 effective landmarks. The initial implementation required "
        "3.044/3.216 seconds for CPU4/Metal KNN and 56.540/331.781 seconds for PLS-LDA. A backend "
        "lifecycle audit then replaced one-slot Metal fold residency with a bounded 16-entry LRU "
        "and made independent-run concurrency memory-aware. With the same serialized CPU graph, "
        "seed, and mathematics, corrected CPU4/Metal times were 2.963/2.442 seconds for KNN and "
        "56.484/171.675 seconds for PLS-LDA. Metal therefore became 1.21x faster for KNN and "
        "1.93x faster than its own pre-fix PLS-LDA path, although low-dimensional PLS-LDA remained "
        "3.04x slower than CPU4. Median CV accuracy remained closely matched (KNN 0.99834/0.99772; "
        "PLS-LDA 0.94883/0.94893), no run collapsed, and PLS-LDA selected-label ARI was 0.493/0.581. "
        "KODAMA truth-label silhouettes remained below classic UMAP, so this is lifecycle and "
        "crossover evidence rather than a universal visualization-improvement claim."
    ),
    (
        "A subsequent fold-identity audit found that sequential Metal PLSLDACV calls without a persistent "
        "fold cache could reuse a temporary host address and incorrectly retain the previous fold matrix. "
        "Fold-specific worker-local residency epochs corrected the issue without changing SIMPLS or LDA. "
        "On raw COIL20 with five folds, 20 components, and 16,384 predictors, one-worker accuracy increased "
        "from the erroneous 0.555556 to 0.959028 and then matched four workers across seeds 4, 17, and 42. "
        "Median one/four-worker kernel times were 0.634/0.406 seconds. In a nested predictor-width study, "
        "median four-core CPU/Metal kernel speedups were 1.71x, 5.19x, 13.82x, and 10.78x at 256, 1,024, "
        "4,096, and 16,384 predictors. These results characterize the kernel crossover; they do not remove "
        "the much larger cost of a complete M-by-Tcycle high-dimensional KODAMA search. A later lifecycle "
        "audit extended the identity rule across calls: uncached stratified invocations now receive a fresh "
        "residency generation, while fixed non-stratified KODAMA folds retain stable cache generations. A "
        "sequential changed-label regression reproduces CPU fold assignments, predictions, and accuracy on "
        "physical Metal."
    ),
    (
        "A preserved-executable follow-up evaluated streamed CPU LDA sufficient statistics with five "
        "folds, 50 fixed components, four workers, seed 7, and three fresh processes. Predictions and "
        "accuracy were identical on MetRef, USPS, COIL20, and TabulaMuris. TabulaMuris (100,102 by 50) "
        "improved from 4.110 to 2.794 seconds (1.47x), and the reported peak-memory ratio was 0.587. "
        "Full MetRef KODAMA at M=5 and Tcycle=5 retained best CV 0.854962, ARI 0.590789, and 55 classes "
        "while median CPU time improved from 1.659 to 1.611 seconds."
    ),
    (
        "Full MNIST70k then exposed a float32 range failure in the randomized SIMPLS norm rather than a "
        "rank limitation: the former CPU path returned a one-component majority fallback at accuracy "
        "0.112529, and Metal failed. An overflow-only norm fallback restored all 50 requested components "
        "at accuracy 0.866457 on CPU and 0.866286 on Metal; median Metal time was 2.296 seconds across "
        "three fresh processes. MetRef, COIL20, and TabulaMuris accuracies were unchanged; Metal USPS "
        "improved by six predictions among 11,000 samples. A compact synthetic regression passes normal "
        "CPU tests, AddressSanitizer, and the physical Metal suite. These local worktree results require "
        "confirmation on the frozen CUDA tag before they become release evidence."
    ),
    (
        "A subsequent Metal PLS-LDA transfer audit retained MPS projection on-device, reduced training "
        "scores to class sums with a label-aware kernel, formed T'T with MPS matrix multiplication, and scored validation rows on Metal. "
        "Complete projected score matrices no longer return to the host; the small fitted LDA model is "
        "uploaded and only final integer labels return after validation scoring. The "
        "device score buffer, SIMPLS fit, LDA covariance, ridge sequence, priors, and component count are "
        "unchanged. A direct regression reproduces the former host calculation from the same projected "
        "scores. In three-process seed-7 comparisons, accuracy and 50 selected components were unchanged "
        "on five datasets. Metal PLSLDACV improved from 1.799 to 0.492 seconds on TabulaMuris (3.65x) and "
        "from 2.296 to 1.704 seconds on MNIST70k (1.35x), was 1.14x faster on USPS, neutral on raw COIL20, "
        "and 0.94x on tiny MetRef. A three-seed CPU4/Metal screen retained a maximum absolute accuracy "
        "difference of 0.001145. A follow-up replaced the serial-per-component-pair T'T kernel with the "
        "MPS product. Across three alternating-order seed pairs, median paired speedups were 1.14x on "
        "TabulaMuris and 1.10x on MNIST70k with 50 components and no PLSLDACV accuracy decrease; "
        "seven-repeat checks gave 1.13x and 1.01x. Tiny MetRef was 0.99x. These are local worktree "
        "engineering results, not frozen release claims."
    ),
    (
        "A later full-cycle Metal lifecycle audit identified two orchestration defects rather than a "
        "classifier limitation. A single-slot resident matrix cache repeatedly evicted the immutable "
        "five-fold matrices, and the independent-M scheduler was capped at four lanes irrespective of "
        "the device working-set budget. A bounded 16-entry least-recently-used fold cache and a "
        "device-budgeted scheduler of at most 32 lanes preserve the SIMPLS, LDA, proposal, and "
        "temperature equations. At M = Tcycle = 100, MetRef PLS-LDA improved from 631.415 to 312.029 "
        "seconds (26 selected lanes), versus 316.329 seconds on four CPU cores. Best CV accuracy "
        "1.000000, selected ARI 0.860185, 32 selected classes, median 25 classes, range 16--35, and "
        "zero collapse rate were unchanged. The rejected resident predictor-Gram experiment is not "
        "present in the implementation."
    ),
    (
        "Three in-process MetRef Metal graph calls required 0.136, 0.023, and 0.023 seconds wall "
        "time, showing that the small-data first-use difference is one-time pipeline compilation; "
        "the four-core CPU graph required 0.152 seconds."
    ),
    (
        "A standalone Metal UMAP follow-up supplied one identical CPU graph and PCA start to the CPU "
        "and Metal optimizers for 200 epochs. Warm Metal speedups were 1.08x on MetRef, 6.10x on USPS, "
        "and 6.52x on TabulaMuris; truth-label edge agreement@15 was 0.6225/0.6309, "
        "0.7077/0.7084, and 0.8828/0.8828 for CPU/Metal. Against a current fastEmbedR source snapshot, "
        "fuzzy-graph edge counts and maximum weights matched exactly on all three datasets. Physical C++, "
                "R, and Python Metal tests pass. A subsequent native Metal openTSNE validation used "
                "matched graphs and PCA starts on MetRef, USPS, and full TabulaMuris. KODAMA/fastEmbedR "
        "Metal openTSNE pair-distance correlations were 0.9863, 0.9981, and 0.9928, while KODAMA "
                "Metal/CPU correlations were 1.0000, 0.9997, and 0.9990. A later audit of fastEmbedR "
                "main at 5248ee0 found no changed native UMAP optimizer file and an identical SHA-256 "
                "for the complete exported float32 openTSNE function relative to the pinned parity "
                "snapshot; unrelated R-control, PCA, and graph-clustering additions were not copied. "
                "Warm KODAMA Metal times were "
                "0.054, 0.201, and 0.556 seconds. These are local worktree results."
    ),
    (
        "A full local MetRef comparison then used one shared CPU graph, seed 4, 655 effective landmarks, "
        "M = 100, Tcycle = 100, 50 requested PLS components, and identical CPU UMAP visualization. "
        "CPU4/Metal KNN optimization required 2.729/1.609 seconds and produced truth-label silhouettes "
        "0.3293/0.3333, but both selected two classes and had low ARI (0.0430/0.0472). CPU4/Metal "
        "PLS-LDA optimization initially required 335.734/651.767 seconds, while best CV accuracy was 1.000 for "
        "both, median CV accuracy was 0.993893 for both, and truth-label silhouette was 0.8598/0.8656 "
        "against 0.1029 for classic UMAP. A single-lane Metal follow-up was stopped after two runs "
        "required 46.8 seconds and was rejected. These values predate the general lifecycle correction "
        "described above; the corrected Metal PLS-LDA time is 312.029 seconds versus a matched "
        "316.329-second CPU4 reference. No dataset-specific scheduling rule or mathematical change "
        "was accepted."
    ),
]


METAL_VALIDATION_ROWS = [
    ("MetRef", "PLSLDACV, 50 components, five-run median", "0.790", "0.202", "3.90x", "accuracy 0.990836 / 0.989691"),
    ("MetRef", "KODAMA KNN, M=10 T=20", "21.709", "0.171", "127x", "CV 0.978626; ARI 0.093922; 15 classes"),
    ("MetRef", "KODAMA PLS-LDA, M=8 T=20", "14.930", "6.900", "2.16x", "best CV 0.963359 / 0.969466; ARI 0.741026 / 0.849732"),
    ("MetRef", "KODAMA KNN, M=100 T=100, shared graph", "2.729", "1.609", "1.70x", "best CV 1.000 / 1.000; two classes"),
    ("MetRef", "KODAMA PLS-LDA, M=100 T=100, corrected shared graph", "316.329", "312.029", "1.01x", "median CV 0.993893 / 0.993893; selected ARI 0.860185; no collapse"),
    ("USPS", "KODAMA KNN, M=100 T=100, shared graph", "12.451", "7.888", "1.58x", "median CV 0.99333 / 0.99345; sil. 0.4550 / 0.4555"),
    ("USPS", "KODAMA PLS-LDA, M=100 T=100, shared graph", "707.057", "1105.998", "0.64x", "median CV 0.95376 / 0.95370; sil. 0.4227 / 0.4174"),
    ("Image Seg.", "KODAMA KNN, M=100 T=100, shared graph", "1.107", "1.672", "0.66x", "median CV 0.99596 / 0.99481; sil. 0.5387 / 0.5396"),
    ("Image Seg.", "KODAMA PLS-LDA, M=100 T=100, shared graph", "12.557", "125.108", "0.10x", "median CV 0.97346 / 0.97259; sil. 0.6418 / 0.6503"),
    ("Page Blocks", "KODAMA KNN, M=100 T=100, shared graph", "2.778", "2.722", "1.02x", "median CV 0.99683 / 0.99732; sil. -0.1627 / -0.1802"),
    ("Page Blocks", "KODAMA PLS-LDA, M=100 T=100, shared graph", "13.655", "118.383", "0.12x", "median CV 0.90512 / 0.90524; sil. -0.1835 / -0.1694"),
    ("SatImage", "KODAMA KNN, M=100 T=100, corrected shared graph", "2.963", "2.442", "1.21x", "median CV 0.99834 / 0.99772; no collapse"),
    ("SatImage", "KODAMA PLS-LDA, M=100 T=100, corrected shared graph", "56.484", "171.675", "0.33x", "median CV 0.94883 / 0.94893; ARI 0.493 / 0.581"),
]


PILOT_DATASET_ROWS = [
    ("COIL20", "1,440", "16,384", "20"),
    ("FashionMNIST", "70,000", "784", "10"),
    ("FlowRepository", "5,220,347", "32", "13"),
    ("MNIST", "70,000", "784", "10"),
    ("Macosko2015_retina", "44,808", "50", "12"),
    ("MetRef", "873", "375", "22"),
    ("TabulaMuris", "70,118", "50", "56"),
    ("USPS", "11,000", "256", "10"),
    ("flow18", "1,000,021", "11", "16"),
    ("imagenet", "1,281,167", "1,024", "1,000"),
    ("mass41", "965,282", "14", "17"),
]


EVALUATION_GUARDRAIL_ROWS = [
    (
        "Circularity",
        "CV accuracy is the internal objective, not an external truth score.",
        "Report ARI, silhouette, local purity, active classes, and standard embedding baselines after optimization.",
    ),
    (
        "Overclaiming",
        "Present kodama-cpp as a KODAMA implementation and acceleration library, not a new universal learning paradigm.",
        "State methodological continuity with the 2014 and 2017 KODAMA publications and limit novelty claims to architecture and numerics.",
    ),
    (
        "Parameter dependence",
        "M, Tcycle, landmarks, splitting, KNN k, graph k, and PLS components can affect results.",
        "Use predefined sensitivity experiments and report parameter settings in every benchmark table and plot.",
    ),
    (
        "Visualization bias",
        "A visually appealing layout is not itself evidence of label quality.",
        "Run classic UMAP/openTSNE and KODAMA-corrected UMAP/openTSNE with the same implementation and record compactness metrics.",
    ),
    (
        "Runtime attribution",
        "End-to-end speedups can hide which kernel produced the benefit.",
        "Separate CV kernels, core optimizer, KODAMA.matrix, graph construction, and visualization timings.",
    ),
    (
        "Reproducibility",
        "Wrapper behavior must match the standalone core.",
        "Record CMake settings, backend metadata, R CMD check, pytest, CTest, seeds, and benchmark command lines.",
    ),
]


ABLATION_MATRIX_ROWS = [
    (
        "Search evolution",
        "Raw-accuracy proposals versus the standard grouped/guarded search",
        "Separates the current search extension from backend acceleration and from the unchanged CV classifiers.",
    ),
    (
        "Classifier",
        "KNN versus PLS-LDA",
        "Tests whether local-neighbor and latent-linear predictability give complementary label vectors.",
    ),
    (
        "Graph correction",
        "Standard graph versus KODAMA-corrected graph",
        "Isolates whether the optimized label ensemble improves downstream UMAP/openTSNE structure.",
    ),
    (
        "Search depth",
        "Tcycle = 20, 50, 100",
        "Measures whether additional proposal cycles improve CV accuracy or reduce fragmentation.",
    ),
    (
        "Ensemble size",
        "M = 20, 50, 100",
        "Measures whether more independent runs stabilize the final graph and best-run diagnostics.",
    ),
    (
        "Landmark and splitting controls",
        "Predeclared values, not per-dataset visual tuning",
        "Checks sensitivity of the initial label space and directly optimized sample subset.",
    ),
    (
        "Backend",
        "Single-core CPU, 4-core MultiCPU, CUDA, and Apple Metal where available",
        "Separates mathematical equivalence from scheduling and hardware acceleration.",
    ),
    (
        "Wrapper parity",
        "C++ core, R wrapper, Python wrapper",
        "Ensures language bindings are thin and do not introduce independent numerical behavior.",
    ),
]


PILOT_CV_ROWS = [
    ("COIL20", "KNNCV", "5.455", "1.435", "3.8x", "0.917/0.916", "-0.000695", "-"),
    ("MNIST", "KNNCV", "76.769", "4.753", "16.2x", "0.974/0.973", "-0.000928", "-"),
    ("MetRef", "KNNCV", "0.163", "0.297", "0.6x", "0.821/0.814", "-0.006873", "-"),
    ("MetRef", "PLSLDACV", "3.196", "0.118", "27.2x", "0.992/0.992", "0.000000", "50"),
    ("USPS", "KNNCV", "2.273", "0.429", "5.3x", "0.688/0.688", "-0.000819", "-"),
    ("USPS", "PLSLDACV", "3.760", "0.158", "23.8x", "0.684/0.687", "+0.002727", "50"),
    ("mass41", "KNNCV", "366.883", "3.692", "99.4x", "0.912/0.912", "-0.000066", "-"),
]


NATIVE_CUDA_VALIDATION_ROWS = [
    ("MNIST70k", "native IVF-Flat", "0.973857", "4.233", "237", "32", "0.994531"),
    ("MetRef", "native IVF-Flat", "0.816724", "0.195", "27", "20", "0.991406"),
]


PILOT_CORE_ROWS = [
    ("MetRef", "CoreKNN", "0.147", "0.122", "1.2x", "0.963/0.947", "0.271/0.313", "16/16"),
    ("MetRef", "CorePLSLDA", "16.597", "0.626", "26.5x", "0.885/0.838", "0.357/0.366", "53/48"),
    ("USPS", "CoreKNN", "2.282", "0.242", "9.4x", "0.969/0.968", "0.216/0.240", "59/53"),
    ("USPS", "CorePLSLDA", "20.189", "1.950", "10.4x", "0.938/0.910", "0.140/0.146", "80/77"),
]


KODAMA_BACKEND_COMPARISON_PARAGRAPHS = [
    (
        "The backend comparison is separated by scope. Core-optimizer rows measure the repeated "
        "CV label-search kernels without final graph construction, while the KODAMA.matrix row "
        "includes landmark selection, label search, projection, and graph correction on a small "
        "end-to-end validation dataset. This prevents kernel timings from being reported as full "
        "pipeline timings."
    ),
    (
        "The rows below therefore state their scope explicitly. CPU rows use the configured CPU "
        "backend, CUDA rows use the CUDA backend, and differences in CPU/CUDA accuracy are treated "
        "as reproducibility tolerances caused by stochastic search, float32 arithmetic, and approximate "
        "nearest-neighbor behavior rather than as evidence of different mathematical objectives."
    ),
]


KODAMA_BACKEND_COMPARISON_ROWS = [
    ("MetRef", "KNN", "CPU", "0.147", "0.963", "0.271", "16", "core optimizer"),
    ("MetRef", "KNN", "CUDA", "0.122", "0.947", "0.313", "16", "core optimizer"),
    ("MetRef", "PLS-LDA", "CPU", "16.597", "0.885", "0.357", "53", "core optimizer"),
    ("MetRef", "PLS-LDA", "CUDA", "0.626", "0.838", "0.366", "48", "core optimizer"),
    ("USPS", "KNN", "CPU", "2.282", "0.969", "0.216", "59", "core optimizer"),
    ("USPS", "KNN", "CUDA", "0.242", "0.968", "0.240", "53", "core optimizer"),
    ("USPS", "PLS-LDA", "CPU", "20.189", "0.938", "0.140", "80", "core optimizer"),
    ("USPS", "PLS-LDA", "CUDA", "1.950", "0.910", "0.146", "77", "core optimizer"),
]


PILOT_MATRIX_ROWS = [
    ("CPU PLS-LDA", "341.648", "341.612", "1.0000/0.9924", "0.7654/0.8206", "24", "351.4"),
    ("CUDA PLS-LDA", "270.408", "270.220", "1.0000/0.9924", "0.7732/0.7856", "27", "729.7"),
]


HISTORICAL_KNN_PARAGRAPHS = [
    (
        "A KODAMA 2.4.1/2.4 KNN-compatible predecessor experiment used MetRef, KNN, M = 100, Tcycle = 100, "
        "655 effective landmarks under the historical 75% rule, splitting = 50, k = 30, "
        "graph k = 100, and seed 1234. Splitting = 50 matches the internal k-means "
        "initialization in KODAMA R 2.4.1. KODAMA R 2.4.1/2.4 is the designated and sole "
        "predecessor performance baseline. KNN provides the classifier-matched systems comparison. "
        "Historical PLS-DA is also measured below as a contextual predecessor baseline, but it is not "
        "used to calculate a classifier-matched speedup against current SIMPLS plus latent-space LDA. "
        "These are single-run preliminary "
        "measurements from the retained 16 July archive, not final replicated release estimates."
    ),
    (
        "The current grouped and guarded proposals extend the historical stochastic search, "
        "so this is a matched-settings predecessor and systems comparison rather than label-trajectory "
        "parity. Current single-core execution was 1.58 times slower than KODAMA R 2.4.1 on this "
        "small case, whereas four-core execution was 2.59 times faster and CUDA was 258.6 times "
        "faster. Every row reached best raw CV accuracy 1.000; the different median scores and label "
        "diagnostics are reported rather than interpreted as backend equivalence. A frozen rerun will "
        "report replicated within-host medians and dispersion under the same settings."
    ),
]


HISTORICAL_KNN_ROWS = [
    ("KODAMA R 2.4.1", "CPU", "1", "610.543", "1.00x", "1.000/0.961", "0.000/0.054", "2/6"),
    ("kodama-cpp", "CPU", "1", "965.348", "0.63x", "1.000/0.874", "0.045/0.071", "2/8.5"),
    ("kodama-cpp", "CPU", "4", "235.547", "2.59x", "1.000/0.874", "0.045/0.071", "2/8.5"),
    ("kodama-cpp", "CUDA", "4", "2.361", "258.6x", "1.000/0.882", "0.052/0.069", "3/8.5"),
]


HISTORICAL_PLS_PARAGRAPHS = [
    (
        "KODAMA R 2.4.1 PLS-DA and current PLS-LDA are compared on the same MetRef data, "
        "M = 100, Tcycle = 100, 655 effective landmarks, graph k = 100, and fixed seeds. "
        "The historical row uses the archived FUN = 'PLS-DA' interface at its documented default "
        "of five components; current validation requests 50 components. The current "
        "route computes label-aware SIMPLS and then fits LDA in latent component space. Because the "
        "decoders and search implementations differ, elapsed times and CV traces are reported side by "
        "side as historical context; their ratio is not presented as an implementation-only speedup."
    ),
    (
        "On chiamaka, the historical five-component PLS-DA run required 965.780 seconds, reached "
        "best/median CV accuracy 1.000/0.9746, and its CV-selected vector had ARI 0.0195 with two "
        "active classes (median ARI 0.0576 and median four classes across runs). Independently retained "
        "current 50-component validations required 316.329 seconds on four CPU cores and 270.408 "
        "seconds on CUDA; their selected ARI values were 0.860 and 0.773. These observations show the "
        "practical difference between released routes, but they do not isolate classifier choice from "
        "search, component count, parallelism, or backend."
    ),
]


# Filled from the retained 2.4.1 run and independently retained current-backend
# validations. No speedup column is used because the evaluators are different.
HISTORICAL_PLS_ROWS = [
    ("KODAMA R 2.4.1", "PLS-DA", "CPU", "1", "5", "965.780", "1.000/0.975", "ARI 0.020/0.058; classes 2/4"),
    ("kodama-cpp", "SIMPLS + latent LDA", "CPU", "4", "50", "316.329", "1.000/0.994", "ARI 0.860; 32 classes"),
    ("kodama-cpp", "SIMPLS + latent LDA", "CUDA", "4", "50", "270.408", "1.000/0.992", "ARI 0.773; 27 classes"),
]


VISUALIZATION_COMPARISON_PARAGRAPHS = [
    (
        "The uploaded HPC archive enables a direct comparison between classic fastEmbedR UMAP or "
        "openTSNE and the corresponding visualization of a KODAMA-corrected graph. Both sides use "
        "the CUDA backend, seeds 4, 17, and 42, k = 30 for UMAP, and perplexity = 30 for openTSNE. "
        "KODAMA uses M = 100, Tcycle = 100, graph k = 100, and the historical effective-landmark "
        "rule. The comparison is paired by dataset, seed, and embedding."
    ),
    (
        "The archive spans 11 datasets from 873 to 5,220,347 observations and 11 to 16,384 input "
        "variables, covering object and clothing images, handwritten digits, DINOv2 image features, "
        "metabolomics, flow and mass cytometry, and single-cell transcriptomics. Complete three-seed "
        "classic/KODAMA pairs were available for 10 datasets per classifier: the KNN default was "
        "missing for FlowRepository and the PLS-LDA default was missing for ImageNet. No unmatched "
        "row contributes to a comparative value."
    ),
    (
        "The archived manifests retain dataset, classifier, backend, seed, M, Tcycle, ncomp, "
        "neighbor parameters, graph checkpoint, R version, and timestamp, but git_commit is NA. "
        "These results therefore support exploratory software evaluation but cannot be attributed "
        "to a uniquely identified source revision. The confirmatory release protocol requires a "
        "clean core tag, wrapper commit, source-archive SHA-256 digest, container digest, and a "
        "nonmissing commit field before execution."
    ),
    (
        "The dataset, rather than the seed, is the inferential unit. We first take the median of the "
        "three seed-paired silhouette changes within each dataset, then report the cross-dataset "
        "median, a 100,000-resample nonparametric bootstrap interval, an exact two-sided paired "
        "Wilcoxon test, a two-sided sign test, and Holm adjustment across the four classifier and "
        "embedding contrasts. Post-archive robustness analyses enumerate all dataset-level sign "
        "flips for the mean effect, report matched rank-biserial effect size, repeat the summary "
        "after removing MetRef, and recompute the median after leaving out each dataset in turn. "
        "These diagnostics are sensitivity analyses rather than additional confirmatory tests."
    ),
    (
        "PLS-LDA KODAMA followed by UMAP provided the clearest cross-dataset benefit: truth-label "
        "silhouette increased on 6 of 10 datasets, and the median dataset-level change was +0.025. "
        "The gains were +0.663 on MetRef, +0.102 on Tabula Muris, +0.084 on mass41, +0.038 on "
        "COIL20, +0.029 on FlowRepository, and +0.021 on flow18. PLS-LDA openTSNE improved 3 of 10 "
        "datasets, including +0.500 on MetRef. The PLS-LDA UMAP bootstrap 95% interval was "
        "[-0.067, +0.084] and its Holm-adjusted paired Wilcoxon p-value was 0.557, so this is an "
        "exploratory pattern rather than confirmatory evidence. Its leave-one-dataset-out medians "
        "nevertheless remained positive, from +0.021 to +0.029; excluding MetRef gave +0.021. "
        "The exact sign-flip mean-effect test was nonsignificant both with all datasets (p=0.637) "
        "and without MetRef (p=0.730), confirming that directional robustness does not establish "
        "statistical evidence."
    ),
    (
        "The benefit was not universal. KNN KODAMA improved UMAP silhouette on 2 of 10 datasets and "
        "openTSNE on none; their cross-dataset median changes were -0.094 and -0.084. PLS-LDA "
        "openTSNE also had a negative median change (-0.056). Trustworthiness and neighborhood "
        "preservation did not improve in aggregate for PLS-LDA, showing a trade-off between external "
        "class compactness and fidelity to the original local geometry. External labels are never used "
        "during KODAMA optimization, and silhouette is reported only afterward. Accordingly, the "
        "results support a classifier- and embedding-dependent capacity to expose additional "
        "label-aligned structure in some datasets, not a universal superiority claim. KNN "
        "openTSNE was consistently worse across all 10 datasets (Holm-adjusted p = 0.0078). "
        "Because this analysis was specified after the archive was complete, all inference is "
        "explicitly exploratory. A frozen confirmatory rerun will use PLS-LDA UMAP silhouette as "
        "the primary endpoint and trustworthiness and Preserve@30 as secondary endpoints."
    ),
]


AUDIT_IMPROVEMENT_ROWS = [
    ("Graph CV fallback", "Training-fold majority replaces any validation-label fallback", "Leakage regression tests"),
    ("Graph indexing", "Explicit zero/one-based metadata at public materialization boundaries", "C++ and wrapper API tests"),
    ("PLS-LDA numerics", "Removed unreachable double path; maintained SIMPLS-LDA is float32-only", "CPU/CUDA/Metal numerical tests"),
    ("CUDA PLS-LDA", "Class sums and pooled covariance use resident sufficient statistics", "CUDA CTest and accuracy parity"),
    ("CPU PLS-LDA", "Streams class sums and latent Gram statistics without full score matrices", "Four-dataset preserved-executable comparison"),
    ("Metal PLS-LDA", "Uses a label-aware class-sum kernel, MPS T'T, and on-device validation scoring", "Five-dataset preserved-executable and direct numerical tests"),
    ("SIMPLS range", "Overflow-safe CPU/Metal power norm preserves the finite float32 path", "Five-dataset and synthetic regression tests"),
    ("Graph residency", "Resident CUDA/Metal voting and one owned graph avoid host copies", "Lifecycle and backend tests"),
    ("Repeated search", "Fixed folds, backend-specific matrix/Gram storage, class buffers, and workspaces are reused", "M/Tcycle benchmark timings"),
    ("Wrapper conversion", "Blocked R transpose and contiguous Python graph export", "R CMD check and pytest"),
    ("Backend provenance", "Graph, optimization, and dissimilarity backends are reported separately", "No-fallback public API tests"),
]


HPC_RERUN_PLAN_ROWS = [
    ("Freeze", "Tag core and wrappers; record commits, source SHA-256, compiler/toolkit, container digest, GPU and CPU models"),
    ("Matrix", "All 11 archived datasets; KNN and PLS-LDA; seeds 4, 17, 42; M=Tcycle=100; no post hoc parameter substitution"),
    ("Backends", "CUDA for the complete matrix; CPU4 and Metal on a predeclared feasible subset with identical seeds and options"),
    ("Visualization", "Classic and KODAMA UMAP at k=30 and openTSNE at perplexity=30, using matched backend-specific starts"),
    ("Timing", "Graph, optimization, projection, correction, visualization, total wall time, peak host RAM, and peak device memory"),
    ("Quality", "CV accuracy and classes internally; silhouette, ARI, trustworthiness, Preserve@30, and local purity only afterward"),
    ("Failures", "Retain failed/time-limited cells explicitly; do not replace datasets, lower M/Tcycle, or alter landmarks silently"),
    ("Inference", "Dataset is the unit; frozen primary PLS-LDA UMAP silhouette, secondary trustworthiness and Preserve@30"),
]


def _fmt_metric(value: str, signed: bool = False) -> str:
    number = float(value)
    return f"{number:+.3f}" if signed else f"{number:.3f}"


def load_hpc_visualization_rows() -> list[tuple[str, ...]]:
    if not HPC_VISUALIZATION_SUMMARY.exists():
        return []
    rows: list[tuple[str, ...]] = []
    with HPC_VISUALIZATION_SUMMARY.open(newline="", encoding="utf-8") as handle:
        for row in csv.DictReader(handle):
            if int(float(row["matched_seeds"])) != 3:
                continue
            classifier = "KNN" if row["classifier"] == "knn" else "PLS-LDA"
            rows.append((
                row["dataset"].replace("FlowRepository_FR-FCM-ZYRM_files", "FlowRepository"),
                classifier,
                row["visualization"],
                f"{_fmt_metric(row['classic_silhouette_median'])} -> "
                f"{_fmt_metric(row['kodama_silhouette_median'])}",
                _fmt_metric(row["delta_silhouette_median"], signed=True),
                f"{_fmt_metric(row['classic_trustworthiness_median'])} -> "
                f"{_fmt_metric(row['kodama_trustworthiness_median'])}",
                f"{_fmt_metric(row['selected_ari_median'])} / "
                f"{float(row['selected_clusters_median']):.0f}",
                f"{float(row['kodama_total_sec_median']):.1f}",
            ))
    return rows


def load_hpc_dataset_rows() -> list[tuple[str, ...]]:
    if not HPC_DATASET_INVENTORY.exists():
        return []
    rows: list[tuple[str, ...]] = []
    with HPC_DATASET_INVENTORY.open(newline="", encoding="utf-8") as handle:
        for row in csv.DictReader(handle):
            rows.append((
                row["dataset"].replace("FlowRepository_FR-FCM-ZYRM_files", "FlowRepository"),
                f"{int(float(row['n'])):,}", f"{int(float(row['p'])):,}", row["domain"],
            ))
    return rows


def load_hpc_runtime_rows() -> list[tuple[str, ...]]:
    if not HPC_RUNTIME_SUMMARY.exists():
        return []
    values: dict[tuple[str, str], float] = {}
    with HPC_RUNTIME_SUMMARY.open(newline="", encoding="utf-8") as handle:
        for row in csv.DictReader(handle):
            dataset = row["dataset"].replace("FlowRepository_FR-FCM-ZYRM_files", "FlowRepository")
            values[(dataset, row["classifier"])] = float(row["kodama_core_sec"])
    rows: list[tuple[str, ...]] = []
    for dataset, *_ in HPC_DATASET_ROWS:
        knn = values.get((dataset, "knn"))
        pls = values.get((dataset, "pls_lda"))
        rows.append((
            dataset,
            "--" if knn is None else f"{knn:.1f}",
            "--" if pls is None else f"{pls:.1f}",
            "complete" if knn is not None and pls is not None else "one classifier missing",
        ))
    return rows


def load_hpc_inference_rows() -> list[tuple[str, ...]]:
    if not HPC_INFERENCE_SUMMARY.exists():
        return []
    rows: list[tuple[str, ...]] = []
    with HPC_INFERENCE_SUMMARY.open(newline="", encoding="utf-8") as handle:
        for row in csv.DictReader(handle):
            classifier = "KNN" if row["classifier"] == "knn" else "PLS-LDA"
            rows.append((
                classifier,
                row["visualization"],
                f"{float(row['positive_datasets']):.0f}/{int(float(row['complete_datasets']))}",
                _fmt_metric(row["median_delta_silhouette"], signed=True),
                f"[{_fmt_metric(row['bootstrap_ci_low'], signed=True)}, "
                f"{_fmt_metric(row['bootstrap_ci_high'], signed=True)}]",
                f"{float(row['wilcoxon_p']):.4f}",
                f"{float(row['wilcoxon_p_holm']):.4f}",
                f"{float(row['sign_test_p']):.4f}",
            ))
    return rows


VISUALIZATION_COMPARISON_ROWS = load_hpc_visualization_rows()
HPC_DATASET_ROWS = load_hpc_dataset_rows()
HPC_INFERENCE_ROWS = load_hpc_inference_rows()
HPC_RUNTIME_ROWS = load_hpc_runtime_rows()


KODAMA_PARAMETER_ROWS = [
    (
        "M / runs",
        "Number of independent KODAMA runs. KODAMAMatrixOptions::runs defaults to 100. Run r uses seed seed + r and writes one length-n label vector to the result matrix.",
    ),
    (
        "Tcycle / cycles",
        "Number of proposal/evaluation cycles inside each run. KODAMAMatrixOptions::cycles defaults to 20 for compatibility with the historical interface; the benchmark-quality setting used in this manuscript is 100. CoreOptions accepts the same role when the core optimizer is called directly.",
    ),
    (
        "landmarks",
        "Exact number of samples optimized directly in one run after validation. The default is 10000; if the requested value is at least n, the implementation uses ceil(0.75 n), then caps the value to [2, n-1]. Quota sampling always returns this effective count.",
    ),
    (
        "splitting",
        "Initial number of label classes and matrix-input landmark strata. If not supplied, splitting is 100 for n < 40000 and 300 otherwise. After landmark sampling, a separate matrix-space k-means initializes labels on the selected rows so sampling strata remain distinct from labels being optimized.",
    ),
    (
        "constrain",
        "Optional group vector. Empty means one movable group per sample. Non-empty constraints keep all members of a group in the same CV fold and make the group the unit of label proposal.",
    ),
    (
        "fixed",
        "Optional binary vector. Fixed samples are included in CV but are not relabeled by proposal moves.",
    ),
    (
        "graph_neighbors",
        "Number of neighbors retained in the final sparse graph. The actual k is min(graph_neighbors, landmarks, floor(0.75 n - 1)) after enforcing k >= 1.",
    ),
]


LABEL_SEARCH_DETAIL_PARAGRAPHS = [
    (
        "The object optimized by KODAMA is the label vector itself. At any time the algorithm holds a candidate vector y on the landmark samples. A classifier is not used to predict an external truth; instead it is trained to reproduce y under cross-validation. A label vector is therefore good when the labels assigned to held-out samples are predictable from the measured variables. The raw accuracy attached to a candidate, acc(y), is the fraction of landmark samples for which the held-out prediction pred_i(y) equals the candidate label y_i."
    ),
    (
        "The first CV pass is run on the initial label vector. This gives both an initial accuracy and a prediction vector. The prediction vector is then used as a proposal guide. If a sample or constrained group repeatedly receives a different predicted label from the classifier, that predicted label is evidence that the current label vector is not aligned with the structure that the classifier can reproduce. KODAMA uses this information to propose a new label vector rather than drawing completely blind random swaps."
    ),
    (
        "Within a cycle, the algorithm chooses a proposal base. In evolutionary mode this is the current accepted vector; otherwise it is the best vector found so far. For one-based cycle t, it sets p_t = t/(Tcycle + 1), s_t = p_t^2(3 - 2p_t), and q_max = 1 + floor((G - 1)(1 - s_t)), then samples q uniformly from 1,...,q_max. Early cycles can therefore alter many groups and later cycles fewer. For each sampled group, the replacement label is drawn from the empirical distribution of the previous CV predictions among its non-fixed members."
    ),
    (
        "The class-transition step aggregates the same information at label level. It builds a transition matrix N_ab counting samples whose current label is a and whose held-out prediction is b. Large off-diagonal counts indicate that the classifier consistently maps one candidate class into another. KODAMA may therefore propose merging unstable source labels into predicted destination labels, while rejecting moves that collapse the solution to a single class or do not reduce unstable fragmentation."
    ),
    (
        "The many-to-one transition proposal uses the same matrix N. For a source class a and a destination b, the expected count under independence is n_a n_b / n. A candidate absorption is considered only when N_ab exceeds this expectation and also exceeds the number of samples that remain in class a. Candidate destinations are sampled with probability proportional to their positive surplus, and the selected source classes are mapped to the destination only if the move reduces the number of active classes without collapsing all labels into one class."
    ),
    (
        "Under the standard KODAMA.matrix policy, the PLS-LDA path applies an additional automatic coarsening proposal before the CV evaluation. It computes the class entropy H = -sum_k p_k log p_k and the effective number of classes K_eff = exp(H). Fragmentation is measured as max(0, log(K / max(1, K_eff))). Classes with high transition instability and small movable size are proposed for merging into their strongest predicted destination; the merge budget is ceil(T_frag max(0, K - K_eff)), where T_frag = fragmentation / (fragmentation + weighted_transition_entropy + 1). This keeps the proposal tied to CV confusion rather than external labels. The branch is PLS-LDA-specific because latent LDA estimates a q-dimensional mean for every active class and a pooled covariance from the remaining degrees of freedom; many small classes therefore affect a global fitted model. KNN voting has no analogous global per-class parameter fit. The lower-level CoreOptions interface exposes the coarsening and diversity switches for controlled studies."
    ),
    (
        "After these proposal moves, the proposed vector is evaluated by exactly one new CV pass. The classifier-common score is S_0 = A sqrt(1 - sum_k p_k^2), where A is raw CV accuracy and p_k is the class proportion. KNN uses S_KNN = S_0. PLS-LDA uses S_PLS-LDA = S_0 - (1 - A) max(H / log n, fragmentation / (1 + fragmentation)) when automatic coarsening is active. The labels stored in clbest maximize the applicable classifier-specific score, and accbest is the raw A evaluated for exactly those labels. Thus the CV objective and evaluation count are common, but the state-selection policy is not identical."
    ),
    (
        "The best vector is updated whenever the proposal score improves. In evolutionary mode, a separate current vector is also maintained. A worse proposal can become current with probability exp((score_new - score_current) / tau_t), where, for one-based cycle t, tau_t = max(1e-9, 0.10 max(0, 1 - current_accuracy) (1 - t/Tcycle)). The 1e-9 floor makes the terminal cycle numerically defined and effectively greedy. Repeating this process for M independent seeds gives an ensemble of guarded high-predictability label vectors rather than relying on a single local optimum."
    ),
]


SCORE_INTERPRETATION_PARAGRAPHS = [
    (
        "It is useful to distinguish three quantities that appear in the implementation. The first is the raw held-out CV accuracy A(y), which is the classifier reproducibility score reported for a label vector. The second is the proposal acceptance score S(y), which is the stochastic-search objective actually used to choose moves when degeneracy guards are enabled. The third is an external diagnostic, such as ARI or embedding compactness, which is never optimized by KODAMA but is reported when reference labels are available."
    ),
    (
        "This distinction is necessary because a high raw CV accuracy alone can be achieved by a collapsed or overly coarse label vector. Such a vector may be easy for the classifier to reproduce but poor as a representation of latent structure. The diversity factor is common to KNN and PLS-LDA. Transition coarsening and the fragmentation subtraction are PLS-LDA-specific standard policies; they change which proposals are selected, while the reported accuracy remains the raw held-out classifier accuracy."
    ),
    (
        "The fixed-M = 100 sensitivity study provides empirical motivation, not a causal ablation. At Tcycle = 100, median active-class counts were 9 for KNN versus 24 for PLS-LDA on MetRef and 32 versus 71 on USPS, showing that the classifier paths occupy different fragmentation regimes despite high held-out accuracy. The archived experiment did not toggle only the PLS-LDA penalty, so the manuscript does not claim that this term universally improves ARI, silhouette, or runtime. It is reported as a classifier-specific implementation policy whose isolated benefit remains to be tested."
    ),
]


REPRODUCIBLE_ALGORITHM_STEPS = [
    (
        "Copy the input matrix once into float32 working storage. Before the M loop, build one full-data neighbor graph G0 for label projection and final agreement correction. When visualization initialization is requested, run one backend-native float32 PCA on the same working matrix and derive the separately scaled UMAP and openTSNE starts from those scores. The graph and both starts are returned with the matrix result."
    ),
    (
        "For each run r = 1,...,M, create a random generator with seed seed + r. Let L be the validated landmark count. Run k-means on all rows with splitting centers and 10 iterations to obtain coarse matrix-derived sampling strata."
    ),
    (
        "For each stratum c, compute q_c = L n_c / n, take floor(q_c) samples without replacement, and assign the remaining slots by randomized systematic rounding of q_c - floor(q_c). This returns exactly L distinct landmarks. If starting labels are supplied, subset them and apply constrained-majority relabeling when needed. Otherwise run matrix-space k-means on the landmarks with init_k = max(2, min(splitting, L)), independently of the sampling partition."
    ),
    (
        "Build CV folds on the landmark set. With no constrain vector, samples are assigned to folds directly. With a constrain vector, each group is assigned as a whole; stratified folds use the group-majority label and non-stratified folds shuffle groups. KODAMA.matrix currently sets the internal CV folds to non-stratified for both KNN and PLS-LDA core paths."
    ),
    (
        "At cycle t, choose the proposal base. In evolutionary-chain mode, the base label vector and prediction vector are the current accepted state; otherwise they are the current best state. These carried predictions drive the label proposal, and the proposed label vector is then evaluated by one CV pass."
    ),
    (
        "Sample proposal groups. Let G be the number of groups and index cycles by t = 1,...,Tcycle. With adaptive proposal size enabled, p_t = t/(Tcycle + 1), s_t = p_t^2 (3 - 2 p_t), theta_t = 1 - s_t, and q_max(t) = 1 + floor((G - 1) theta_t). Draw q uniformly from {1,...,q_max(t)} and sample q groups without replacement. Without adaptive proposal size, q is uniform on {1,...,G}."
    ),
    (
        "For each sampled group, collect the non-fixed members. Their candidate replacement labels are the previous CV predictions for those members. The replacement label is drawn from the empirical frequency of those predicted labels, and all eligible members of the group are relabeled together."
    ),
    (
        "Apply the classifier-common class-transition proposal. The transition matrix has entries N_ab = #{i: current label a, CV prediction b}. The many-to-one proposal requires N_ab - n_a n_b/n > 0 and N_ab > N_aa, prefers destinations supported by multiple source classes, and samples among eligible destinations in proportion to summed positive surplus. It rejects moves that collapse to one class or fail to reduce class count. Only for PLS-LDA, an additional coarsening move may merge small unstable classes according to the same transition matrix."
    ),
    (
        "Evaluate the proposed labels with exactly one CV pass. Let A_t be the fraction of landmark samples whose proposed label equals the held-out CV prediction. Both classifiers multiply A_t by sqrt(1 - sum_k p_k^2), where p_k is the current class proportion. Only the PLS-LDA policy subtracts the additional entropy/effective-class fragmentation term when class coarsening is enabled. Record raw A_t alongside the selected vector."
    ),
    (
        "Update the best vector whenever the proposal score improves. In evolutionary-chain mode, also update the current vector if the score improves the current score, or with probability exp((score - current_score) / tau_t), where tau_t = max(1e-9, 0.10 max(0, 1 - current_accuracy) (1 - t/Tcycle))."
    ),
    (
        "After Tcycle cycles, project the optimized landmark labels back to all samples. For the KNN path, each non-landmark receives the majority label among its labeled landmark neighbors in the global graph. For the PLS-LDA path, the final PLS-LDA model predicts the non-landmarks. If constraints are present, the final labels are replaced by the majority label inside each constraint group."
    ),
]


GRAPH_CONSTRUCTION_PARAGRAPHS = [
    (
        "The C++ implementation stores the final KODAMA representation as a sparse neighbor graph rather than materializing a dense n by n matrix by default. KODAMAGraph constructs one global KNN graph using the selected metric and graph_neighbors and computes backend-matched PCA starts. KODAMAGraphResult owns only indices, float32 distances, two initialization arrays, dimensions, backend provenance, and timings; it does not own, alias, or copy the source MatrixView. KODAMAMatrix can consume that result alone or alongside a separately supplied MatrixView; raw-matrix input invokes the same preparation logic internally. The graph owner is kept through all independent runs, compacted in place, and moved into the result. The default applies agreement correction to that same storage. If correction is deferred, the result retains the base graph and KODAMADissimilarityInPlace can correct it lazily; no simultaneous base_knn and corrected-knn copies are retained."
    ),
    (
        "For large CUDA and Metal inputs, graph preparation builds one resident IVF-Flat index and reuses it throughout KODAMA. Automatic nlist is bounded independently of nprobe and follows the matrix scale rather than the accelerator probe limit. Cosine inputs are normalized once in float32 before the resident upload, preserving the same metric contract as the nonresident path. A deterministic 128-query exact pilot increases nprobe until recall reaches 0.99 or the backend probe bound; retained candidates receive exact float32 distances. Self-search writes indices and distances directly into the resident graph buffers. KODAMA.matrix wrappers do not materialize those graph matrices by default. The R wrapper can retain one float32 graph behind an owning external pointer through optimization and visualization, or materialize conventional index and double-distance matrices with a cache-blocked row-major/column-major conversion. On flow18 (1,000,021 samples and 100 neighbors), the blocked conversion reduced matched median conversion time from 0.7359 to 0.5095 seconds; handle storage avoided the 1.20 GB R graph object until explicit materialization. Thus large graph arrays cross the wrapper boundary only when requested."
    ),
    (
        "Let c_i^(r) be the final label of sample i from run r. For each retained edge (i,j) with original nonnegative distance d_ij, compute V_ij as the number of runs where both endpoint labels are nonzero, and S_ij as the number of those valid runs where the two labels agree. If V_ij = 0 or S_ij = 0, the corrected edge distance is set to infinity. Otherwise a_ij = S_ij / V_ij and d'_ij = (1 + d_ij) / a_ij^2. Each neighbor row is then sorted by d'_ij. This implementation contract is not invariant to an arbitrary rescaling of d, so matched comparisons fix preprocessing and metric. The resulting graph is the KODAMA-corrected graph used by KODAMA.visualization."
    ),
    (
        "Agreement counting is backend-specific but mathematically identical. CPU transposes the M by n label ensemble once into sample-major contiguous storage before its parallel edge pass. At M = 100 and k = 100 this reduced matched four-thread correction time from 0.3170 to 0.1031 seconds at n = 20,000 and from 1.5520 to 0.6323 seconds at n = 100,000, with byte-identical output. Metal applies the same correction through a native compute kernel and is covered by CPU-distance parity tests. Host-transposed and warp-ballot CUDA variants were exact but slower in matched tests, and were therefore rejected in favor of the original run-major single-kernel CUDA path."
    ),
    (
        "A dense dissimilarity can be derived from the same agreement statistic by evaluating the formula for all sample pairs. The library keeps the sparse graph form because downstream UMAP, openTSNE, and graph clustering only need local neighborhoods, and the sparse form is the scalable object shared by the available backends."
    ),
]


MAIN_ALGORITHM_PSEUDOCODE_LINES = [
    "Algorithm: KODAMA with guided evolution and landmark projection",
    "",
    "Input: X or prepared graph P, classifier F, runs M, cycles T, landmarks L",
    "       splitting K0, folds Pi, seed s",
    "if P is absent then P <- KODAMA.Graph(X)",
    "G0 <- P.graph; (Z_U, Z_T) <- P.PCA_starts",
    "for r = 1,...,M do",
    "  strata <- CoarsePartition(X, K0, seed=s+r)",
    "  L_r <- ExactQuotaSample(strata, L, seed=s+r)",
    "  y <- InitialPartition(X[L_r,], K0, seed=s+r)",
    "  (A, pred) <- CrossValidate(F, X[L_r,], y, Pi_r)",
    "  (y_best, A_best, S_best) <- (y, A, Score_F(y, A))",
    "  for t = 1,...,T do",
    "    y_prop <- GuidedProposal(y, pred, t, T, constraints)",
    "    if F == PLS-LDA then y_prop <- TransitionCoarsen(y_prop, pred)",
    "    (A_prop, pred_prop) <- CrossValidate(F, X[L_r,], y_prop, Pi_r)",
    "    update (y_best, A_best, S_best) when Score_F(y_prop, A_prop) improves",
    "    (y, pred, A) <- AcceptOrCool(y_prop, pred_prop, A_prop, t, T)",
    "  end for",
    "  c[,r] <- ProjectLabels(F, X[L_r,], y_best, X)",
    "end for",
    "for each edge (i,j) in G0 do",
    "  V_ij <- {r : c[i,r] and c[j,r] are valid}; a_ij <- mean over r in V_ij of 1[c[i,r] = c[j,r]]",
    "  d_ij' <- infinity if a_ij = 0 else (1 + d_ij) / a_ij^2",
    "end for",
    "return corrected graph G, projected label matrix c, and paired raw CV accuracies",
]


SINGLE_RUN_PSEUDOCODE_LINES = [
    "Algorithm 1. One independent KODAMA run",
    "",
    "Input:",
    "  X_L: landmark matrix for this run",
    "  y0: initial label vector on landmarks",
    "  groups: movable units, either samples or constrained groups",
    "  fixed: optional fixed-label mask",
    "  folds: CV fold assignment",
    "  F: classifier in {KNN, PLS-LDA}",
    "  Tcycle: proposal/evaluation cycles",
    "",
    "Output:",
    "  y_best: optimized labels for this run",
    "  pred_best: CV predictions attached to y_best",
    "  acc_best: raw CV accuracy of y_best",
    "",
    "pred0 <- CrossValidate(F, X_L, y0, folds)",
    "acc0 <- Accuracy(y0, pred0)",
    "score0 <- ObjectiveScore(F, y0, acc0)",
    "y_best <- y0; pred_best <- pred0; acc_best <- acc0; score_best <- score0",
    "y_current <- y0; pred_current <- pred0; acc_current <- acc0; score_current <- score0",
    "",
    "for t = 1,...,Tcycle do",
    "  if evolutionary_search then",
    "    y_base <- y_current; pred_base <- pred_current",
    "  else",
    "    y_base <- y_best; pred_base <- pred_best",
    "  end if",
    "",
    "  y_prop <- y_base",
    "  qmax <- ProposalBudget(t, Tcycle, number_of_groups)",
    "  q <- UniformInteger(1, qmax)",
    "  selected_groups <- SampleWithoutReplacement(groups, q)",
    "",
    "  for each group g in selected_groups do",
    "    E <- {i in members(g): fixed[i] != 1}",
    "    if E is empty then continue",
    "    replacement <- SampleEmpirical(pred_base[E])",
    "    y_prop[E] <- replacement",
    "  end for",
    "",
    "  y_prop <- ClassTransitionProposals(y_prop, pred_base, fixed)",
    "  if F == PLS-LDA then",
    "    y_prop <- PLSLDATransitionCoarsening(y_prop, pred_base, fixed)",
    "  end if",
    "  pred_prop <- CrossValidate(F, X_L, y_prop, folds)  # exactly one CV pass",
    "  acc_prop <- Accuracy(y_prop, pred_prop)",
    "  score_prop <- ObjectiveScore(F, y_prop, acc_prop)",
    "",
    "  if score_prop > score_best then",
    "    y_best <- y_prop; pred_best <- pred_prop",
    "    acc_best <- acc_prop; score_best <- score_prop",
    "  end if",
    "",
    "  if evolutionary_search and",
    "     Accept(score_prop, score_current, acc_current, t, Tcycle, rng) then",
    "    y_current <- y_prop; pred_current <- pred_prop",
    "    acc_current <- acc_prop; score_current <- score_prop",
    "  end if",
    "",
    "  if acc_prop == 1 and (!guarded_diversity or score_prop >= score_best) then break",
    "end for",
    "",
    "return y_best, pred_best, acc_best",
]


KODAMA_PSEUDOCODE_LINES = [
    "Algorithm 2. KODAMA.matrix ensemble",
    "",
    "Input:",
    "  X: n by p data matrix",
    "  F: classifier in {KNN, PLS-LDA}",
    "  M: number of independent runs",
    "  Tcycle: proposal/evaluation cycles per run",
    "  landmarks, splitting, graph_neighbors",
    "  optional starting_labels, constrain, fixed",
    "",
    "Output:",
    "  C: M by n matrix of optimized label vectors",
    "  A: length-M vector of best raw CV accuracies",
    "  G_K: KODAMA-corrected neighbor graph",
    "  Z_umap, Z_tsne: optional backend-native PCA starts",
    "",
    "X32 <- float32(X)",
    "L_eff <- ValidateLandmarks(landmarks, n)  # ceil(0.75*n) when request >= n",
    "G0 <- KNNGraph(X32, graph_neighbors)  # exactly once, before M",
    "Z_umap, Z_tsne <- VisualizationPCAStarts(X32)  # one PCA, if requested",
    "for r = 1,...,M do",
    "  rng <- RNG(seed + r)",
    "  coarse <- KMeans(X32, splitting, rng)",
    "  L <- QuotaLandmarks(coarse, L_eff, rng)",
    "  if starting_labels supplied then",
    "    y0 <- starting_labels[L]",
    "  else",
    "    y0 <- ExpressionKMeans(X32[L], splitting, rng)",
    "  end if",
    "  groups <- ConstrainedGroups(constrain[L])",
    "            or one singleton group per landmark",
    "  y0 <- ConstrainedMajority(y0, groups)",
    "  folds <- AssignCVFolds(groups, y0)",
    "",
    "  y_best, pred_best, acc_best <- OneIndependentRun(",
    "      X32[L], y0, groups, fixed[L], folds, F, Tcycle, rng)",
    "",
    "  C[r, ] <- ProjectLabelsToAllSamples(y_best, X32, L, F, groups, G0)",
    "  A[r] <- acc_best",
    "end for",
    "",
    "G_K <- ReweightGraphByLabelAgreementInPlace(G0, C)",
    "return C, A, G_K, Z_umap, Z_tsne",
    "",
    "QuotaLandmarks(strata, L, rng, ordered):",
    "  for each stratum c do",
    "    q[c] <- L * size(c) / n",
    "    take[c] <- floor(q[c])",
    "    frac[c] <- q[c] - take[c]",
    "  end for",
    "  R <- L - sum_c take[c]",
    "  u <- Uniform(0, 1)",
    "  SystematicRound(frac, R, u, ordered, take)",
    "  return union_c SampleWithoutReplacement(c, take[c], rng)",
    "",
    "ObjectiveScore(F, y, a):",
    "  score <- a",
    "  if guarded_diversity then",
    "    p_l <- class proportions in y",
    "    score <- score * sqrt(1 - sum_l p_l^2)",
    "  end if",
    "  if F == PLS-LDA and class_coarsening then",
    "    score <- score - (1 - a) * Parsimony(y)",
    "  end if",
    "  return score",
    "",
    "Accept(s_new, s_old, a_old, t, Tcycle, rng):",
    "  if s_new >= s_old then return TRUE",
    "  tau <- max(1e-9, 0.10 * max(0, 1 - a_old) * (1 - t / Tcycle))",
    "  return Uniform(0, 1, rng) < exp((s_new - s_old) / tau)",
]


GRAPH_INPUT_PSEUDOCODE_LINES = [
    "Algorithm 3. KODAMA.matrix.graph from supplied neighbors",
    "",
    "Input:",
    "  I, D: n by k neighbor indices and distances",
    "  optional X: original n by p data matrix",
    "  F: classifier in {KNN, PLS-LDA}",
    "  M, Tcycle, landmarks, splitting",
    "  optional starting_labels, constrain, fixed",
    "",
    "Output:",
    "  C: M by n matrix of optimized label vectors",
    "  A: length-M vector of best raw CV accuracies",
    "  G_K: KODAMA-corrected graph on the supplied neighborhoods",
    "",
    "G0 <- NormalizeExternalGraph(I, D)",
    "if X is supplied then",
    "  X_work <- float32(X)",
    "else",
    "  X_work <- SelfTuningLaplacianFeatures(G0, ncomp)",
    "end if",
    "# For KNN, X_work supplies selection/initialization geometry only;",
    "# graph voting, projection, and correction continue to use G0.",
    "",
    "for r = 1,...,M do",
    "  rng <- RNG(seed + r)",
    "  L <- SelectLandmarksFromGraphOrData(G0, X_work, landmarks, rng)",
    "  groups <- ConstrainedGroups(constrain[L])",
    "            or one singleton group per landmark",
    "  y0 <- InitialLabelsFromGraphOrData(G0[L], X_work[L], starting_labels[L],",
    "                                    splitting, groups, rng)",
    "  folds <- AssignCVFolds(groups, y0)",
    "",
    "  if F == KNN then",
    "    y_best, pred_best, acc_best <- Algorithm1(G0[L], y0, groups, fixed[L], folds, KNN, Tcycle)",
    "  else",
    "    y_best, pred_best, acc_best <- Algorithm1(X_work[L], y0, groups, fixed[L], folds, PLS-LDA, Tcycle)",
    "  end if",
    "",
    "  C[r, ] <- ProjectLabelsToAllSamples(y_best, X_work, L, F, groups, G0)",
    "  A[r] <- acc_best",
    "end for",
    "",
    "G_K <- ReweightGraphByLabelAgreementInPlace(G0, C)",
    "return C, A, G_K",
]


IMPLEMENTATION_EVIDENCE_ROWS = [
    (
        "CPU SIMD k-means",
        "src/kodama_matrix.cpp evaluates package-owned float32 k-means norms and point-centroid products with four-way unrolled AArch64 NEON or x86-64 SSE2 kernels and a portable scalar fallback. Seeds, initial centers, centroid updates, tie-breaking, iteration count, and landmark quotas are unchanged.",
        "The full CPU and public-API suites pass after the change; native Metal smoke tests pass independently. The matched end-to-end benchmark records CV accuracy, ARI, active classes, and silhouettes rather than treating runtime alone as acceptance evidence.",
        "On COIL20 (1,440 by 16,384), four-core KODAMA-KNN at M=20/Tcycle=20 decreased from 188.579 to 42.421 seconds (4.45x) for seed 4. Across paired seeds 4, 17, and 42, median speedup was 4.42x and median SIMD-minus-scalar changes were 0.00000 for best CV, +0.00324 for median CV, +0.00559 for best ARI, -0.00060 for truth-label UMAP silhouette, and +0.00365 for label silhouette. SIMD reassociates float32 sums, so trajectory identity is not claimed; paired distributions, rather than bitwise labels, are the validation target.",
    ),
    (
        "Float32 numerical paths",
        "MatrixView float overload in include/kodama/kodama.hpp; DenseF PLS-LDA workspaces in src/plscv.cpp; CUDA and Metal device buffers remain float32. CPU streams latent LDA statistics, and CPU/Metal use an overflow-only double norm reduction while retaining float32 model state.",
        "tests/test_cv.cpp and tests/test_metal.cpp check float32 KNNCV/PLSLDACV outputs, backend metadata, requested components, and a large-finite-power-vector regression; normal CPU, AddressSanitizer, and physical Metal suites pass.",
        "On MNIST70k, CPU changed from an erroneous one-component fallback at 0.112529 accuracy to 50 components at 0.866457; Metal changed from failure to 50 components at 0.866286. Frozen CUDA confirmation remains required.",
    ),
    (
        "Package-owned CPU HNSW",
        "src/native_knn.cpp implements HNSW over contiguous float32 vectors and flat adjacency arrays. It completes a deterministic connected base seed of min(n, max(n_threads + 1, 2 m_HNSW + 1)) points before per-node-locked parallel insertion, cached reciprocal-edge distances, reusable batched candidate workspaces, and batched parallel queries. KNNCV, CoreKNN, graph construction, and KODAMA.matrix use it without linking FAISS.",
        "tests/test_cv.cpp requires valid ordered self-excluding graphs and at least 0.99 brute-force recall for both one-thread and four-thread builds. A 100-build stress test records minimum exact recall, while dependency-free CTest covers folds, prediction accuracy, graph construction, and KODAMA optimization; the full CPU suite also passes ThreadSanitizer.",
        "On a 4,000 by 32 Apple M3 benchmark with k = 30, the three-run median was 6.722 s with one thread and 1.791 s with four (3.75x scaling), versus 7.248 s for the previous four-thread path (4.05x improvement). Median one/four-thread graph overlap was 0.99966. Across 100 repeated 600 by 24 builds with k = 20, minimum and mean recall versus brute force were 0.99975 and 0.99987. On MetRef, native CPU KNNCV reproduced accuracy 0.827033.",
    ),
    (
        "Single-owner graph lifecycle",
        "src/kodama_matrix.cpp retains global_graph through the independent runs, compacts and optionally fuses rows in place, applies agreement correction in the same buffers, and moves the graph into KODAMAMatrixResult. The result no longer owns base_knn beside knn; R and Python serialize one graph.",
        "tests/test_cv.cpp compares lazy in-place correction with the edge-agreement formula, verifies stable index and distance storage addresses, and checks storage accounting. Wrapper tests require one knn payload and correction-state metadata.",
        "For the flow18 dimensions (1,000,021 samples, k = 100), the lifecycle benchmark reduced retained graph capacity from 2,408,050,568 to 808,016,968 bytes and maximum RSS from 2,354,304 to 791,808 KiB (66.4%). The graph ownership/copy phase fell from 0.561 to 0.059 s (89.4%); complete process wall time, including graph creation and a fixed RSS observation hold, fell from 1.60 to 1.07 s.",
    ),
    (
        "Final agreement correction",
        "src/kodama_matrix.cpp uses one parallel sample-major CPU label layout; src/metal_backend.mm exposes a native Metal in-place correction; src/kodama_matrix_cuda.cu retains its measured-faster run-major fused correction and row sort.",
        "tests/test_cv.cpp checks the sparse edge-agreement formula and CPU/CUDA parity; tests/test_metal.cpp checks CPU/Metal distance parity and guards against stale matrix-buffer reuse across operations.",
        "At M=100 and k=100, the accepted CPU layout was 3.08x faster at n=20,000 and 2.45x faster at n=100,000 with identical hashes. Exact CUDA transpose and warp-ballot experiments were rejected because they were slower. Matched rows are retained in the reproducibility archive.",
    ),
    (
        "CUDA nearest-neighbor search",
        "src/native_cuda_backend.cu implements package-owned float32 exact KNN, signed-hash IVF-Flat, GPU k-means, resident inverted lists, exact-pilot recall tuning, and bounded row-batched exact k-means assignment. src/kodama_matrix.cpp selects exact or IVF full-data graph construction from matrix scale and reports the selected index metadata.",
        "CUDA tests explicitly exercise exact and IVF index types, tuning metadata, constrained folds, float32 CoreKNN, and end-to-end KODAMAMatrix_CUDA initialization. Graph tests check explicit IVF dispatch and returned nlist, nprobe, and pilot recall. A clean build and ldd audit exclude FAISS/cuVS/RAFT/RMM.",
        "On flow18 (1,000,021 by 11, k=100), graph time fell from 296.482 s with exact search to 25.938 s with resident IVF (11.43x). At M=100/Tcycle=100, the graph-returning KNN/PLS-LDA pipelines required 47.358/66.135 s and labels-only execution required 46.945/64.987 s. Best CV accuracy remained 1.000 and PLS-LDA ARI changed by -0.001974 between exact and IVF search, not because of residency.",
    ),
    (
        "Resident accelerator IVF lifecycle",
        "src/resident_ivf.cpp exposes a move-only ResidentIVFIndex. CUDA and Metal retain the float32 input, projection, centroids, list offsets, IDs, and graph buffers. Automatic nlist is independent of the bounded nprobe search; a 128-query exact pilot targets recall 0.99. Self-search writes exact candidate distances directly to resident output buffers.",
        "tests/test_cv.cpp and tests/test_metal.cpp build one resident index, run repeated self-searches, verify self-exclusion and metadata, reject null non-empty views and non-positive k, and exercise move construction/assignment on physical Metal. A CUDA regression requests more than 256 lists to keep nlist independent of the probe cap. Wrapper tests cover labels-only execution without graph serialization, and tests/test_public_api_0_1.cpp compile-links the ownership and search surface.",
        "On a deterministic 50,000 by 32 systems microbenchmark, 1,000-query reuse was 3.45x faster than rebuild-plus-search on CUDA and 3.86x faster on Metal, with neighbor overlap 1.000. On flow18, the direct resident graph required 25.938 s; labels-only KNN/PLS-LDA saved the final download and required 46.945/64.987 s end to end.",
    ),
    (
        "Label-aware SIMPLS PLS-LDA",
        "src/plscv.cpp uses label/class inputs in CPU, CUDA, and Metal float32 SIMPLS paths; accelerator implementations avoid dense one-hot responses where possible.",
        "CPU/CUDA/Metal tests check PLSLDACV sizes, constrained folds, selected component reporting, backend metadata, and accuracy thresholds.",
        "At audited core commit 9da48ee, the matched MetRef M = 100, Tcycle = 100 PLS-LDA run required 341.648 s on four CPU cores and 270.408 s on CUDA, a 1.26x end-to-end speedup with identical median raw CV accuracy 0.992366; peak host-process memory was 351.4/729.7 MB. In a five-run fixed-50-component kernel check, median PLSLDACV time was 0.790 s on four Mac CPU threads and 0.202 s on Metal (3.90x), with accuracy 0.990836/0.989691. On chiamaka, corresponding kernel medians were 1.294 s on four CPU threads and 0.308 s on CUDA (4.20x), with accuracy differing by at most one sample.",
    ),
    (
        "Reusable fold and data buffers",
        "PLSFoldXCacheF in src/plscv.cpp caches fold assignments, train/validation indices, and scaled matrices. CUDA forms and reuses each fixed fold's float32 X'X; Metal instead retains fold matrices and persistent MPS buffers and encodes paired Xw and X'Xw products without allocating a predictor Gram matrix. Metal assigns a worker-local epoch to each fold so allocator address reuse cannot select a stale matrix. CUDA orders label uploads against the persistent SIMPLS stream with events rather than host barriers. CoreKNN precomputes fold neighbors.",
        "tests/test_cv.cpp and tests/test_metal.cpp check result sizes, worker/scheduler metadata, automatic-lane bounds, and that optimization does not decrease initial CV accuracy. The Metal suite additionally requires exact one/four-worker PLS-LDA prediction parity on a deterministic multi-class matrix; R and Python tests verify scheduler metadata parity.",
        "In matched M=100, Tcycle=100 CUDA wrapper runs, resident Gram reuse reduced PLS-LDA time from 175.497 to 164.441 s on MetRef and from 89.921 to 88.576 s on flow18. On Apple M3, persistent worker queues reduced median MetRef M=20/Tcycle=10 time from 7.390 to 7.014 s. After the fold-epoch fix, raw COIL20 one/four-worker accuracies matched across three seeds and median five-fold kernel time was 0.634/0.406 s. Four-core CPU/Metal speedup rose from 1.71x at 256 predictors to 10.78x at 16,384 predictors. The automatic lane study remains preliminary and sensitive to thermal/run order.",
    ),
    (
        "Exact-quota landmarking",
        "src/kodama_matrix.cpp forms coarse matrix-derived strata, allocates exact population-proportional quotas by randomized systematic rounding, and samples rows without replacement.",
        "tests/test_cv.cpp checks the exact effective count, distinct and represented samples, fixed-seed repeatability, and CPU/CUDA landmark-diagnostic parity.",
        "On flow18, 750,016 landmarks were selected from 300/300 coarse matrix strata in 0.321 s. The earlier exact graph bottleneck motivated resident IVF; the accepted direct graph measurement is 25.938 s versus 296.482 s exact.",
    ),
    (
        "Native Apple Metal backend",
        "src/metal_backend.mm owns persistent Metal state, exact and IVF-Flat KNN, Lloyd k-means, MPS matrix products, a bounded resident fold cache, device-budgeted independent-M scheduling, label-aware SIMPLS, resident score reductions, latent-space LDA scoring, native clean-sampler fixed-point atomic UMAP, and native FFT-grid openTSNE.",
        "tests/test_metal.cpp exercises exact/IVF KNN, graphs, CoreKNN/CorePLSLDA, matrix and graph-input KODAMA, UMAP, openTSNE, and direct class-sum, T'T, and discriminant-label equivalence against the former host calculation. R and Python tests cover public Metal UMAP/openTSNE dispatch.",
        "MetRef M=Tcycle=100 PLS-LDA improved from 631.415 to 312.029 s with exact optimization diagnostics and is slightly faster than the 316.329-s CPU4 reference; warm MetRef Metal graph construction is 0.023 s versus 0.152 s on CPU4. A native systems fixture measured exact Metal graph construction 5.56x faster than CPU4 at recall 1.000, while its tiny PLS-LDA CV remained 1.52x slower. BreastCancerDiagnostic M=Tcycle=100 confirmed the low-p crossover: 105.285 s Metal versus 6.990 s CPU4 with comparable diagnostics. Resident PLS-LDA statistics improved TabulaMuris from 1.799 to 0.492 s and MNIST70k from 2.296 to 1.704 s. Warm Metal openTSNE pair-distance correlation remains at least 0.9863 against fastEmbedR Metal.",
    ),
    (
        "Graph-input KODAMA",
        "The graph-input entry points in src/kodama_matrix.cpp accept supplied neighbor indices/distances; the KNN path reuses the graph and the PLS-LDA graph-only path builds self-tuning Laplacian features.",
        "R and Python wrapper tests exercise matrix_graph with KNN and PLS-LDA, and C++ tests cover graph construction, clustering, and visualization outputs.",
        "Graph-input is presented as an optional API. The manuscript separately states that graph-only PLS-LDA should not be interpreted as equivalent to data-input PLS-LDA without matched benchmark evidence.",
    ),
    (
        "Float32 randomized PCA",
        "src/pca.cpp supplies package-owned float32 preprocessing, sketching, QR, compact eigensolve, orientation, and variance reporting; src/pca_cuda.cu and the Metal matrix-multiply adapter execute dominant products on the selected accelerator.",
        "tests/test_cv.cpp checks dimensions, finite values, centered scores, orthonormal loadings, ordered singular values, determinism, float/double input parity, and CUDA agreement; tests/test_metal.cpp checks Metal agreement; R and Python tests exercise the public wrappers.",
        "PCA is an auxiliary public primitive rather than part of the KODAMA objective. Performance claims will be reported separately from KODAMA.matrix timings when a matched PCA benchmark is archived.",
    ),
    (
        "Visualization parity and initialization",
        "src/visualization.cpp exposes KODAMAVisualizationPCAInit and CPU UMAP/openTSNE; the CUDA embedding kernels supply the corresponding accelerator optimizers. KODAMA.matrix builds the full-data graph once before M and, when requested, uses one native PCA to retain both scaled starts. R and Python wrappers enforce explicit/stored/raw/fallback precedence and backend matching, so supplying raw data again does not recompute an already compatible stored start.",
        "C++ tests verify a single graph build for M > 1, both retained initialization matrices, centering, openTSNE scale, backend metadata, and raw-data overloads; wrapper tests verify default stored starts and graph-only fallbacks. A controlled CPU openTSNE comparison at the pinned fastEmbedR revision produced zero maximum and mean coordinate difference.",
        "On flow18 (1,000,021 samples), prior matched exact-graph CUDA KNN runs with Tcycle = 100 took 300.314 s for M = 1 and 301.091 s for M = 2; both reported graph_builds = 1, showing that M did not rebuild the graph. The accepted resident-IVF path reduces that shared graph to 25.938 s. UMAP fuzzy-graph edge counts and maximum weights match the pinned source on the three local parity datasets. CPU, CUDA, and Metal intentionally retain the backend-specific CSR, COO/CSR atomic, and clean-row sampling strategies of the audited source. Later coordinates are assessed by embedding diagnostics rather than bitwise equality because float32 reduction and atomic update order differ.",
    ),
    (
        "Standalone CUDA namespace",
        "The bundled CUDA visualization entry points and runtime controls use KODAMA-owned kodama_cuda_* and KODAMA_* names. No fastEmbedR-prefixed CUDA identifier remains in the standalone source tree.",
        "A standalone namespace audit scans source identifiers on every Unix CTest run and inspects the linked library symbol table when available. The dependency-free and physical Metal suites pass locally after the rename.",
        "This is a linkage-isolation correction, not a speed or numerical change. The tagged CUDA release must archive the linked symbol table and a same-process kodama-cpp plus fastEmbedR UMAP/openTSNE smoke test before the claim is closed.",
    ),
]


NOVELTY_ROWS = [
    (
        "State ownership",
        "Moves folds, labels, classifier workspaces, and graph outputs from repeated R-level calls into reusable C++ objects.",
    ),
    (
        "Backend-native classifiers",
        "Keeps each classifier's mathematics consistent across CPU, CUDA, and Metal while implementing its dominant operations separately for each backend.",
    ),
    (
        "Float32 execution",
        "Uses float32 analysis matrices and accelerator workspaces end to end while accepting wrapper-provided double inputs at the boundary.",
    ),
    (
        "Label-aware SIMPLS",
        "Computes PLS-LDA response cross-products from compact labels instead of dense one-hot matrices and evaluates the requested feasible component count.",
    ),
    (
        "Accelerated search",
        "Provides package-owned CPU HNSW, CUDA exact/recall-tuned IVF-Flat search, and Metal exact/recall-tuned IVF-Flat search.",
    ),
    (
        "Reproducible heterogeneity",
        "Keeps M runs independent and reports actual backend and search metadata so acceleration can be tested without altering the objective.",
    ),
    (
        "Search evolution",
        "Makes grouped adaptive proposals, transition-driven class moves, and the label-only degeneracy guard explicit and separately reports their raw CV accuracy and acceptance score.",
    ),
]


REFERENCES = [
    (
        "Cacciatore, S., Luchinat, C., and Tenori, L. Knowledge discovery by accuracy maximization. "
        "Proceedings of the National Academy of Sciences, 111(14), 5117-5122, 2014."
    ),
    (
        "Cacciatore, S., Tenori, L., Luchinat, C., Bennett, P. R., and MacIntyre, D. A. "
        "KODAMA: an R package for knowledge discovery and data mining. Bioinformatics, "
        "33(4), 621-623, 2017."
    ),
    (
        "Cacciatore, S. and Tenori, L. KODAMA: Knowledge Discovery by Accuracy Maximization. "
        "R package documentation, CRAN."
    ),
    (
        "Chapelle, O., Scholkopf, B., and Zien, A. Semi-Supervised Learning. MIT Press, 2006."
    ),
    (
        "Zhu, X., Ghahramani, Z., and Lafferty, J. Semi-supervised learning using Gaussian fields "
        "and harmonic functions. Proceedings of ICML, 2003."
    ),
    (
        "Zhou, D., Bousquet, O., Lal, T. N., Weston, J., and Scholkopf, B. Learning with local "
        "and global consistency. Advances in Neural Information Processing Systems, 2003."
    ),
    (
        "Lee, D.-H. Pseudo-label: The simple and efficient semi-supervised learning method for "
        "deep neural networks. ICML Workshop on Challenges in Representation Learning, 2013."
    ),
    (
        "Ratner, A., De Sa, C., Wu, S., Selsam, D., and Re, C. Data programming: Creating large "
        "training sets, quickly. Advances in Neural Information Processing Systems, 2016."
    ),
    (
        "Ben-Hur, A., Elisseeff, A., and Guyon, I. A stability based method for discovering "
        "structure in clustered data. Pacific Symposium on Biocomputing, 2002."
    ),
    (
        "Tibshirani, R. and Walther, G. Cluster validation by prediction strength. Journal of "
        "Computational and Graphical Statistics, 14(3), 511-528, 2005."
    ),
    (
        "Hubert, L. and Arabie, P. Comparing partitions. Journal of Classification, 2, "
        "193-218, 1985."
    ),
    (
        "Rousseeuw, P. J. Silhouettes: a graphical aid to the interpretation and validation of "
        "cluster analysis. Journal of Computational and Applied Mathematics, 20, 53-65, 1987."
    ),
    (
        "Kohavi, R. A study of cross-validation and bootstrap for accuracy estimation and model "
        "selection. Proceedings of IJCAI, 1995."
    ),
    (
        "Ambroise, C. and McLachlan, G. J. Selection bias in gene extraction on the basis of "
        "microarray gene-expression data. Proceedings of the National Academy of Sciences, "
        "99(10), 6562-6566, 2002."
    ),
    (
        "Varma, S. and Simon, R. Bias in error estimation when using cross-validation for model "
        "selection. BMC Bioinformatics, 7, 91, 2006."
    ),
    (
        "Kriegeskorte, N., Simmons, W. K., Bellgowan, P. S. F., and Baker, C. I. Circular "
        "analysis in systems neuroscience: the dangers of double dipping. Nature Neuroscience, "
        "12, 535-540, 2009."
    ),
    (
        "de Jong, S. SIMPLS: an alternative approach to partial least squares regression. "
        "Chemometrics and Intelligent Laboratory Systems, 18(3), 251-263, 1993."
    ),
    (
        "Johnson, J., Douze, M., and Jegou, H. Billion-scale similarity search with GPUs. "
        "IEEE Transactions on Big Data, 2019."
    ),
    (
        "Malkov, Y. A. and Yashunin, D. A. Efficient and robust approximate nearest neighbor "
        "search using Hierarchical Navigable Small World graphs. IEEE Transactions on Pattern "
        "Analysis and Machine Intelligence, 2020."
    ),
    (
        "McInnes, L., Healy, J., and Melville, J. UMAP: Uniform Manifold Approximation and "
        "Projection for dimension reduction. arXiv:1802.03426, 2018."
    ),
    (
        "van der Maaten, L. and Hinton, G. Visualizing data using t-SNE. Journal of Machine "
        "Learning Research, 9, 2579-2605, 2008."
    ),
    (
        "Linderman, G. C., Rachh, M., Hoskins, J. G., Steinerberger, S., and Kluger, Y. "
        "Fast interpolation-based t-SNE for improved visualization of single-cell RNA-seq data. "
        "Nature Methods, 16, 243-245, 2019."
    ),
    (
        "Sonnenburg, S., Braun, M. L., Ong, C. S., Bengio, S., Bottou, L., Holmes, G., LeCun, Y., "
        "Muller, K.-R., Pereira, F., Rasmussen, C. E., Ratsch, G., Scholkopf, B., Smola, A., "
        "Vincent, P., Weston, J., and Williamson, R. C. The need for open source software in "
        "machine learning. Journal of Machine Learning Research, 8, 2443-2466, 2007."
    ),
    (
        "Pineau, J., Vincent-Lamarre, P., Sinha, K., Lariviere, V., Beygelzimer, A., d'Alche-Buc, "
        "F., Fox, E., and Larochelle, H. Improving reproducibility in machine learning research. "
        "Journal of Machine Learning Research, 22(164), 1-20, 2021."
    ),
]


def set_run_font(run, size: float | None = None, bold: bool | None = None, color=None) -> None:
    run.font.name = TOKENS["font"]
    run._element.rPr.rFonts.set(qn("w:ascii"), TOKENS["font"])
    run._element.rPr.rFonts.set(qn("w:hAnsi"), TOKENS["font"])
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D0D7DE", size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_width(table, widths_in) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for idx, width in enumerate(widths_in):
            cell = row.cells[idx]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def apply_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = TOKENS["font"]
    normal.font.size = Pt(TOKENS["body_size"])
    normal._element.rPr.rFonts.set(qn("w:ascii"), TOKENS["font"])
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), TOKENS["font"])
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in [
        ("Heading 1", TOKENS["h1_size"], TOKENS["heading_blue"], 16, 8),
        ("Heading 2", TOKENS["h2_size"], TOKENS["heading_blue"], 12, 6),
        ("Heading 3", TOKENS["h3_size"], TOKENS["heading_dark"], 8, 4),
    ]:
        style = styles[name]
        style.font.name = TOKENS["font"]
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:ascii"), TOKENS["font"])
        style._element.rPr.rFonts.set(qn("w:hAnsi"), TOKENS["font"])
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def apply_compact_memo_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    normal = doc.styles["Normal"]
    normal.font.size = Pt(10.3)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[name]
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(3)


def add_title(
    doc: Document,
    title: str,
    subtitle: str,
    authors: str,
    author_records=None,
    affiliations=None,
    author_notes=None,
    corresponding_author: str | None = None,
) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    set_run_font(run, 22, True, TOKENS["title"])

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(subtitle)
    set_run_font(run, 11)
    run.italic = True

    if author_records:
        table = doc.add_table(rows=0, cols=2)
        for name, markers, email in author_records:
            cells = table.add_row().cells
            for cell in cells:
                set_cell_margins(cell, top=0, bottom=0, start=0, end=0)
            p = cells[0].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(name)
            set_run_font(run, 10, True, TOKENS["muted"])
            run = p.add_run(markers)
            set_run_font(run, 8, False, TOKENS["muted"])
            run.font.superscript = True
            p = cells[1].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(email)
            set_run_font(run, 8.5, False, TOKENS["muted"])
            run.font.small_caps = True
        set_table_width(table, [3.1, 3.35])
        set_table_borders(table, color="FFFFFF", size="0")
    else:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(authors)
        set_run_font(run, 10, False, TOKENS["muted"])

    if affiliations:
        for marker, affiliation in affiliations:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(marker)
            set_run_font(run, 7.5, False, TOKENS["muted"])
            run.font.superscript = True
            run = p.add_run(" " + affiliation)
            set_run_font(run, 8.5, False, TOKENS["muted"])
            run.italic = True

    if author_notes:
        for marker, note in author_notes:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(marker)
            set_run_font(run, 7.5, False, TOKENS["muted"])
            run.font.superscript = True
            run = p.add_run(" " + note)
            set_run_font(run, 8.5, False, TOKENS["muted"])

    if corresponding_author:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run("Corresponding author: " + corresponding_author)
        set_run_font(run, 8.5, False, TOKENS["muted"])


def add_callout(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.45])
    set_table_borders(table, color=TOKENS["border"])
    cell = table.cell(0, 0)
    set_cell_shading(cell, TOKENS["callout_fill"])
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(label + ": ")
    set_run_font(run, 10.5, True, RGBColor(0x1F, 0x3A, 0x5F))
    run = p.add_run(text)
    set_run_font(run, 10.5)


def add_bullets(doc: Document, items) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_run_font(run)


def new_numbering_id(doc: Document, style_id: str = "ListNumber") -> int:
    numbering = doc.part.numbering_part.element
    abstract_num_id = None
    for abstract in numbering.findall(qn("w:abstractNum")):
        for p_style in abstract.iter(qn("w:pStyle")):
            if p_style.get(qn("w:val")) == style_id:
                abstract_num_id = abstract.get(qn("w:abstractNumId"))
                break
        if abstract_num_id is not None:
            break
    if abstract_num_id is None:
        abstract_num_id = "0"

    existing = [
        int(num.get(qn("w:numId")))
        for num in numbering.findall(qn("w:num"))
        if num.get(qn("w:numId")) is not None
    ]
    num_id = (max(existing) + 1) if existing else 1

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), str(abstract_num_id))
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return num_id


def apply_numbering_id(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    num_id_node = num_pr.find(qn("w:numId"))
    if num_id_node is None:
        num_id_node = OxmlElement("w:numId")
        num_pr.append(num_id_node)
    num_id_node.set(qn("w:val"), str(num_id))


def add_numbered(doc: Document, items) -> None:
    num_id = new_numbering_id(doc)
    for item in items:
        p = doc.add_paragraph(style="List Number")
        apply_numbering_id(p, num_id)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_run_font(run)


def add_pseudocode_block(
    doc: Document,
    lines,
    *,
    font_size: float = 7.2,
    line_spacing: float = 0.88,
) -> None:
    for idx, line in enumerate(lines):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.16)
        paragraph.paragraph_format.right_indent = Inches(0.16)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = line_spacing
        paragraph.paragraph_format.keep_together = False
        set_paragraph_shading(paragraph, "F7F9FC")
        run = paragraph.add_run(line if line else " ")
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
        run.font.size = Pt(font_size)
        if line.startswith("Algorithm "):
            run.font.bold = True
            run.font.color.rgb = TOKENS["title"]
        if idx == len(lines) - 1:
            paragraph.paragraph_format.space_after = Pt(6)


def format_table_text(table, font_size=8.5) -> None:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.1
                for run in p.runs:
                    set_run_font(run, font_size)


def add_table(doc: Document, headers, rows, widths, font_size=8.5) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_borders(table)
    set_table_width(table, widths)
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = text
        set_cell_shading(cell, TOKENS["table_fill"])
        for p in cell.paragraphs:
            p.paragraph_format.keep_with_next = True
            for run in p.runs:
                set_run_font(run, font_size, True)
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    table_header = OxmlElement("w:tblHeader")
    table_header.set(qn("w:val"), "true")
    header_properties.append(table_header)
    # Rows are appended after the initial width pass, so apply the row-level
    # no-split property again once the table is complete.
    set_table_width(table, widths)
    format_table_text(table, font_size=font_size)


def build_architecture_figure() -> None:
    from PIL import Image, ImageDraw, ImageFont

    scale = 2
    width, height = 1800, 1120
    img = Image.new("RGB", (width * scale, height * scale), "white")
    draw = ImageDraw.Draw(img)

    def font(size, bold=False):
        candidates = [
            "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
            "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
            "/System/Library/Fonts/Supplemental/Helvetica Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Helvetica.ttf",
        ]
        for candidate in candidates:
            if candidate and Path(candidate).exists():
                return ImageFont.truetype(candidate, size * scale)
        return ImageFont.load_default()

    title_font = font(42, True)
    subtitle_font = font(25)
    box_title_font = font(27, True)
    small_font = font(20)

    def xy(coords):
        return tuple(int(v * scale) for v in coords)

    def box(coords, fill, outline, lines, text_fill=(20, 35, 55), radius=20):
        draw.rounded_rectangle(xy(coords), radius=radius * scale, fill=fill, outline=outline, width=3 * scale)
        x0, y0, x1, y1 = xy(coords)
        rendered = []
        total_h = 0
        for index, line in enumerate(lines):
            line_font = box_title_font if index == 0 else small_font
            bbox = draw.textbbox((0, 0), line, font=line_font)
            height = bbox[3] - bbox[1]
            rendered.append((line, line_font, bbox[2] - bbox[0], height))
            total_h += height
        gap = int(7 * scale)
        total_h += gap * (len(rendered) - 1)
        y = y0 + ((y1 - y0) - total_h) // 2
        for line, line_font, width_px, height_px in rendered:
            draw.text((x0 + ((x1 - x0) - width_px) // 2, y), line, font=line_font, fill=text_fill)
            y += height_px + gap

    def arrow(start, end, color=(85, 98, 115), width_px=5):
        sx, sy = start
        ex, ey = end
        draw.line(xy((sx, sy, ex, ey)), fill=color, width=width_px * scale)
        import math
        angle = math.atan2(ey - sy, ex - sx)
        length = 18
        spread = 0.42
        p1 = (ex - length * math.cos(angle - spread), ey - length * math.sin(angle - spread))
        p2 = (ex - length * math.cos(angle + spread), ey - length * math.sin(angle + spread))
        draw.polygon([xy((ex, ey))[0:2], xy(p1)[0:2], xy(p2)[0:2]], fill=color)

    # Palette chosen to keep the figure readable when printed in grayscale.
    navy = (11, 37, 69)
    blue = (232, 241, 252)
    blue_line = (91, 139, 190)
    green = (232, 246, 239)
    green_line = (84, 150, 111)
    amber = (252, 243, 224)
    amber_line = (196, 143, 60)
    purple = (242, 239, 253)
    purple_line = (118, 101, 203)
    gray = (246, 248, 251)
    gray_line = (180, 193, 208)

    draw.text(xy((70, 48)), "kodama-cpp architecture", font=title_font, fill=navy)
    draw.text(
        xy((70, 105)),
        "Standalone C++17 core with dependency-light CPU/Metal and opt-in CUDA backends",
        font=subtitle_font,
        fill=(75, 85, 100),
    )

    box((170, 165, 630, 275), blue, blue_line, ["R wrapper", "thin language binding"])
    box((1170, 165, 1630, 275), blue, blue_line, ["Python wrapper", "thin language binding"])
    box(
        (480, 320, 1320, 500),
        gray,
        gray_line,
        ["kodama-cpp C++17 core", "MatrixView + typed options/results", "KNNCV | PLSLDACV | CoreKNN | CorePLSLDA | KODAMA.matrix | PCA"],
    )

    box((60, 555, 530, 765), green, green_line, ["CPU backend", "package-owned float32 HNSW", "SIMPLS/LDA + randomized PCA", "optional OpenMP"])
    box((665, 555, 1135, 765), purple, purple_line, ["Apple Metal backend", "exact + resident IVF-Flat KNN", "MPS SIMPLS/LDA + PCA", "system frameworks only"])
    box((1270, 555, 1740, 765), amber, amber_line, ["CUDA backend", "native exact + resident IVF-Flat", "k-means + SIMPLS/LDA + PCA", "CUDA Toolkit only"])

    box((430, 830, 1370, 930), blue, blue_line, ["Graph, PCA, embedding, and clustering utilities", "CPU/CUDA/Metal PCA + graph + UMAP + openTSNE | CPU random walks"])
    box((565, 980, 1235, 1070), gray, gray_line, ["Typed outputs", "labels, accuracy traces, graphs, embeddings", "timings, memory, backend metadata"])

    arrow((630, 220), (700, 320), color=blue_line)
    arrow((1170, 220), (1100, 320), color=blue_line)
    arrow((650, 500), (400, 555), color=green_line)
    arrow((900, 500), (900, 555), color=purple_line)
    arrow((1150, 500), (1400, 555), color=amber_line)
    arrow((295, 765), (590, 830), color=green_line)
    arrow((900, 765), (900, 830), color=purple_line)
    arrow((1505, 765), (1210, 830), color=amber_line)
    arrow((900, 930), (900, 980), color=blue_line)

    img = img.resize((width, height), Image.Resampling.LANCZOS)
    img.save(ARCH_FIGURE)


def build_docx() -> None:
    build_architecture_figure()
    doc = Document()
    apply_styles(doc)
    add_title(
        doc,
        "kodama-cpp: Cross-Validated Accuracy Maximization on CPU, CUDA, and Apple Metal",
        "Manuscript for the Journal of Machine Learning Research Machine Learning Open Source Software track",
        "",
        author_records=AUTHORS,
        affiliations=AFFILIATIONS,
        author_notes=AUTHOR_NOTES,
    )

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(ABSTRACT)

    for heading, paragraphs in SECTIONS:
        level = 1 if heading[0].isdigit() else 2
        if heading == "7. Limitations":
            doc.add_page_break()
        doc.add_heading(heading, level=level)
        for paragraph in paragraphs:
            doc.add_paragraph(paragraph)
        if heading == "1. Introduction":
            doc.add_heading("Main contributions", level=2)
            add_table(doc, ("Contribution", "Role in kodama-cpp"), CONTRIBUTION_ROWS, [1.75, 4.7])
        if heading == "2. KODAMA objective":
            add_callout(
                doc,
                "Objective",
                "c* = argmax_c A(c; X, F, Pi), where A is cross-validated accuracy under the chosen classifier and fold assignment.",
            )
            doc.add_heading("Landmark-projection strategy", level=2)
            for paragraph in LANDMARKING_NOVELTY_PARAGRAPHS:
                doc.add_paragraph(paragraph)
            add_table(
                doc,
                ("Validation dimension", "Current manuscript evidence"),
                LANDMARK_VALIDATION_ROWS,
                [1.45, 5.0],
                font_size=8.0,
            )
            doc.add_heading("Label-vector search mechanics", level=2)
            for paragraph in LABEL_SEARCH_DETAIL_PARAGRAPHS:
                doc.add_paragraph(paragraph)
            doc.add_heading("Raw accuracy, acceptance score, and diagnostics", level=2)
            for paragraph in SCORE_INTERPRETATION_PARAGRAPHS:
                doc.add_paragraph(paragraph)
            doc.add_heading("Reproducible KODAMA.matrix procedure", level=2)
            doc.add_paragraph(
                "The public KODAMA.matrix routine is reproducible from the following option definitions and run-level procedure. "
                "The description below names the C++ options because the R and Python wrappers should expose the same contract."
            )
            add_table(doc, ("Option", "Role in the algorithm"), KODAMA_PARAMETER_ROWS, [1.35, 5.1], font_size=8.0)
            doc.add_heading("Run-level optimization rule", level=3)
            add_numbered(doc, REPRODUCIBLE_ALGORITHM_STEPS)
            doc.add_heading("Final graph and dissimilarity", level=3)
            for paragraph in GRAPH_CONSTRUCTION_PARAGRAPHS:
                doc.add_paragraph(paragraph)
            doc.add_heading("Main algorithm pseudocode", level=2)
            doc.add_paragraph(
                "The following compact algorithm connects landmark selection, guided label evolution, "
                "projection, and agreement-graph construction. Each inner cycle contributes exactly "
                "one new CV evaluation; the initial state contributes one baseline CV evaluation."
            )
            add_pseudocode_block(doc, MAIN_ALGORITHM_PSEUDOCODE_LINES)
        if heading == "3. Implementation architecture":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(4)
            p.add_run().add_picture(str(ARCH_FIGURE), width=Inches(6.45))
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.paragraph_format.space_after = Pt(8)
            run = caption.add_run(
                "Figure 1. Architecture of kodama-cpp. The C++17 core owns the numerical kernels, "
                "CPU, CUDA, and Apple Metal backends, graph utilities, and typed outputs; R and Python remain thin wrappers."
            )
            set_run_font(run, 9.2, False, TOKENS["muted"])
            doc.add_heading("Implementation novelty", level=2)
            for paragraph in IMPLEMENTATION_NOVELTY_PARAGRAPHS:
                doc.add_paragraph(paragraph)
            doc.add_heading("CUDA backend", level=2)
            for paragraph in CUDA_BACKEND_PARAGRAPHS:
                doc.add_paragraph(paragraph)
            doc.add_heading("Apple Metal backend", level=2)
            for paragraph in METAL_BACKEND_PARAGRAPHS:
                doc.add_paragraph(paragraph)
            doc.add_heading("Shared accelerator contract", level=2)
            add_table(
                doc,
                ("Invariant", "CUDA realization", "Metal realization"),
                BACKEND_CONTRACT_ROWS,
                [1.25, 2.6, 2.6],
                font_size=7.5,
            )
            doc.add_heading("Public API summary", level=2)
            add_table(doc, ("Layer", "Public API", "Purpose"), API_ROWS, [1.35, 1.75, 3.35])
            doc.add_heading("Graph-first input contract", level=2)
            doc.add_paragraph(
                "KODAMA.graph returns a matrix- or external-handle-backed graph, backend-matched PCA "
                "starts, metadata, and timings, but never the source matrix. Raw features remain "
                "caller-owned when supplied separately."
            )
            add_table(
                doc,
                ("Input", "Working geometry", "KNN evolution", "PLS-LDA evolution", "Graph builds"),
                GRAPH_INPUT_CONTRACT_ROWS,
                [0.8, 1.55, 1.35, 1.45, 0.8],
                font_size=7.2,
            )
            doc.add_page_break()
            add_pseudocode_block(doc, GRAPH_PREPARATION_PSEUDOCODE_LINES, font_size=7.0, line_spacing=0.86)
            doc.add_heading("Graph-API smoke validation", level=3)
            add_table(
                doc,
                ("Classifier", "Input", "Builds", "Core s", "Best CV", "Agreement"),
                GRAPH_API_SMOKE_ROWS,
                [1.05, 0.8, 0.65, 0.75, 0.85, 0.9],
                font_size=7.2,
            )
            doc.add_paragraph(GRAPH_API_SMOKE_NOTE)
            doc.add_heading("Computational scaling", level=2)
            add_table(doc, ("Component", "Dominant work", "Implementation implication"), COMPLEXITY_ROWS, [1.25, 2.7, 2.5], font_size=7.6)
            doc.add_heading("Graph-input KODAMA validation boundary", level=2)
            add_table(doc, ("Path", "Implementation", "Interpretation"), GRAPH_INPUT_VALIDATION_ROWS, [1.25, 2.75, 2.45], font_size=7.4)
            doc.add_heading("KODAMA directly from a KNN matrix", level=2)
            for paragraph in GRAPH_INPUT_PROCEDURE_PARAGRAPHS:
                doc.add_paragraph(paragraph)
            add_pseudocode_block(doc, GRAPH_INPUT_USAGE_LINES, font_size=6.9, line_spacing=0.84)
        if heading == "4. Relationship to public KODAMA literature":
            doc.add_heading("Compatibility with the R implementation", level=2)
            add_table(doc, ("Aspect", "Public R/literature contract", "kodama-cpp"), COMPATIBILITY_ROWS, [1.25, 2.65, 2.55])
            doc.add_heading("Implementation novelty relative to the R package", level=2)
            add_table(doc, ("Innovation", "What changes in kodama-cpp"), NOVELTY_ROWS, [1.65, 4.8], font_size=8.2)
            doc.add_heading("Relationship to semi-supervised learning and cluster validation", level=2)
            for paragraph in RELATED_WORK_PARAGRAPHS:
                doc.add_paragraph(paragraph)
            add_table(
                doc,
                ("Literature area", "Typical goal", "Role of KODAMA"),
                RELATED_WORK_ROWS,
                [1.45, 2.35, 2.65],
                font_size=7.8,
            )
        if heading == "6. Availability and reproducibility":
            doc.add_heading("Installation paths", level=2)
            add_table(doc, ("Target", "Command or check"), INSTALLATION_ROWS, [1.1, 5.35], font_size=7.4)
            doc.add_heading("Wrapper validation", level=2)
            add_table(doc, ("Target", "Validation result"), WRAPPER_VALIDATION_ROWS, [1.5, 4.95], font_size=8.0)
            doc.add_heading("License and dependency notes", level=2)
            add_table(doc, ("Component", "Release note"), LICENSE_DEPENDENCY_ROWS, [1.8, 4.65], font_size=8.0)
        if heading == "5. Evaluation":
            doc.add_heading("Audited implementation improvements", level=2)
            doc.add_paragraph(
                "These accepted changes preserve the KODAMA objective while correcting backend "
                "behavior, memory ownership, or repeated-work overhead. The complete audit is "
                "retained in docs/full-code-audit-2026-08-07.md."
            )
            add_table(
                doc,
                ("Area", "Accepted implementation", "Evidence"),
                AUDIT_IMPROVEMENT_ROWS,
                [1.25, 3.45, 1.65],
                font_size=7.1,
            )
            doc.add_heading("Benchmark protocol and data coverage", level=2)
            add_table(doc, ("Item", "Value"), BENCHMARK_PROTOCOL_ROWS, [1.45, 5.0], font_size=8.0)
            doc.add_heading("Dataset inventory", level=3)
            add_table(
                doc,
                ("Dataset", "Samples", "Variables", "Classes"),
                PILOT_DATASET_ROWS,
                [1.8, 1.25, 1.25, 1.25],
                font_size=7.6,
            )
            doc.add_heading("Evaluation guardrails and ablations", level=2)
            add_table(
                doc,
                ("Concern", "Manuscript position", "Evidence reported or controlled"),
                EVALUATION_GUARDRAIL_ROWS,
                [1.15, 2.25, 3.05],
                font_size=7.5,
            )
            add_table(
                doc,
                ("Ablation", "Comparison", "Question answered"),
                ABLATION_MATRIX_ROWS,
                [1.2, 1.9, 3.35],
                font_size=7.5,
            )
            doc.add_heading("Dependency-free CPU and Apple Metal validation", level=2)
            for paragraph in METAL_VALIDATION_PARAGRAPHS:
                doc.add_paragraph(paragraph)
            add_table(
                doc,
                ("Dataset", "Scope", "CPU s", "Metal s", "Speedup", "Quality CPU / Metal"),
                METAL_VALIDATION_ROWS,
                [0.7, 1.35, 0.65, 0.7, 0.65, 2.4],
                font_size=6.8,
            )
            doc.add_heading("CUDA workstation validation", level=2)
            for paragraph in PILOT_EXPERIMENT_PARAGRAPHS:
                doc.add_paragraph(paragraph)
            add_table(
                doc,
                ("Dataset", "Path", "Accuracy", "Seconds", "nlist", "max nprobe", "pilot recall"),
                NATIVE_CUDA_VALIDATION_ROWS,
                [1.0, 1.25, 0.85, 0.8, 0.7, 0.9, 0.95],
                font_size=7.2,
            )
            doc.add_heading("Earlier accelerator validation snapshot", level=3)
            add_table(
                doc,
                ("Dataset", "Kernel", "CPU s", "CUDA s", "Speedup", "Acc CPU/CUDA", "Delta", "Comp."),
                PILOT_CV_ROWS,
                [0.75, 0.75, 0.7, 0.7, 0.65, 1.65, 0.7, 0.45],
                font_size=6.4,
            )
            doc.add_heading("Core optimizer medians", level=3)
            add_table(
                doc,
                ("Dataset", "Core", "CPU s", "CUDA s", "Speedup", "Acc CPU/CUDA", "ARI CPU/CUDA", "Classes"),
                PILOT_CORE_ROWS,
                [0.78, 0.95, 0.72, 0.72, 0.65, 1.0, 1.0, 0.63],
                font_size=6.7,
            )
            doc.add_heading("MultiCPU versus GPU KODAMA comparison", level=3)
            for paragraph in KODAMA_BACKEND_COMPARISON_PARAGRAPHS:
                doc.add_paragraph(paragraph)
            add_table(
                doc,
                ("Dataset", "Classifier", "Backend", "Seconds", "CV acc", "ARI", "Classes", "Scope"),
                KODAMA_BACKEND_COMPARISON_ROWS,
                [0.75, 0.75, 0.65, 0.65, 0.65, 0.55, 0.55, 1.0],
                font_size=6.9,
            )
            doc.add_heading("KODAMA 2.4.1/2.4 predecessor comparison", level=3)
            for paragraph in HISTORICAL_KNN_PARAGRAPHS:
                doc.add_paragraph(paragraph)
            add_table(
                doc,
                ("Implementation", "Backend", "Workers", "Seconds", "Speedup vs legacy R", "Best/median CV", "Selected/median ARI", "Selected/median classes"),
                HISTORICAL_KNN_ROWS,
                [1.15, 0.65, 0.55, 0.7, 0.8, 1.0, 1.05, 1.05],
                font_size=6.8,
            )
            for paragraph in HISTORICAL_PLS_PARAGRAPHS:
                doc.add_paragraph(paragraph)
            add_table(
                doc,
                ("Implementation", "Evaluator", "Backend", "Workers", "Components", "Seconds", "Best/median CV", "Post hoc diagnostic"),
                HISTORICAL_PLS_ROWS,
                [1.05, 1.25, 0.6, 0.55, 0.65, 0.65, 1.0, 1.25],
                font_size=6.7,
            )
            doc.add_heading("Current-commit KODAMA.matrix MetRef validation", level=3)
            add_table(
                doc,
                ("Backend/classifier", "Wall s", "Core s", "Max/median CV", "Selected ARI/truth sil.", "Selected classes", "Peak MB"),
                PILOT_MATRIX_ROWS,
                [0.85, 0.75, 0.75, 1.25, 1.2, 0.8, 0.85],
                font_size=7.0,
            )
            if METREF_BACKEND_FIGURE.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(6)
                p.add_run().add_picture(str(METREF_BACKEND_FIGURE), width=Inches(6.45))
                cap = doc.add_paragraph(
                    "Figure 2. Historical 2026-07-20 MetRef embedding-parity diagnostic after the float32 "
                    "PLS-LDA optimization. Classic data, CPU KODAMA, and CUDA KODAMA were deliberately "
                    "rendered through the same CPU embedding implementation and PCA initialization to isolate "
                    "the KODAMA backend. This is not the current operational backend-native visualization "
                    "contract. Colors are external reference labels withheld from optimization."
                )
                cap.style = doc.styles["Caption"]
            doc.add_heading("Classic versus KODAMA visualization validation", level=3)
            for paragraph in VISUALIZATION_COMPARISON_PARAGRAPHS:
                doc.add_paragraph(paragraph)
            if HPC_INFERENCE_ROWS:
                add_table(
                    doc,
                    ("Classifier", "Embedding", "Positive", "Median delta", "Bootstrap 95% CI", "Wilcoxon p", "Holm p", "Sign p"),
                    HPC_INFERENCE_ROWS,
                    [0.75, 0.75, 0.65, 0.8, 1.1, 0.75, 0.65, 0.65],
                    font_size=6.9,
                )
                separator = doc.add_paragraph()
                separator.paragraph_format.space_after = Pt(3)
            if HPC_DATASET_ROWS:
                add_table(
                    doc,
                    ("Dataset", "n", "p", "Domain / representation"),
                    HPC_DATASET_ROWS,
                    [1.55, 0.75, 0.65, 3.05],
                    font_size=7.4,
                )
                separator = doc.add_paragraph()
                separator.paragraph_format.space_after = Pt(3)
            if HPC_RUNTIME_ROWS:
                runtime_heading = doc.add_heading("All-dataset KODAMA timing coverage", level=3)
                runtime_heading.paragraph_format.keep_with_next = True
                doc.add_paragraph(
                    "Entries are median CUDA KODAMA core wall times in seconds over seeds 4, 17, and 42 "
                    "at M=Tcycle=100. A dash marks a classifier run absent from the uploaded archive."
                )
                add_table(
                    doc,
                    ("Dataset", "KNN core s", "PLS-LDA core s", "Coverage"),
                    HPC_RUNTIME_ROWS,
                    [2.25, 1.1, 1.25, 1.75],
                    font_size=7.3,
                )
                separator = doc.add_paragraph()
                separator.paragraph_format.space_after = Pt(3)
                doc.add_heading("Frozen HPC rerun plan", level=3)
                add_table(
                    doc,
                    ("Stage", "Predeclared requirement"),
                    HPC_RERUN_PLAN_ROWS,
                    [1.15, 5.2],
                    font_size=7.4,
                )
            add_table(
                doc,
                ("Dataset", "Classifier", "Embedding", "Sil. classic -> KODAMA", "Delta sil.", "Trust classic -> KODAMA", "ARI / classes", "KODAMA s"),
                VISUALIZATION_COMPARISON_ROWS,
                [0.82, 0.72, 0.72, 1.22, 0.62, 1.25, 0.8, 0.7],
                font_size=6.5,
            )
            if IMAGENET_COMPARISON_FIGURE.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(4)
                p.add_run().add_picture(str(IMAGENET_COMPARISON_FIGURE), width=Inches(5.5))
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.paragraph_format.space_after = Pt(8)
                run = caption.add_run(
                    "ImageNet seed-4 visualization comparison. The three-seed aggregate diagnostics are "
                    "reported in the accompanying table and are not measurements of this single displayed run."
                )
                set_run_font(run, 9.0, False, TOKENS["muted"])
            if HPC_VISUALIZATION_FIGURE.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(4)
                p.add_run().add_picture(str(HPC_VISUALIZATION_FIGURE), width=Inches(6.45))
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.paragraph_format.space_after = Pt(8)
                run = caption.add_run(
                    "Figure 3. Paired three-seed CUDA comparison of classic fastEmbedR and KODAMA-corrected "
                    "UMAP/openTSNE at M = 100 and Tcycle = 100. Points are dataset-level median silhouette "
                    "changes; positive values favor KODAMA. External labels were used only after optimization."
                )
                set_run_font(run, 9.0, False, TOKENS["muted"])
            if PENDIGITS_ADVERSE_FIGURE.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(4)
                p.add_run().add_picture(str(PENDIGITS_ADVERSE_FIGURE), width=Inches(6.45))
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.paragraph_format.space_after = Pt(8)
                run = caption.add_run(
                    "Figure 4. PenDigits seed-4 CPU4 adverse control at M = Tcycle = 100. "
                    "Classic UMAP retains compact truth classes, whereas both KODAMA-corrected "
                    "layouts are diffuse despite high internal CV accuracy. External labels were "
                    "used only for coloring and post hoc diagnostics."
                )
                set_run_font(run, 9.0, False, TOKENS["muted"])
            doc.add_heading("Sensitivity of M and Tcycle", level=2)
            for paragraph in PARAMETER_SENSITIVITY_PARAGRAPHS:
                doc.add_paragraph(paragraph)
            if RELEASE_SENSITIVITY_READY:
                add_table(
                    doc,
                    ("Dataset", "Classifier", "Best acc T20 -> T100", "Median acc T20 -> T100", "Median classes T20 -> T100", "Seconds"),
                    TCYCLE_SENSITIVITY_ROWS,
                    [0.8, 0.8, 1.25, 1.35, 1.0, 0.9],
                    font_size=7.6,
                )
                separator = doc.add_paragraph()
                separator.paragraph_format.space_after = Pt(3)
                add_table(
                    doc,
                    ("Dataset", "Classifier", "Best acc M20/M50/M100", "Median acc M20/M50/M100", "Selected ARI", "Seconds"),
                    M_SENSITIVITY_ROWS,
                    [0.8, 0.8, 1.4, 1.4, 1.1, 1.0],
                    font_size=7.4,
                )
                if SENSITIVITY_FIGURE.exists():
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(4)
                    p.add_run().add_picture(str(SENSITIVITY_FIGURE), width=Inches(6.45))
                    caption = doc.add_paragraph()
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.paragraph_format.space_after = Pt(8)
                    run = caption.add_run(
                        "Figure 4. Sensitivity of KODAMA quality and runtime to Tcycle and M on named matrix-input datasets."
                    )
                    set_run_font(run, 9.0, False, TOKENS["muted"])
                doc.add_heading("Agreement-graph convergence", level=3)
                add_table(
                    doc,
                    ("Dataset", "Classifier", "RMSE to M100 at M10/M20/M50", "Correlation at M50"),
                    ENSEMBLE_CONVERGENCE_ROWS,
                    [0.9, 0.9, 2.5, 1.25],
                    font_size=7.5,
                )
                if ENSEMBLE_CONVERGENCE_FIGURE.exists():
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(4)
                    p.add_run().add_picture(str(ENSEMBLE_CONVERGENCE_FIGURE), width=Inches(6.45))
                    caption = doc.add_paragraph()
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.paragraph_format.space_after = Pt(8)
                    run = caption.add_run(
                        "Figure 5. Convergence of edge-agreement weights as independent KODAMA runs are added. "
                        "RMSE is measured against the M = 100 ensemble; the right panel shows the Bernoulli "
                        "worst-case Monte Carlo standard error."
                    )
                    set_run_font(run, 9.0, False, TOKENS["muted"])
            else:
                doc.add_paragraph(
                    "Named matrix-input release results are pending. Preliminary anonymized development rows are deliberately excluded from this submission draft."
                )
            doc.add_heading("Implementation claims and evidence", level=2)
            add_table(
                doc,
                ("Feature claim", "Where implemented", "Test proving it", "Benchmark proving benefit"),
                IMPLEMENTATION_EVIDENCE_ROWS,
                [1.15, 1.8, 1.85, 1.65],
                font_size=7.4,
            )
            doc.add_heading("Release-validation evidence", level=2)
            add_table(doc, ("Area", "Current evidence"), VALIDATION_ROWS, [1.45, 5.0])

    algorithm_heading = doc.add_heading("Algorithm 1: one independent M run", level=1)
    algorithm_heading.paragraph_format.page_break_before = True
    doc.add_paragraph(
        "The optimization unit of KODAMA is one independent run inside the M-run ensemble. "
        "The pseudocode below isolates that unit: it starts from one landmark label vector, "
        "uses the previous CV predictions to propose label changes, evaluates the proposal "
        "with exactly one CV pass, and returns the best label vector found by that run."
    )
    add_pseudocode_block(doc, SINGLE_RUN_PSEUDOCODE_LINES)

    algorithm_heading = doc.add_heading("Algorithm 2: KODAMA.matrix ensemble pseudocode", level=1)
    algorithm_heading.paragraph_format.page_break_before = True
    doc.add_paragraph(
        "The matrix-level routine repeats Algorithm 1 independently M times, projects each "
        "optimized run back to all samples, and builds the final KODAMA-corrected graph from "
        "agreement across the optimized label vectors."
    )
    add_pseudocode_block(doc, KODAMA_PSEUDOCODE_LINES)

    algorithm_heading = doc.add_heading("Algorithm 3: KODAMA.matrix.graph pseudocode", level=1)
    algorithm_heading.paragraph_format.page_break_before = True
    doc.add_paragraph(
        "The graph-input routine keeps the applicable classifier-specific label-evolution policy "
        "while replacing global neighbor-search construction with caller-supplied indices and "
        "distances. KNN voting, label projection, and agreement correction use the supplied graph. "
        "When the original feature matrix is absent, self-tuning Laplacian features provide "
        "selection and initialization geometry and, for PLS-LDA, the classifier representation."
    )
    add_pseudocode_block(doc, GRAPH_INPUT_PSEUDOCODE_LINES)

    doc.add_heading("References", level=1)
    for ref in REFERENCES:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(ref)
        set_run_font(run, 9.2)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("kodama-cpp JMLR MLOSS manuscript")
    doc.save(MANUSCRIPT)


def build_self_review() -> None:
    doc = Document()
    apply_styles(doc)
    apply_compact_memo_styles(doc)
    add_title(
        doc,
        "JMLR MLOSS re-review: kodama-cpp",
        "Assessment after revision of the evidence and reproducibility claims",
        "Review date: 8 August 2026",
    )
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(
        "This submission presents a substantial standalone implementation of KODAMA rather than a new "
        "learning objective. The manuscript clearly distinguishes continuity with the published KODAMA "
        "method from the new C++17 systems contribution: float32 state ownership, CPU/CUDA/Metal backends, "
        "label-aware SIMPLS plus latent-space LDA, package-owned neighbor search, supplied-graph entry points, "
        "and thin R/Python bindings. The algorithm is reproducible from the equations and pseudocode: proposal "
        "generation, one-CV-pass-per-cycle evaluation, acceptance score, cooling rule, independent M runs, "
        "landmark and splitting semantics, projection, and agreement-graph correction are all specified."
    )
    doc.add_paragraph(
        "The empirical section now includes paired three-seed classic-versus-KODAMA CUDA results across an "
        "11-dataset archive. The dataset is the inferential unit. For PLS-LDA plus UMAP, 6 of 10 complete "
        "datasets improve, the median silhouette change is +0.025, the 100,000-resample bootstrap interval is "
        "[-0.067, +0.084], and the Holm-adjusted exact Wilcoxon p-value is 0.557. KNN plus openTSNE is worse "
        "on all 10 datasets (Holm-adjusted p=0.0078). These post-specified analyses are labeled exploratory; "
        "the manuscript does not claim universal visualization improvement."
    )
    doc.add_paragraph(
        "A new four-panel ImageNet figure makes the software output concrete, and the supplement now reports "
        "preliminary KODAMA core timing for every completed dataset/classifier cell. The figure is appropriately "
        "described as illustrative: its archived classic and KODAMA renders do not yet document one identical "
        "plotting-index vector. The frozen rerun plan requires matched plotting indices and fills the missing "
        "FlowRepository KNN and ImageNet PLS-LDA timing cells."
    )
    doc.add_paragraph(
        "A local Apple Metal scheduler study additionally shows that the optimal number of independent worker "
        "lanes depends on the retained fold workspace: three lanes were selected for raw COIL20 and four for "
        "MetRef. This supports memory-aware automatic scheduling, but the timing is preliminary because thermal "
        "state and run order were not controlled as rigorously as the planned confirmatory experiment."
    )
    doc.add_paragraph(
        "A new paired CPU4/Metal screen removes graph-search confounding by supplying one identical CPU graph "
        "to both optimizers for each dataset and seed. Across three seeds at M=20/Tcycle=20, median Metal "
        "speedups range from 0.98x for MetRef PLS-LDA to 2.33x for USPS PLS-LDA. The study properly retains "
        "the raw COIL20 50-component PLS-LDA cell as censored after it crossed the local feasibility budget. "
        "The wrapper now also invalidates stale native objects when the selected core archive or public headers "
        "change. These are useful engineering revisions, but the results remain dirty-worktree evidence."
    )
    doc.add_paragraph(
        "Two further local corrections strengthen the numerical implementation. CPU PLS-LDA now streams class "
        "sums and the latent Gram matrix instead of retaining full score matrices; TabulaMuris PLSLDACV improved "
        "by 1.47x with identical predictions and a 0.587 peak-memory ratio. An overflow-only SIMPLS power-norm "
        "guard restored all 50 requested components on MNIST70k: CPU accuracy changed from an erroneous 0.112529 "
        "fallback to 0.866457, while Metal changed from failure to 0.866286 in a 2.296-second median. Normal CPU, "
        "AddressSanitizer, physical Metal, and a compact synthetic regression pass. Frozen CUDA confirmation is "
        "still required."
    )
    doc.add_paragraph(
        "Cycle 15 adds a full shared-graph MetRef run at M=Tcycle=100 and clean wrapper checks. Metal "
        "accelerated KNN optimization from 2.729 to 1.609 seconds, but four-lane Metal PLS-LDA required "
        "651.767 seconds versus 335.734 seconds on CPU4. Both PLS-LDA backends reached median CV accuracy "
        "0.993893 and strong truth-label silhouette (0.8598/0.8656), while KNN selected two classes. A "
        "single-lane Metal proposal was rejected after two runs required 46.8 seconds. This strengthens the "
        "paper because it reports an unfavorable systems result and avoids a benchmark-specific scheduler rule."
    )
    doc.add_paragraph(
        "Cycle 16 adds the corresponding full shared-graph USPS controls. Metal accelerated KNN "
        "optimization from 12.451 to 7.888 seconds, but PLS-LDA required 1105.998 seconds on Metal "
        "versus 707.057 seconds on CPU4. KODAMA truth-label silhouettes were 0.4550/0.4555 for KNN "
        "and 0.4227/0.4174 for PLS-LDA, compared with approximately 0.13 for classic UMAP. The "
        "result reinforces the need for stage-resolved, classifier-specific performance claims."
    )
    doc.add_paragraph(
        "Cycle 22 added the first complete shared-graph SatImage control at M=Tcycle=100. A later "
        "lifecycle audit showed that its Metal result was dominated by one-slot fold-cache eviction "
        "and conservative independent-run scheduling. The corrected implementation reduced KNN "
        "from 3.216 to 2.442 seconds and PLS-LDA from 331.781 to 171.675 seconds. KNN is now 1.21x "
        "faster than CPU4; Metal PLS-LDA is 1.93x faster than its own pre-fix path but remains 3.04x "
        "slower than CPU4 for this 36-variable case. CV and class-count summaries remain close and "
        "no run collapses. Both KODAMA truth-label silhouettes remain below classic UMAP. The same "
        "coverage cycle raises CPU line/branch coverage to 80.48%/69.33% by testing the public binary "
        "UMAP graph route. Production mathematics is unchanged."
    )
    doc.add_paragraph(
        "Cycle 26 separates Metal graph performance from low-dimensional PLS-LDA orchestration. "
        "A native Apple M3 fixture builds the exact graph in 0.0448 seconds versus 0.2493 seconds "
        "with four-core HNSW at recall 1.000; KNN CV is 25.0x faster on Metal. In contrast, the "
        "small PLS-LDA CV probe is 1.52x slower. Removing nested fold workers from independent M "
        "lanes reduces BreastCancerDiagnostic M=Tcycle=100 PLS-LDA from 109.557 to 105.285 seconds "
        "with unchanged core diagnostics, but CPU4 remains faster at 6.990 seconds. A one-command "
        "SIMPLS candidate was slower and was removed. These negative results delimit the Metal "
        "crossover without altering the classifier or selecting a backend by dataset."
    )

    doc.add_heading("Disposition of Previous Comments", level=1)
    disposition_rows = [
        ("Uncertainty and paired inference", "Addressed", "Dataset-level seed medians, bootstrap intervals, exact paired Wilcoxon tests, sign tests, and Holm adjustment are now reported."),
        ("Primary endpoint", "Addressed for a future confirmatory study", "PLS-LDA UMAP silhouette is predeclared as primary; trustworthiness and Preserve@30 are secondary. Existing archive analyses remain explicitly exploratory."),
        ("Overstatement of silhouette", "Addressed", "The text now describes reference-label compactness and label-aligned structure, not recovery of more information."),
        ("Unsupported Metal timing", "Addressed", "The historical 425x row was removed because warm-up and synchronization boundaries were not retained."),
        ("Timing boundaries", "Addressed", "The protocol defines wall-clock call boundaries, synchronization, allocation and transfer inclusion, warm-up handling, and thread reporting."),
        ("Release identity and HPC provenance", "Open", "Public commit identities are recorded, but no immutable tag exists and uploaded manifests contain git_commit=NA; a clean tagged confirmatory rerun is still required."),
        ("Coverage and CUDA audit", "Partially addressed", "A direct public-entry-point map records physical Metal and current-source CUDA execution. Binary UMAP, matrix orchestration, and resident-IVF ownership/error contracts raise measured CPU line/branch coverage to 80.48%/69.33%, but this remains below the MLOSS target and public CI still does not exercise CUDA hardware."),
        ("Independent adoption", "Open", "Predecessor-package use is documented separately, but the new library still lacks verifiable independent adoption evidence."),
        ("Python source packaging", "Addressed", "The tested binding is intentionally versioned with the core tree, and the documentation no longer promises a separate unpublished repository."),
        ("Metal worker scheduling", "Addressed locally", "A bounded 16-entry resident fold cache and memory-aware automatic independent-run concurrency reduce full MetRef PLS-LDA from 631.415 to 312.029 seconds and SatImage from 331.781 to 171.675 seconds. Nested fold workers are now suppressed inside independent M lanes. The retained low-p BreastCancerDiagnostic crossover is reported explicitly. Thermally controlled tagged replication and matched CUDA confirmation remain open."),
        ("Backend-comparison confounding", "Addressed locally", "A shared-graph mode now feeds one identical graph to CPU4 and Metal and is reported separately from native end-to-end timing; the tagged CUDA counterpart remains open."),
        ("R wrapper stale relinking", "Addressed locally", "A checksum signature over the selected archive and public headers forces both Rcpp objects and the shared library to rebuild after a core change; unchanged installs remain incremental."),
        ("CPU PLS-LDA score storage", "Addressed locally", "The CPU path streams class sums and latent Gram statistics; four-dataset predictions match and the large-n case is faster and smaller."),
        ("Metal PLS-LDA score transfers", "Addressed locally", "Training scores remain device-resident, a label-aware kernel forms class sums, MPS forms T'T, validation scoring is on-device, and five-dataset tests retain the small-data crossover."),
        ("Large-sample SIMPLS range", "Addressed locally", "MNIST70k and a compact overflow regression pass on CPU/Metal; the frozen CUDA confirmation remains open."),
        ("Full-cycle local wrapper evidence", "Addressed locally", "R CMD check passed with no errors or warnings and one new-submission NOTE; an installed-core Python wheel passed eight CPU/physical-Metal tests; full MetRef, USPS, ImageSegmentation, PageBlocks, PenDigits, SatImage, and paired COIL20 controls are retained. Landmark timing is exposed separately."),
    ]
    add_table(doc, ("Comment", "Status", "Revision or remaining action"), disposition_rows, [1.8, 1.15, 4.0], font_size=8.0)

    doc.add_heading("Major Comments", level=1)
    add_numbered(
        doc,
        [
            "The reviewed software is still a release candidate rather than an immutable reviewed release. There is no public v0.1.0 tag or GitHub release, and no archived source snapshot. Freeze the exact reviewed commit, create the source archive required by MLOSS, publish its checksum, and cite that identity consistently. A DOI is useful but not itself mandatory.",
            "The new HPC archive is broad, but its KODAMA run manifests record git_commit as NA. Therefore the central 11-dataset result cannot yet be tied unambiguously to the reviewed source. Reproduce or formally identify these runs from a tagged core and wrapper commit, archive the benchmark configuration and hardware record, and preserve the same seed-level CSV files.",
            "The ImageNet panel is useful, but the archived KODAMA render explicitly uses a fixed 250,000-point plotting subset while the retained classic plot does not record the same index vector. Keep the current figure qualitative. In the definitive rerun, save one stratified plotting-index vector and apply it, the same class palette, point size, axis treatment, and rasterization policy to all four layouts.",
            "The all-dataset timing table is preliminary and incomplete: FlowRepository KNN and ImageNet PLS-LDA are absent, and the archive is not tied to a release commit. Complete both cells or report a predeclared resource-limit failure, then provide graph, optimization, projection, correction, visualization, total wall time, host RAM, and device memory from the same tagged benchmark protocol.",
            "CPU source coverage is 80.48% by line and 69.33% by branch, still below the MLOSS expectation of coverage close to 100%. Public binary UMAP, matrix orchestration, invalid-input, unavailable-backend, fixed-seed replay, monotone best-trace, guarded one-class, anti-collapse, shake-recovery, constant-predictor PLS-LDA fallback, and resident-IVF lifecycle/error contracts are now exercised. Graph materialization and matrix/visualization integration are covered on CPU and physical Metal; accelerator implementation branches remain outside the portable denominator. CUDA is not exercised by public CI; archive hardware logs or add a maintained GPU runner.",
            "The cover letter documents substantial use of the predecessor KODAMA R package, but this is not evidence of an active community around kodama-cpp. The new repository currently has no stars, forks, releases, public issues, or independently documented adopters. Obtain and cite verifiable beta-user, downstream-package, issue, or external-reproduction evidence rather than transferring predecessor adoption to the new implementation.",
            "The Python binding is versioned and tested with the reviewed core source. Before submission, archive the exact source release and test installation from that archive in a clean environment.",
            "Raw high-dimensional PLS-LDA is now characterized at kernel level: the COIL20 predictor-width study demonstrates a CPU/Metal crossover and exposed a stale-fold residency bug that is now fixed and regression-tested. Metal also keeps projected scores resident, uses a label-aware class-sum kernel and MPS T'T, and scores validation rows on-device. In the current full protocol, raw COIL20 CPU4 at M=Tcycle=100 exceeded 300 seconds before the first four independent runs completed and is explicitly censored. Complete release-tagged CPU4/CUDA scaling with the newly separated landmark timing, peak memory, and the distinct CUDA cached-Gram versus Metal retained-matrix formulations is still required.",
            "The robust SIMPLS norm fixes a severe local MNIST70k failure and is supported by replicated CPU/Metal results, but it has not been exercised on the frozen CUDA artifact. Run fixed-50-component MNIST70k PLSLDACV and the compact overflow regression on tagged CUDA, record selected components and predictions, and retain any bounded float32 difference rather than assuming cross-backend parity.",
        ],
    )

    doc.add_heading("Minor Comments", level=1)
    add_numbered(
        doc,
        [
            "Complete the author-only cover-letter fields: coauthor consent, funding, competing interests, and conflict-free editor/reviewer suggestions. Refresh all dated repository and download metrics immediately before submission.",
            "Use KODAMA R 2.4.1/2.4 as the only predecessor timing baseline. Keep the comparison classifier-matched where possible, use M=100/Tcycle=100, and state explicitly that revised proposal dynamics prevent trajectory-level parity.",
            "Document the native CUDA graph-builder limit k <= 256 prominently and report unsupported requested graph sizes as missing rather than silently truncating them.",
            "Retain the explicit distinction between internal CV accuracy, the search score, and external-label diagnostics. The new visualization table makes the non-equivalence especially clear and should not be simplified into a claim that maximizing CV accuracy necessarily recovers reference classes.",
            "Clarify that classic rows are full fastEmbedR package calls whereas KODAMA rows use the library's pinned standalone embedding port on a corrected graph. The mathematical lineage and parameters are matched, but the software-call boundaries and total runtimes are intentionally different.",
            "Keep the four-page paper self-contained while treating the long technical document as a supplement. The present layout now satisfies the four-description-page limit, with references beginning on page 5.",
            "The current technical supplement is comprehensive but long. Remove duplicated narrative, retain auditable implementation details and complete tables, and move operational logs into the release archive.",
            "The inherited CUDA helper symbols and environment controls now use KODAMA-owned prefixes and a source/link audit prevents regression. Close the remaining accelerator-specific part by inspecting the tagged CUDA symbol table and loading kodama-cpp beside fastEmbedR in one process.",
            "Keep KODAMA R 2.4.1/2.4 as the sole predecessor. The retained 16 July table and a later 6 August diagnostic have different single-run timings; identify the final tagged protocol unambiguously and report replicated within-host medians rather than combining or selecting between those preliminary rows.",
        ],
    )

    doc.add_heading("Strengths", level=1)
    add_bullets(
        doc,
        [
            "The distinction between raw CV accuracy, proposal acceptance score, and external diagnostics is unusually clear and prevents circular quality claims.",
            "The single-run and ensemble pseudocode make the stochastic optimization independently implementable.",
            "Backend identity is strict and testable; unavailable accelerators do not masquerade as successful CPU execution.",
            "The evidence matrix links implementation claims to source locations, tests, and measured results.",
            "Cross-platform GitHub Actions, measured CPU coverage, a frozen API test, provenance audit, and hosted documentation substantially strengthen MLOSS readiness.",
            "The manuscript retains unfavorable datasets and now quantifies that KODAMA correction is classifier-, embedding-, and dataset-dependent.",
            "The paired HPC analysis spans image, cytometry, metabolomics, and single-cell representations and avoids using unmatched runs in comparative summaries.",
        ],
    )

    doc.add_heading("Reviewer Recommendation", level=1)
    doc.add_paragraph(
        "Recommendation: Major Revision. The software is suitable in scope and technically promising for JMLR "
        "MLOSS, and the four-page description now contains a substantive multi-dataset result. However, the "
        "reviewed artifact is not yet frozen, the new HPC runs lack a recorded git commit, ImageNet plotting "
        "indices are not yet matched across all four panels, the timing matrix has two missing cells, coverage remains "
        "substantially below the track's stated expectation, active adoption of the new library has not yet "
        "been demonstrated, complete high-dimensional PLS-LDA pipeline scaling is not yet characterized, the CUDA scheduler "
        "and robust-norm experiments remain outstanding. "
        "The earlier statistical-interpretation concern is now adequately addressed for an exploratory study."
    )
    doc.add_paragraph(
        "The full MetRef, USPS, ImageSegmentation, PageBlocks, PenDigits, and SatImage results, plus paired COIL20 KNN and censored raw PLS-LDA, do not change this recommendation. They close a useful local "
        "full-cycle gap, but also demonstrate that native Metal availability is not a universal acceleration "
        "claim: Metal can accelerate KNN while CPU4 is faster for complete PLS-LDA in several matrix regimes. The main "
        "text and supplement must preserve this classifier- and workload-specific interpretation. PenDigits, COIL20, and SatImage further show that high internal CV accuracy can coexist with a degraded final visualization."
    )
    doc.add_paragraph(
        "I would encourage resubmission after these software-release and evidence issues are addressed. I do not "
        "request a new KODAMA objective or an algorithmic redesign: the principal work is to freeze and test the "
        "release, bind the HPC archive to that release, strengthen coverage and accelerator auditability, "
        "archive and test the bundled Python wrapper, and establish independent usage evidence."
    )
    doc.save(SELF_REVIEW)


def tex_escape(text: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "^": r"\textasciicircum{}",
        "~": r"\textasciitilde{}",
        ">": r"\textgreater{}",
        "<": r"\textless{}",
    }
    out = text
    for key, value in replacements.items():
        out = out.replace(key, value)
    return out


def build_tex() -> None:
    body = []
    body.append(r"\documentclass[twoside,11pt]{article}")
    body.append(r"\usepackage{jmlr2e}")
    body.append(r"\usepackage{booktabs}")
    body.append(r"\usepackage{tabularx}")
    body.append(r"\usepackage{amsmath}")
    body.append(r"\usepackage{graphicx}")
    body.append(r"\usepackage{url}")
    body.append("")
    body.append(r"\jmlrheading{1}{2026}{1-XX}{1/26}{7/26}{26-0000}{Kassim et al.}")
    body.append(r"\ShortHeadings{kodama-cpp}{Kassim et al.}")
    body.append(r"\firstpageno{1}")
    body.append("")
    body.append(r"\begin{document}")
    body.append(r"\title{kodama-cpp: Cross-Validated Accuracy Maximization on CPU, CUDA, and Apple Metal}")
    body.append(r"\author{%")
    for name, markers, email in AUTHORS:
        markers_tex = markers.replace("†", r"\dagger")
        body.append(
            r"\name "
            + tex_escape(name)
            + r"$^{"
            + markers_tex
            + r"}$ \email "
            + tex_escape(email)
            + r" \\"
        )
    for marker, affiliation in AFFILIATIONS:
        affiliation_tex = tex_escape(affiliation).replace('"Ugo Schiff"', r"``Ugo Schiff''")
        if marker == "2":
            affiliation_tex = affiliation_tex.replace(
                "Sciences, Institute",
                r"Sciences, \\ \addr Institute",
            ).replace(
                "Medicine (IDM), University",
                r"Medicine (IDM), \\ \addr University",
            )
        body.append(r"\addr $^{" + marker + r"}$ " + affiliation_tex + r" \\")
    body.append(r"\addr $^{*}$ Moussa Kassim and Martin Ocharo contributed equally. \\")
    body.append(
        r"\addr $^{\dagger}$ Leonardo Tenori and Stefano Cacciatore are co-corresponding authors. \\"
    )
    body.append(r"}")
    body.append(r"\editor{To be assigned}")
    body.append(r"\maketitle")
    body.append(r"\begin{abstract}")
    body.append(tex_escape(ABSTRACT))
    body.append(r"\end{abstract}")
    body.append(r"\begin{keywords}")
    body.append("KODAMA, unsupervised learning, cross-validation, heterogeneous computing, CUDA, Metal, C++ library, nearest neighbors, PLS-LDA")
    body.append(r"\end{keywords}")
    body.append("")
    for heading, paragraphs in SECTIONS:
        title = heading.split(". ", 1)[1] if ". " in heading else heading
        body.append(r"\section{" + tex_escape(title) + "}")
        for paragraph in paragraphs:
            body.append(tex_escape(paragraph))
            body.append("")
        if heading == "1. Introduction":
            body.append(r"\begin{table}[h]")
            body.append(r"\caption{Main contributions of kodama-cpp.}")
            body.append(r"\small")
            body.append(r"\begin{tabularx}{\linewidth}{p{0.24\linewidth}X}")
            body.append(r"\toprule")
            body.append(r"Contribution & Role in kodama-cpp \\")
            body.append(r"\midrule")
            for contribution, detail in CONTRIBUTION_ROWS:
                body.append(f"{tex_escape(contribution)} & {tex_escape(detail)} \\\\")
            body.append(r"\bottomrule")
            body.append(r"\end{tabularx}")
            body.append(r"\end{table}")
            body.append("")
        if heading == "2. KODAMA objective":
            body.append(r"\begin{equation}")
            body.append(r"c^\star = \arg\max_c \; A(c; X, F, \Pi),")
            body.append(r"\end{equation}")
            body.append("where A is the held-out accuracy of classifier family F under fold assignment Pi.")
            body.append("")
            body.append(r"\subsection{Landmark-projection strategy}")
            for paragraph in LANDMARKING_NOVELTY_PARAGRAPHS:
                body.append(tex_escape(paragraph))
                body.append("")
            body.append(r"\begin{table}[h]")
            body.append(r"\caption{Validation of the landmark-projection implementation.}")
            body.append(r"\small")
            body.append(r"\begin{tabularx}{\linewidth}{p{0.22\linewidth}X}")
            body.append(r"\toprule")
            body.append(r"Validation dimension & Current manuscript evidence \\")
            body.append(r"\midrule")
            for dimension, evidence in LANDMARK_VALIDATION_ROWS:
                body.append(f"{tex_escape(dimension)} & {tex_escape(evidence)} \\\\")
            body.append(r"\bottomrule")
            body.append(r"\end{tabularx}")
            body.append(r"\end{table}")
            body.append("")
            body.append(r"\subsection{Label-vector search mechanics}")
            for paragraph in LABEL_SEARCH_DETAIL_PARAGRAPHS:
                body.append(tex_escape(paragraph))
                body.append("")
            body.append(r"\subsection{Raw accuracy, acceptance score, and diagnostics}")
            for paragraph in SCORE_INTERPRETATION_PARAGRAPHS:
                body.append(tex_escape(paragraph))
                body.append("")
            body.append(r"\subsection{Reproducible KODAMA.matrix procedure}")
            body.append(tex_escape(
                "The public KODAMA.matrix routine is reproducible from the following option definitions and run-level procedure. "
                "The description names the C++ options because the R and Python wrappers should expose the same contract."
            ))
            body.append("")
            body.append(r"\begin{table}[h]")
            body.append(r"\caption{KODAMA.matrix options that determine the optimization.}")
            body.append(r"\small")
            body.append(r"\begin{tabularx}{\linewidth}{p{0.18\linewidth}X}")
            body.append(r"\toprule")
            body.append(r"Option & Role in the algorithm \\")
            body.append(r"\midrule")
            for option, role in KODAMA_PARAMETER_ROWS:
                body.append(f"{tex_escape(option)} & {tex_escape(role)} \\\\")
            body.append(r"\bottomrule")
            body.append(r"\end{tabularx}")
            body.append(r"\end{table}")
            body.append("")
            body.append(r"\subsubsection{Run-level optimization rule}")
            body.append(r"\begin{enumerate}")
            for item in REPRODUCIBLE_ALGORITHM_STEPS:
                body.append(r"\item " + tex_escape(item))
            body.append(r"\end{enumerate}")
            body.append("")
            body.append(r"\subsubsection{Final graph and dissimilarity}")
            for paragraph in GRAPH_CONSTRUCTION_PARAGRAPHS:
                body.append(tex_escape(paragraph))
                body.append("")
            body.append(r"\subsection{Main algorithm pseudocode}")
            body.append(tex_escape(
                "The compact algorithm below connects landmark selection, guided label evolution, "
                "projection, and agreement-graph construction. Each inner cycle contributes exactly "
                "one new CV evaluation; the initial state contributes one baseline CV evaluation."
            ))
            body.append("")
            body.append(r"\begin{scriptsize}")
            body.append(r"\begin{verbatim}")
            body.extend(MAIN_ALGORITHM_PSEUDOCODE_LINES)
            body.append(r"\end{verbatim}")
            body.append(r"\end{scriptsize}")
            body.append("")
        if heading == "3. Implementation architecture":
            body.append(r"\begin{figure}[h]")
            body.append(r"\centering")
            body.append(r"\includegraphics[width=\linewidth]{kodama_cpp_architecture.png}")
            body.append(r"\caption{Architecture of kodama-cpp. The C++17 core owns the numerical kernels, CPU, CUDA, and Apple Metal backends, graph utilities, and typed outputs; R and Python remain thin wrappers.}")
            body.append(r"\end{figure}")
            body.append("")
            body.append(r"\subsection{Implementation novelty}")
            for paragraph in IMPLEMENTATION_NOVELTY_PARAGRAPHS:
                body.append(tex_escape(paragraph))
                body.append("")
            body.append(r"\subsection{CUDA backend}")
            for paragraph in CUDA_BACKEND_PARAGRAPHS:
                body.append(tex_escape(paragraph))
                body.append("")
            body.append(r"\subsection{Apple Metal backend}")
            for paragraph in METAL_BACKEND_PARAGRAPHS:
                body.append(tex_escape(paragraph))
                body.append("")
            body.append(r"\subsection{Shared accelerator contract}")
            body.append(r"\begin{table}[h]")
            body.append(r"\caption{Backend-specific execution under a shared KODAMA contract.}")
            body.append(r"\small")
            body.append(r"\begin{tabularx}{\linewidth}{p{0.18\linewidth}X X}")
            body.append(r"\toprule")
            body.append(r"Invariant & CUDA realization & Metal realization \\")
            body.append(r"\midrule")
            for invariant, cuda_impl, metal_impl in BACKEND_CONTRACT_ROWS:
                body.append(f"{tex_escape(invariant)} & {tex_escape(cuda_impl)} & {tex_escape(metal_impl)} \\\\")
            body.append(r"\bottomrule")
            body.append(r"\end{tabularx}")
            body.append(r"\end{table}")
            body.append("")
            body.append(tex_escape(
                "The public surface groups cross-validation, core optimization, graph preparation, "
                "matrix and graph-input KODAMA, PCA, visualization, and graph utilities. The installed "
                "C++ headers and generated R/Python references are the authoritative signature list."
            ))
            body.append("")
            body.append(r"\subsection{Graph-first input contract}")
            body.append(tex_escape(
                "KODAMA.graph returns a matrix- or external-handle-backed graph, backend-matched PCA "
                "starts, metadata, and timings, but never the source matrix. Raw features remain "
                "caller-owned when supplied separately."
            ))
            body.append("")
            body.append(r"\begin{table}[h]")
            body.append(r"\caption{Graph-first KODAMA.matrix input contract.}")
            body.append(r"\small")
            body.append(r"\begin{tabularx}{\linewidth}{p{0.12\linewidth}X X X p{0.10\linewidth}}")
            body.append(r"\toprule")
            body.append(r"Input & Working geometry & KNN evolution & PLS--LDA evolution & Graph builds \\")
            body.append(r"\midrule")
            for input_form, geometry, knn_path, pls_path, builds in GRAPH_INPUT_CONTRACT_ROWS:
                body.append(
                    f"{tex_escape(input_form)} & {tex_escape(geometry)} & {tex_escape(knn_path)} & "
                    f"{tex_escape(pls_path)} & {tex_escape(builds)} \\\\"
                )
            body.append(r"\bottomrule")
            body.append(r"\end{tabularx}")
            body.append(r"\end{table}")
            body.append("")
            body.append(r"\begin{scriptsize}")
            body.append(r"\begin{verbatim}")
            body.extend(GRAPH_PREPARATION_PSEUDOCODE_LINES)
            body.append(r"\end{verbatim}")
            body.append(r"\end{scriptsize}")
            body.append("")
            body.append(tex_escape(GRAPH_API_SMOKE_NOTE))
            body.append("")
            body.append(r"\subsection{Computational scaling}")
            body.append(r"\begin{table}[h]")
            body.append(r"\caption{Dominant computational work in the KODAMA pipeline.}")
            body.append(r"\small")
            body.append(r"\begin{tabularx}{\linewidth}{p{0.18\linewidth}X X}")
            body.append(r"\toprule")
            body.append(r"Component & Dominant work & Implementation implication \\")
            body.append(r"\midrule")
            for component, work, implication in COMPLEXITY_ROWS:
                body.append(f"{tex_escape(component)} & {tex_escape(work)} & {tex_escape(implication)} \\\\")
            body.append(r"\bottomrule")
            body.append(r"\end{tabularx}")
            body.append(r"\end{table}")
            body.append("")
            body.append(r"\subsection{Graph-input KODAMA validation boundary}")
            body.append(r"\begin{table}[h]")
            body.append(r"\caption{Interpretation of graph-input KODAMA paths.}")
            body.append(r"\small")
            body.append(r"\begin{tabularx}{\linewidth}{p{0.18\linewidth}X X}")
            body.append(r"\toprule")
            body.append(r"Path & Implementation & Interpretation \\")
            body.append(r"\midrule")
            for path, implementation, interpretation in GRAPH_INPUT_VALIDATION_ROWS:
                body.append(f"{tex_escape(path)} & {tex_escape(implementation)} & {tex_escape(interpretation)} \\\\")
            body.append(r"\bottomrule")
            body.append(r"\end{tabularx}")
            body.append(r"\end{table}")
            body.append("")
            body.append(r"\subsection{KODAMA directly from a KNN matrix}")
            for paragraph in GRAPH_INPUT_PROCEDURE_PARAGRAPHS:
                body.append(tex_escape(paragraph))
                body.append("")
            body.append(tex_escape(
                "Executable C++, R, and Python calls for every input form are maintained in the "
                "public API tutorial so that wrapper syntax does not duplicate the mathematical "
                "contract in the publication supplement."
            ))
            body.append("")
        if heading == "4. Relationship to public KODAMA literature":
            body.append(r"\begin{table}[h]")
            body.append(r"\caption{Compatibility with the public KODAMA R interface and literature.}")
            body.append(r"\small")
            body.append(r"\begin{tabularx}{\linewidth}{p{0.17\linewidth}X X}")
            body.append(r"\toprule")
            body.append(r"Aspect & Public R/literature contract & kodama-cpp \\")
            body.append(r"\midrule")
            for aspect, prior, current in COMPATIBILITY_ROWS:
                body.append(f"{tex_escape(aspect)} & {tex_escape(prior)} & {tex_escape(current)} \\\\")
            body.append(r"\bottomrule")
            body.append(r"\end{tabularx}")
            body.append(r"\end{table}")
            body.append("")
            body.append(r"\subsection{Relationship to semi-supervised learning and cluster validation}")
            for paragraph in RELATED_WORK_PARAGRAPHS:
                body.append(tex_escape(paragraph))
                body.append("")
            body.append(r"\begin{table}[h]")
            body.append(r"\caption{Literature positioning of KODAMA.}")
            body.append(r"\small")
            body.append(r"\begin{tabularx}{\linewidth}{p{0.22\linewidth}X X}")
            body.append(r"\toprule")
            body.append(r"Literature area & Typical goal & Role of KODAMA \\")
            body.append(r"\midrule")
            for area, goal, role in RELATED_WORK_ROWS:
                body.append(f"{tex_escape(area)} & {tex_escape(goal)} & {tex_escape(role)} \\\\")
            body.append(r"\bottomrule")
            body.append(r"\end{tabularx}")
            body.append(r"\end{table}")
            body.append("")
        if heading == "6. Availability and reproducibility":
            body.append(r"\subsection{Installation, wrappers, and licensing}")
            body.append(
                r"CMake configures the standalone core for CPU, native Apple Metal, or NVIDIA CUDA. "
                r"Thin R and Python packages link to the installed core. Build and compatibility "
                r"records are maintained in \path{README.md}, \path{docs/getting-started.md}, and "
                r"\path{docs/r-api.md}. Licensing boundaries are defined by \path{PROVENANCE.md}, "
                r"\path{THIRD_PARTY_NOTICES.md}, SPDX headers, retained license texts, and the "
                r"automated audit."
            )
            body.append("")
        if heading == "5. Evaluation":
            body.append(r"\subsection{Audited implementation improvements}")
            body.append(
                "Accepted changes preserve the KODAMA objective while correcting fold isolation, "
                "index-base ownership, float32 numerical range, graph residency, repeated fold work, "
                "and wrapper conversion. CPU streams PLS-LDA sufficient statistics; CUDA derives them "
                "from resident class sums and fold Gram matrices; Metal now uses a label-aware class-sum "
                "kernel, MPS T'T, and on-device validation scoring. Leakage, numerical, backend, "
                "and wrapper regressions cover these contracts. The dated repository audit maps every "
                "claim to source, tests, and retained benchmark files."
            )
            body.append("")
            body.append(r"\subsection{Benchmark protocol and data coverage}")
            body.append(r"\begin{table}[h]")
            body.append(r"\caption{Benchmark protocol for the current evaluation.}")
            body.append(r"\begin{tabularx}{\linewidth}{p{0.22\linewidth}X}")
            body.append(r"\toprule")
            body.append(r"Item & Value \\")
            body.append(r"\midrule")
            for item, value in BENCHMARK_PROTOCOL_ROWS:
                body.append(f"{tex_escape(item)} & {tex_escape(value)} \\\\")
            body.append(r"\bottomrule")
            body.append(r"\end{tabularx}")
            body.append(r"\end{table}")
            body.append("")
            body.append(r"\subsection{Evaluation guardrails and ablations}")
            body.append(tex_escape(
                "CV accuracy is the internal search objective, never an external truth score. ARI, "
                "silhouette, trustworthiness, local purity, and active-class count are computed only "
                "after optimization. Classic and KODAMA embeddings are paired by dataset and seed, "
                "timing stages are reported separately, and failures remain explicit. The frozen "
                "release protocol predeclares comparisons for search evolution, classifier, graph "
                "correction, Tcycle, M, landmark/splitting controls, backend, and wrapper parity; the "
                "complete ablation matrix is archived with the benchmark driver."
            ))
            body.append("")
            body.append(r"\subsection{Dependency-free CPU and Apple Metal validation}")
            for paragraph in METAL_VALIDATION_PARAGRAPHS:
                body.append(tex_escape(paragraph))
                body.append("")
            body.append(r"\begin{table}[h]")
            body.append(r"\caption{Dependency-free CPU and native Apple Metal validation. Quality entries are CPU/Metal where two values are shown.}")
            body.append(r"\small")
            body.append(r"\footnotesize")
            body.append(r"\begin{tabularx}{\linewidth}{@{}p{0.08\linewidth}p{0.18\linewidth}p{0.08\linewidth}p{0.08\linewidth}p{0.08\linewidth}X@{}}")
            body.append(r"\toprule")
            body.append(r"Dataset & Scope & CPU s & Metal s & Speedup & Quality \\")
            body.append(r"\midrule")
            for dataset, scope, cpu_s, metal_s, speedup, quality in METAL_VALIDATION_ROWS:
                body.append(
                    f"{tex_escape(dataset)} & {tex_escape(scope)} & {tex_escape(cpu_s)} & "
                    f"{tex_escape(metal_s)} & {tex_escape(speedup)} & {tex_escape(quality)} \\\\"
                )
            body.append(r"\bottomrule")
            body.append(r"\end{tabularx}")
            body.append(r"\end{table}")
            body.append("")
            body.append(r"\subsection{CUDA workstation validation}")
            for paragraph in PILOT_EXPERIMENT_PARAGRAPHS:
                body.append(tex_escape(paragraph))
                body.append("")
            body.append(r"\begin{table}[h]")
            body.append(r"\caption{Dependency-distilled native CUDA KNN spot checks.}")
            body.append(r"\small")
            body.append(r"\begin{tabular}{lllllll}")
            body.append(r"\toprule")
            body.append(r"Dataset & Path & Accuracy & Seconds & nlist & max nprobe & Pilot recall \\")
            body.append(r"\midrule")
            for dataset, path, accuracy, seconds, nlist, nprobe, recall in NATIVE_CUDA_VALIDATION_ROWS:
                body.append(
                    f"{tex_escape(dataset)} & {tex_escape(path)} & {tex_escape(accuracy)} & "
                    f"{tex_escape(seconds)} & {tex_escape(nlist)} & {tex_escape(nprobe)} & "
                    f"{tex_escape(recall)} \\\\"
                )
            body.append(r"\bottomrule")
            body.append(r"\end{tabular}")
            body.append(r"\end{table}")
            body.append("")
            body.append(r"\subsubsection{Archived kernel and core diagnostics}")
            body.append(tex_escape(
                "Earlier kernel-only and core-optimizer measurements are retained in the benchmark "
                "archive for release-history continuity. They are not repeated here because they use "
                "different timing boundaries from the native end-to-end and shared-graph comparisons "
                "below. Their seed-level outputs remain available for audit."
            ))
            body.append("")
            body.append(r"\subsubsection{MultiCPU versus GPU KODAMA comparison}")
            for paragraph in KODAMA_BACKEND_COMPARISON_PARAGRAPHS:
                body.append(tex_escape(paragraph))
                body.append("")
            body.append(r"\begin{table}[h]")
            body.append(r"\caption{KODAMA core optimizer CPU/CUDA comparison. Scope is reported explicitly because these rows do not include final graph construction.}")
            body.append(r"\small")
            body.append(r"\resizebox{\linewidth}{!}{%")
            body.append(r"\begin{tabular}{llllllll}")
            body.append(r"\toprule")
            body.append(r"Dataset & Classifier & Backend & Seconds & CV acc & ARI & Classes & Scope \\")
            body.append(r"\midrule")
            for dataset, classifier, backend, seconds, acc, ari, classes, scope in KODAMA_BACKEND_COMPARISON_ROWS:
                body.append(
                    f"{tex_escape(dataset)} & {tex_escape(classifier)} & {tex_escape(backend)} & "
                    f"{tex_escape(seconds)} & {tex_escape(acc)} & {tex_escape(ari)} & "
                    f"{tex_escape(classes)} & {tex_escape(scope)} \\\\"
                )
            body.append(r"\bottomrule")
            body.append(r"\end{tabular}%")
            body.append(r"}")
            body.append(r"\end{table}")
            body.append("")
            body.append(r"\subsubsection{KODAMA 2.4.1/2.4 predecessor comparison}")
            for paragraph in HISTORICAL_KNN_PARAGRAPHS:
                body.append(tex_escape(paragraph))
                body.append("")
            body.append(r"\begin{table}[h]")
            body.append(r"\caption{KNN-compatible MetRef predecessor comparison at M = 100 and Tcycle = 100. Speedup is relative to KODAMA R 2.4.1/2.4.}")
            body.append(r"\small")
            body.append(r"\resizebox{\linewidth}{!}{%")
            body.append(r"\begin{tabular}{llllllll}")
            body.append(r"\toprule")
            body.append(r"Implementation & Backend & Workers & Seconds & Speedup & CV best/median & ARI selected/median & Classes selected/median \\")
            body.append(r"\midrule")
            for implementation, backend, workers, seconds, speedup, cv_pair, ari_pair, classes_pair in HISTORICAL_KNN_ROWS:
                body.append(
                    f"{tex_escape(implementation)} & {tex_escape(backend)} & {tex_escape(workers)} & "
                    f"{tex_escape(seconds)} & {tex_escape(speedup)} & {tex_escape(cv_pair)} & "
                    f"{tex_escape(ari_pair)} & {tex_escape(classes_pair)} \\\\"
                )
            body.append(r"\bottomrule")
            body.append(r"\end{tabular}%")
            body.append(r"}")
            body.append(r"\end{table}")
            body.append("")
            for paragraph in HISTORICAL_PLS_PARAGRAPHS:
                body.append(tex_escape(paragraph))
                body.append("")
            body.append(r"\begin{table}[h]")
            body.append(r"\caption{Contextual MetRef PLS comparison at M = 100 and Tcycle = 100. Historical PLS--DA and current SIMPLS plus latent-space LDA are different evaluators; no implementation-only speedup is inferred.}")
            body.append(r"\small")
            body.append(r"\resizebox{\linewidth}{!}{%")
            body.append(r"\begin{tabular}{llllllll}")
            body.append(r"\toprule")
            body.append(r"Implementation & Evaluator & Backend & Workers & Components & Seconds & CV best/median & Post hoc diagnostic \\")
            body.append(r"\midrule")
            for implementation, evaluator, backend, workers, components, seconds, cv_pair, diagnostic in HISTORICAL_PLS_ROWS:
                body.append(
                    f"{tex_escape(implementation)} & {tex_escape(evaluator)} & {tex_escape(backend)} & "
                    f"{tex_escape(workers)} & {tex_escape(components)} & {tex_escape(seconds)} & "
                    f"{tex_escape(cv_pair)} & {tex_escape(diagnostic)} \\\\"
                )
            body.append(r"\bottomrule")
            body.append(r"\end{tabular}%")
            body.append(r"}")
            body.append(r"\end{table}")
            body.append("")
            body.append(r"\begin{table}[h]")
            body.append(r"\caption{Matched current-commit KODAMA.matrix PLS--LDA validation on MetRef with M = 100 and Tcycle = 100. ARI and silhouette use external labels only after optimization.}")
            body.append(r"\small")
            body.append(r"\resizebox{\linewidth}{!}{%")
            body.append(r"\begin{tabular}{lllllll}")
            body.append(r"\toprule")
            body.append(r"Backend/classifier & Wall s & Core s & Max/median CV & Selected ARI/truth sil. & Selected classes & Peak MB \\")
            body.append(r"\midrule")
            for classifier, elapsed, runtime, acc_pair, ari_pair, median_classes, class_range in PILOT_MATRIX_ROWS:
                body.append(
                    f"{tex_escape(classifier)} & {tex_escape(elapsed)} & {tex_escape(runtime)} & "
                    f"{tex_escape(acc_pair)} & {tex_escape(ari_pair)} & {tex_escape(median_classes)} & "
                    f"{tex_escape(class_range)} \\\\"
                )
            body.append(r"\bottomrule")
            body.append(r"\end{tabular}%")
            body.append(r"}")
            body.append(r"\end{table}")
            body.append("")
            body.append(r"\subsubsection{Classic versus KODAMA visualization validation}")
            for paragraph in VISUALIZATION_COMPARISON_PARAGRAPHS:
                body.append(tex_escape(paragraph))
                body.append("")
            if HPC_INFERENCE_ROWS:
                body.append(r"\begin{table}[h]")
                body.append(r"\centering\scriptsize")
                body.append(r"\caption{Exploratory dataset-level inference for silhouette change. Bootstrap intervals use 100,000 dataset resamples; Wilcoxon $p$ values are exact and Holm-adjusted across four contrasts.}")
                body.append(r"\begin{tabular}{llrrrrrr}")
                body.append(r"\toprule")
                body.append(r"Classifier & Embedding & Positive & Median & 95\% CI & Wilcoxon & Holm & Sign \\")
                body.append(r"\midrule")
                for classifier, embedding, positive, median, ci, wilcoxon, holm, sign in HPC_INFERENCE_ROWS:
                    body.append(
                        f"{tex_escape(classifier)} & {tex_escape(embedding)} & {tex_escape(positive)} & "
                        f"{tex_escape(median)} & {tex_escape(ci)} & {tex_escape(wilcoxon)} & "
                        f"{tex_escape(holm)} & {tex_escape(sign)} \\\\"
                    )
                body.append(r"\bottomrule")
                body.append(r"\end{tabular}")
                body.append(r"\end{table}")
            if HPC_DATASET_ROWS:
                body.append(r"\begin{table}[h]")
                body.append(r"\caption{Datasets in the uploaded HPC KODAMA comparison archive.}")
                body.append(r"\small")
                body.append(r"\begin{tabularx}{\linewidth}{lrrX}")
                body.append(r"\toprule")
                body.append(r"Dataset & $n$ & $p$ & Domain / representation \\")
                body.append(r"\midrule")
                for dataset, n_value, p_value, domain in HPC_DATASET_ROWS:
                    body.append(
                        f"{tex_escape(dataset)} & {tex_escape(n_value)} & {tex_escape(p_value)} & "
                        f"{tex_escape(domain)} \\\\"
                    )
                body.append(r"\bottomrule")
                body.append(r"\end{tabularx}")
                body.append(r"\end{table}")
                body.append("")
            if HPC_RUNTIME_ROWS:
                body.append(r"\begin{table}[h]")
                body.append(r"\caption{Preliminary archived CUDA KODAMA core wall times in seconds, medians over seeds 4, 17, and 42 at $M=Tcycle=100$. A dash denotes a classifier run absent from the archive; all rows will be rerun from the frozen release tag.}")
                body.append(r"\centering\small")
                body.append(r"\begin{tabular}{lrrl}")
                body.append(r"\toprule")
                body.append(r"Dataset & KNN core s & PLS--LDA core s & Coverage \\")
                body.append(r"\midrule")
                for dataset, knn_seconds, pls_seconds, coverage in HPC_RUNTIME_ROWS:
                    body.append(
                        f"{tex_escape(dataset)} & {tex_escape(knn_seconds)} & "
                        f"{tex_escape(pls_seconds)} & {tex_escape(coverage)} \\\\"
                    )
                body.append(r"\bottomrule")
                body.append(r"\end{tabular}")
                body.append(r"\end{table}")
                body.append("")
                body.append(tex_escape(
                    "The definitive rerun freezes source and wrapper tags, checksums, toolchain and "
                    "container identity, the full dataset/classifier/seed matrix, paired plotting "
                    "indices, timing boundaries, memory metrics, and explicit failure rows. The "
                    "executable protocol is maintained in docs/JMLR_CUDA_HPC_EXPERIMENT_PLAN.md."
                ))
                body.append("")
            body.append(tex_escape(
                "The following tables report every complete dataset-level comparison. The corresponding "
                "machine-readable CSV retains seed-level rows, plotting indices, and additional quality "
                "metrics. Values are three-seed medians; external labels are used only after optimization."
            ))
            body.append("")
            if VISUALIZATION_COMPARISON_ROWS:
                for embedding_name in ("UMAP", "openTSNE"):
                    embedding_rows = [
                        row for row in VISUALIZATION_COMPARISON_ROWS
                        if row[2] == embedding_name
                    ]
                    if not embedding_rows:
                        continue
                    body.append(r"\begin{table}[p]")
                    body.append(r"\centering\scriptsize")
                    body.append(
                        r"\caption{Dataset-level classic versus KODAMA "
                        + tex_escape(embedding_name)
                        + r" comparison on CUDA. Silhouette and trustworthiness entries are classic $\rightarrow$ KODAMA; $\Delta$ silhouette is KODAMA minus classic.}"
                    )
                    body.append(r"\resizebox{\linewidth}{!}{%")
                    body.append(r"\begin{tabular}{lllllll}")
                    body.append(r"\toprule")
                    body.append(r"Dataset & Classifier & Silhouette & $\Delta$ sil. & Trustworthiness & ARI/classes & KODAMA s \\")
                    body.append(r"\midrule")
                    for dataset, classifier, embedding, silhouette, delta, trust, ari_classes, seconds in embedding_rows:
                        body.append(
                            f"{tex_escape(dataset)} & {tex_escape(classifier)} & "
                            f"{tex_escape(silhouette)} & {tex_escape(delta)} & "
                            f"{tex_escape(trust)} & {tex_escape(ari_classes)} & "
                            f"{tex_escape(seconds)} \\\\"
                        )
                    body.append(r"\bottomrule")
                    body.append(r"\end{tabular}%")
                    body.append(r"}")
                    body.append(r"\end{table}")
                    body.append("")
            if IMAGENET_COMPARISON_FIGURE.exists():
                body.append(r"\begin{figure}[h]")
                body.append(r"\centering")
                body.append(r"\includegraphics[width=0.78\linewidth]{jmlr_imagenet_comparison_20260807/imagenet_classic_vs_kodama.png}")
                body.append(r"\caption{ImageNet seed-4 CUDA layouts for classic and KODAMA-KNN openTSNE/UMAP. Colors are 1,000 external classes withheld from KODAMA. Aggregate three-seed diagnostics are archived with the seed-level rows and are not measurements of this single displayed realization.}")
                body.append(r"\end{figure}")
                body.append("")
            if HPC_VISUALIZATION_FIGURE.exists():
                body.append(r"\begin{figure}[h]")
                body.append(r"\centering")
                body.append(r"\includegraphics[width=\linewidth]{jmlr_hpc_kodama_20260806/kodama_vs_classic_silhouette.png}")
                body.append(r"\caption{Paired three-seed CUDA comparison of classic fastEmbedR and KODAMA-corrected UMAP/openTSNE at M = 100 and Tcycle = 100. Points are dataset-level median silhouette changes; positive values favor KODAMA. External labels were used only after optimization.}")
                body.append(r"\end{figure}")
                body.append("")
            if PENDIGITS_ADVERSE_FIGURE.exists():
                body.append(r"\begin{figure}[h]")
                body.append(r"\centering")
                body.append(r"\includegraphics[width=\linewidth]{jmlr_cycle20_pendigits_m100_t100_20260808/pendigits_cpu_adverse.png}")
                body.append(r"\caption{PenDigits seed-4 CPU4 adverse control at $M=Tcycle=100$. Classic UMAP retains compact truth classes, whereas KNN- and PLS--LDA-corrected layouts are diffuse despite high internal CV accuracy. External labels were used only after optimization.}")
                body.append(r"\end{figure}")
                body.append("")
            if COIL20_KNN_FIGURE.exists():
                body.append(r"\begin{figure}[h]")
                body.append(r"\centering")
                body.append(r"\includegraphics[width=\linewidth]{jmlr_cycle21_coil20_m100_t100_20260808/coil20_knn_cpu_metal.png}")
                body.append(r"\caption{Raw COIL20 seed-4 KNN control at $M=Tcycle=100$. Metal reduced optimization wall time by 2.39x relative to CPU4, but both KODAMA truth silhouettes were below classic UMAP. External labels were used only after optimization.}")
                body.append(r"\end{figure}")
                body.append("")
            if SATIMAGE_FIGURE.exists():
                body.append(r"\begin{figure}[h]")
                body.append(r"\centering")
                body.append(r"\includegraphics[width=\linewidth]{jmlr_cycle24_satimage_metal_corrected_m100_t100_20260808/satimage_cpu_metal_corrected.png}")
                body.append(r"\caption{Corrected SatImage seed-4 CPU4/Metal control at $M=Tcycle=100$. A bounded resident fold cache and memory-aware run scheduler make Metal 1.21x faster for KNN and 1.93x faster than its pre-fix PLS--LDA path; low-dimensional Metal PLS--LDA remains 3.04x slower than CPU4. External labels were used only after optimization.}")
                body.append(r"\end{figure}")
                body.append("")
            body.append(r"\subsection{Sensitivity of M and Tcycle}")
            for paragraph in PARAMETER_SENSITIVITY_PARAGRAPHS:
                body.append(tex_escape(paragraph))
                body.append("")
            if RELEASE_SENSITIVITY_READY:
                body.append(r"\begin{table}[h]")
                body.append(r"\caption{Tcycle sensitivity at fixed M = 100. Accuracy entries show Tcycle = 20 to Tcycle = 100.}")
                body.append(r"\small")
                body.append(r"\resizebox{\linewidth}{!}{%")
                body.append(r"\begin{tabular}{llllll}")
                body.append(r"\toprule")
                body.append(r"Dataset & Classifier & Best acc & Median acc & Median classes & Seconds \\")
                body.append(r"\midrule")
                for dataset, classifier, best_acc, median_acc, median_classes, seconds in TCYCLE_SENSITIVITY_ROWS:
                    body.append(
                        f"{tex_escape(dataset)} & {tex_escape(classifier)} & {tex_escape(best_acc)} & "
                        f"{tex_escape(median_acc)} & {tex_escape(median_classes)} & {tex_escape(seconds)} \\\\"
                    )
                body.append(r"\bottomrule")
                body.append(r"\end{tabular}%")
                body.append(r"}")
                body.append(r"\end{table}")
                body.append("")
                body.append(r"\begin{table}[h]")
                body.append(r"\caption{M sensitivity at fixed Tcycle = 100. Entries show M = 20 / 50 / 100.}")
                body.append(r"\small")
                body.append(r"\resizebox{\linewidth}{!}{%")
                body.append(r"\begin{tabular}{llllll}")
                body.append(r"\toprule")
                body.append(r"Dataset & Classifier & Best acc & Median acc & Selected ARI & Seconds \\")
                body.append(r"\midrule")
                for dataset, classifier, best_acc, median_acc, best_ari, seconds in M_SENSITIVITY_ROWS:
                    body.append(
                        f"{tex_escape(dataset)} & {tex_escape(classifier)} & {tex_escape(best_acc)} & "
                        f"{tex_escape(median_acc)} & {tex_escape(best_ari)} & {tex_escape(seconds)} \\\\"
                    )
                body.append(r"\bottomrule")
                body.append(r"\end{tabular}%")
                body.append(r"}")
                body.append(r"\end{table}")
                body.append("")
                if SENSITIVITY_FIGURE.exists():
                    body.append(r"\begin{figure}[h]")
                    body.append(r"\centering")
                    body.append(r"\includegraphics[width=\linewidth]{kodama_m_tcycle_sensitivity.png}")
                    body.append(r"\caption{Sensitivity of KODAMA quality and runtime to Tcycle and M on named matrix-input datasets.}")
                    body.append(r"\end{figure}")
                    body.append("")
                body.append(r"\begin{table}[h]")
                body.append(r"\caption{Convergence of agreement-graph edge weights. RMSE entries show M = 10 / 20 / 50 relative to M = 100.}")
                body.append(r"\small")
                body.append(r"\begin{tabular}{llll}")
                body.append(r"\toprule")
                body.append(r"Dataset & Classifier & RMSE to M100 & Correlation at M50 \\")
                body.append(r"\midrule")
                for dataset, classifier, rmse, correlation in ENSEMBLE_CONVERGENCE_ROWS:
                    body.append(
                        f"{tex_escape(dataset)} & {tex_escape(classifier)} & {tex_escape(rmse)} & "
                        f"{tex_escape(correlation)} \\\\"
                    )
                body.append(r"\bottomrule")
                body.append(r"\end{tabular}")
                body.append(r"\end{table}")
                body.append("")
                if ENSEMBLE_CONVERGENCE_FIGURE.exists():
                    body.append(r"\begin{figure}[h]")
                    body.append(r"\centering")
                    body.append(r"\includegraphics[width=\linewidth]{jmlr_pilot_20260716/ensemble_convergence.png}")
                    body.append(r"\caption{Convergence of edge-agreement weights as independent KODAMA runs are added. RMSE is measured against M = 100; the right panel shows the Bernoulli worst-case Monte Carlo standard error.}")
                    body.append(r"\end{figure}")
                    body.append("")
            else:
                body.append(tex_escape(
                    "Named matrix-input release results are pending. Preliminary anonymized development rows are deliberately excluded from this submission draft."
                ))
                body.append("")
            body.append(r"\subsection{Implementation claims and archived evidence}")
            body.append(
                r"Source locations, tests, rejected alternatives, and backend measurements are "
                r"archived in the repository's dated full-code audit, benchmark index, and "
                r"checksum-bearing CSV files. The supplement reports the decisive aggregates rather "
                r"than duplicating that machine-auditable matrix."
            )
            body.append("")
            body.append(r"\begin{table}[h]")
            body.append(r"\caption{Release-validation evidence tracked by the project.}")
            body.append(r"\small")
            body.append(r"\begin{tabularx}{\linewidth}{p{0.18\linewidth}X}")
            body.append(r"\toprule")
            body.append(r"Area & Current evidence \\")
            body.append(r"\midrule")
            for area, evidence in VALIDATION_ROWS:
                body.append(f"{tex_escape(area)} & {tex_escape(evidence)} \\\\")
            body.append(r"\bottomrule")
            body.append(r"\end{tabularx}")
            body.append(r"\end{table}")
            body.append("")
    body.append(r"\section{KODAMA pseudocode}")
    body.append(tex_escape(
        "Algorithm 1 isolates one independent run inside the M-run ensemble. It starts from one "
        "landmark label vector, uses the previous CV predictions to propose label changes, "
        "evaluates each proposal with exactly one CV pass, and returns the best label vector "
        "found by that run."
    ))
    body.append("")
    body.append(r"\begin{scriptsize}")
    body.append(r"\begin{verbatim}")
    body.extend(SINGLE_RUN_PSEUDOCODE_LINES)
    body.append(r"\end{verbatim}")
    body.append(r"\end{scriptsize}")
    body.append("")
    body.append(tex_escape(
        "The compact ensemble pseudocode in Section 2.5 and the graph-input contract in "
        "Section 3.8 define KODAMA.matrix and KODAMA.matrix.graph without repeating the same "
        "steps here. Executable C++, R, and Python examples are maintained in the public API "
        "documentation."
    ))
    body.append("")
    body.append(r"\acks{The authors thank contributors to the KODAMA, fastPLS, and fastEmbedR software projects.}")
    body.append(r"\nocite{*}")
    body.append(r"\bibliography{kodama_cpp_refs}")
    body.append(r"\end{document}")
    TEX.write_text("\n".join(body) + "\n")


def build_bib() -> None:
    BIB.write_text(
        dedent(
            r"""
            @article{cacciatore2014kodama,
              title={Knowledge discovery by accuracy maximization},
              author={Cacciatore, Stefano and Luchinat, Claudio and Tenori, Leonardo},
              journal={Proceedings of the National Academy of Sciences},
              volume={111},
              number={14},
              pages={5117--5122},
              year={2014}
            }

            @manual{kodamaR,
              title={KODAMA: Knowledge Discovery by Accuracy Maximization},
              author={Cacciatore, Stefano and Tenori, Leonardo},
              note={R package documentation, CRAN},
              year={2026}
            }

            @article{cacciatore2017kodama,
              title={KODAMA: an R package for knowledge discovery and data mining},
              author={Cacciatore, Stefano and Tenori, Leonardo and Luchinat, Claudio and Bennett, Phillip R. and MacIntyre, David A.},
              journal={Bioinformatics},
              volume={33},
              number={4},
              pages={621--623},
              year={2017},
              doi={10.1093/bioinformatics/btw705}
            }

            @book{chapelle2006semisupervised,
              title={Semi-Supervised Learning},
              author={Chapelle, Olivier and Sch{\"o}lkopf, Bernhard and Zien, Alexander},
              publisher={MIT Press},
              year={2006}
            }

            @inproceedings{zhu2003gaussianfields,
              title={Semi-supervised learning using Gaussian fields and harmonic functions},
              author={Zhu, Xiaojin and Ghahramani, Zoubin and Lafferty, John},
              booktitle={Proceedings of the Twentieth International Conference on Machine Learning},
              pages={912--919},
              year={2003}
            }

            @inproceedings{zhou2003localglobal,
              title={Learning with local and global consistency},
              author={Zhou, Dengyong and Bousquet, Olivier and Lal, Thomas Navin and Weston, Jason and Sch{\"o}lkopf, Bernhard},
              booktitle={Advances in Neural Information Processing Systems},
              year={2003}
            }

            @inproceedings{lee2013pseudolabel,
              title={Pseudo-label: The simple and efficient semi-supervised learning method for deep neural networks},
              author={Lee, Dong-Hyun},
              booktitle={ICML Workshop on Challenges in Representation Learning},
              year={2013}
            }

            @inproceedings{ratner2016dataprogramming,
              title={Data programming: Creating large training sets, quickly},
              author={Ratner, Alexander and De Sa, Christopher and Wu, Sen and Selsam, Daniel and R{\'e}, Christopher},
              booktitle={Advances in Neural Information Processing Systems},
              year={2016}
            }

            @inproceedings{benhur2002stability,
              title={A stability based method for discovering structure in clustered data},
              author={Ben-Hur, Asa and Elisseeff, Andre and Guyon, Isabelle},
              booktitle={Pacific Symposium on Biocomputing},
              pages={6--17},
              year={2002}
            }

            @article{tibshirani2005predictionstrength,
              title={Cluster validation by prediction strength},
              author={Tibshirani, Robert and Walther, Guenther},
              journal={Journal of Computational and Graphical Statistics},
              volume={14},
              number={3},
              pages={511--528},
              year={2005}
            }

            @article{hubert1985comparing,
              title={Comparing partitions},
              author={Hubert, Lawrence and Arabie, Phipps},
              journal={Journal of Classification},
              volume={2},
              pages={193--218},
              year={1985}
            }

            @article{rousseeuw1987silhouettes,
              title={Silhouettes: A graphical aid to the interpretation and validation of cluster analysis},
              author={Rousseeuw, Peter J.},
              journal={Journal of Computational and Applied Mathematics},
              volume={20},
              pages={53--65},
              year={1987}
            }

            @inproceedings{kohavi1995crossvalidation,
              title={A study of cross-validation and bootstrap for accuracy estimation and model selection},
              author={Kohavi, Ron},
              booktitle={Proceedings of the Fourteenth International Joint Conference on Artificial Intelligence},
              pages={1137--1143},
              year={1995}
            }

            @article{ambroise2002selectionbias,
              title={Selection bias in gene extraction on the basis of microarray gene-expression data},
              author={Ambroise, Christophe and McLachlan, Geoffrey J.},
              journal={Proceedings of the National Academy of Sciences},
              volume={99},
              number={10},
              pages={6562--6566},
              year={2002}
            }

            @article{varma2006bias,
              title={Bias in error estimation when using cross-validation for model selection},
              author={Varma, Sudhir and Simon, Richard},
              journal={BMC Bioinformatics},
              volume={7},
              pages={91},
              year={2006}
            }

            @article{kriegeskorte2009circular,
              title={Circular analysis in systems neuroscience: the dangers of double dipping},
              author={Kriegeskorte, Nikolaus and Simmons, W. Kyle and Bellgowan, Patrick S. F. and Baker, Chris I.},
              journal={Nature Neuroscience},
              volume={12},
              pages={535--540},
              year={2009}
            }

            @article{dejong1993simpls,
              title={SIMPLS: an alternative approach to partial least squares regression},
              author={de Jong, Sijmen},
              journal={Chemometrics and Intelligent Laboratory Systems},
              volume={18},
              number={3},
              pages={251--263},
              year={1993}
            }

            @article{johnson2019faiss,
              title={Billion-scale similarity search with GPUs},
              author={Johnson, Jeff and Douze, Matthijs and Jegou, Herve},
              journal={IEEE Transactions on Big Data},
              year={2019}
            }

            @article{malkov2020hnsw,
              title={Efficient and robust approximate nearest neighbor search using Hierarchical Navigable Small World graphs},
              author={Malkov, Yu. A. and Yashunin, D. A.},
              journal={IEEE Transactions on Pattern Analysis and Machine Intelligence},
              year={2020}
            }

            @article{mcinnes2018umap,
              title={UMAP: Uniform Manifold Approximation and Projection for dimension reduction},
              author={McInnes, Leland and Healy, John and Melville, James},
              journal={arXiv:1802.03426},
              year={2018}
            }

            @article{vandermaaten2008tsne,
              title={Visualizing data using t-SNE},
              author={van der Maaten, Laurens and Hinton, Geoffrey},
              journal={Journal of Machine Learning Research},
              volume={9},
              pages={2579--2605},
              year={2008}
            }

            @article{linderman2019fitsne,
              title={Fast interpolation-based t-SNE for improved visualization of single-cell RNA-seq data},
              author={Linderman, George C. and Rachh, Manas and Hoskins, Jeremy G. and Steinerberger, Stefan and Kluger, Yuval},
              journal={Nature Methods},
              volume={16},
              pages={243--245},
              year={2019}
            }

            @article{sonnenburg2007opensource,
              title={The need for open source software in machine learning},
              author={Sonnenburg, S{\"o}ren and Braun, Mikio L. and Ong, Cheng Soon and Bengio, Samy and Bottou, Leon and Holmes, Geoffrey and LeCun, Yann and M{\"u}ller, Klaus-Robert and Pereira, Fernando and Rasmussen, Carl Edward and R{\"a}tsch, Gunnar and Sch{\"o}lkopf, Bernhard and Smola, Alexander and Vincent, Pascal and Weston, Jason and Williamson, Robert C.},
              journal={Journal of Machine Learning Research},
              volume={8},
              pages={2443--2466},
              year={2007}
            }

            @article{pineau2021reproducibility,
              title={Improving reproducibility in machine learning research},
              author={Pineau, Joelle and Vincent-Lamarre, Philippe and Sinha, Koustuv and Lariviere, Vincent and Beygelzimer, Alina and d'Alche-Buc, Florence and Fox, Emily and Larochelle, Hugo},
              journal={Journal of Machine Learning Research},
              volume={22},
              number={164},
              pages={1--20},
              year={2021}
            }
            """
        ).strip()
        + "\n"
    )


def main() -> None:
    build_docx()
    build_self_review()
    build_bib()
    print(MANUSCRIPT)
    print(SELF_REVIEW)
    print(BIB)


if __name__ == "__main__":
    main()
