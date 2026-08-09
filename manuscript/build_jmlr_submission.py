#!/usr/bin/env python3
# SPDX-FileCopyrightText: 2026 Stefano Cacciatore
# SPDX-License-Identifier: MIT

from __future__ import annotations

from pathlib import Path
from textwrap import dedent

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

import build_kodama_manuscript as supplement


ROOT = Path(__file__).resolve().parent
MAIN_TEX = ROOT / "kodama_cpp_jmlr_mloss.tex"
COVER_DOCX = ROOT / "kodama_cpp_jmlr_cover_letter.docx"
COVER_TXT = ROOT / "kodama_cpp_jmlr_cover_letter.txt"
READINESS_DOCX = ROOT / "kodama_cpp_jmlr_submission_readiness.docx"

SOFTWARE_VERSION = "0.1.0"
REVIEWED_BASELINE = "a32f71844fb0860f8faec0169c2850f800a68da4"
R_WRAPPER_BASELINE = "4d5506bde8918162dae0ad1d19c4f6f7a4df64ba"


def build_main_tex() -> None:
    text = dedent(
        r"""
        \documentclass[twoside,11pt]{article}
        \usepackage[preprint]{jmlr2e}
        \usepackage{amsmath}
        \usepackage{booktabs}
        \usepackage{graphicx}
        \usepackage{lastpage}
        \usepackage{tabularx}
        \usepackage{url}
        \hypersetup{hidelinks}
        \setlength{\textfloatsep}{7pt plus 2pt minus 2pt}
        \setlength{\floatsep}{6pt plus 2pt minus 2pt}
        \setlength{\intextsep}{6pt plus 2pt minus 2pt}
        \setlength{\abovecaptionskip}{3pt}
        \setlength{\belowcaptionskip}{0pt}
        \renewcommand{\arraystretch}{0.96}

        \jmlrheading{}{2026}{}{}{}{}{Kassim et al.}
        \ShortHeadings{kodama-cpp}{Kassim et al.}
        \firstpageno{1}

        \begin{document}

        \title{kodama-cpp: Cross-Validated Accuracy Maximization on CPU, CUDA, and Apple Metal}

        \author{%
        \name Moussa Kassim$^{1,2,\ast}$ \email moussa.kassim@icgeb.org\\
        \name Martin Ocharo$^{1,2,\ast}$ \email martin.ocharo@icgeb.org\\
        \name Dalia Ahmed$^{1}$ \email dalia.ahmed@icgeb.org\\
        \name Dupe Ojo$^{1}$ \email dupe.ojo@icgeb.org\\
        \name Alessia Vignoli$^{3,4}$ \email vignoli@cerm.unifi.it\\
        \name Leonardo Tenori$^{3,4,\dagger}$ \email tenori@cerm.unifi.it\\
        \name Stefano Cacciatore$^{1,2,\dagger}$ \email stefano.cacciatore@icgeb.org\\
        \addr $^1$Bioinformatics Unit, International Centre for Genetic Engineering and Biotechnology (ICGEB), Cape Town 7925, South Africa\\
        \addr $^2$Department of Integrative Biomedical Sciences, Institute of Infectious Disease \& Molecular Medicine (IDM), University of Cape Town, Cape Town 7925, South Africa\\
        \addr $^3$Department of Chemistry ``Ugo Schiff'', University of Florence, Sesto Fiorentino, Italy\\
        \addr $^4$Magnetic Resonance Center (CERM), University of Florence, Sesto Fiorentino, Italy\\
        \addr $^{\ast}$Moussa Kassim and Martin Ocharo contributed equally.\\
        \addr $^{\dagger}$Leonardo Tenori and Stefano Cacciatore are co-corresponding authors.}

        \maketitle

        \begin{abstract}
        KODAMA discovers latent structure by maximizing held-out label predictability. This standalone C++17 implementation is shared by thin R and Python wrappers. Native multicore CPU, NVIDIA CUDA, and Apple Metal backends provide float32 KNN without silent fallback and introduce a label-aware SIMPLS route followed by LDA in latent component space. The library also introduces prediction-guided label evolution and exact-quota landmark projection. In an explicitly exploratory matched CUDA analysis, PLS--LDA KODAMA improved UMAP truth-label silhouette on 6 of 10 datasets (median $+0.025$, bootstrap 95\% interval $[-0.067,+0.084]$); KNN and openTSNE were less favorable. We therefore separate systems validation from external-label diagnostics and make no universal visualization claim.
        \end{abstract}

        \begin{keywords}
        KODAMA, cross-validation, unsupervised learning, heterogeneous computing, manifold learning
        \end{keywords}

        \section{Introduction}

        KODAMA treats labels as optimization variables: useful labels can be reproduced on held-out samples. An ensemble of optimized vectors then corrects neighborhood dissimilarities \citep{cacciatore2014kodama,cacciatore2017kodama}. Unlike stability analysis or prediction strength, predictability is inside the search and no observed labels anchor it \citep{benhur2002stability,tibshirani2005predictionstrength}.

        Repeated cross-validation makes reusable state a systems concern. Our five contributions are \textbf{(1)} a typed C++ core shared by R and Python; \textbf{(2)} CPU, CUDA, and Metal under one float32 contract; \textbf{(3)} a new label-aware SIMPLS plus latent-space LDA evaluator alongside KNN; \textbf{(4)} prediction-guided evolution with adaptive proposals, degeneracy guards, and error-scaled cooling; and \textbf{(5)} exact-quota landmark projection for scalable, reproducible subsampling.

        \section{KODAMA objective and algorithm}

        For data $X\in\mathbb{R}^{n\times p}$, labels $y$, folds $\Pi$, and classifier family $\mathcal F$, let
        \begin{equation}
          A(y;X,\mathcal F,\Pi)=\frac{1}{n}\sum_{i=1}^{n}\mathbf{1}\!\left[y_i=\widehat y_i^{(-\Pi(i))}\right].
        \end{equation}
        KODAMA searches for high $A$ using either KNN or SIMPLS followed by latent-space LDA \citep{dejong1993simpls}. Folds stay fixed within each run, and supplied constraint groups move atomically. Each independent run draws exactly $L$ landmarks by population-proportional quotas from a coarse matrix partition; randomized systematic rounding fills the residual quota without replacement. A separate partition initializes the labels, so sampling strata are not optimization targets.

        The classifiers share an evolution \emph{scaffold}, not an identical state policy. Both use fixed folds, prediction-guided grouped proposals, the common transition matrix, one new CV pass per proposal, error-scaled cooling, and independent $M$ runs. Let $p_k$ be the class proportions, $D=1-\sum_kp_k^2$, $H=-\sum_kp_k\log p_k$, $K_{\rm eff}=e^H$, and $R=\max\{0,\log[K/\max(1,K_{\rm eff})]\}$. State selection uses
        \begin{equation}
          S_{\mathcal F}(y)=A(y)\sqrt D-\mathbf{1}_{\{\mathcal F={\rm PLS\mbox{-}LDA}\}}(1-A(y))
          \max\!\left\{\frac{H}{\log n},\frac{R}{1+R}\right\}.
        \end{equation}
        Standard \texttt{KODAMA.matrix} applies transition-driven coarsening and the subtraction only to PLS--LDA; KNN uses the common diversity term alone. This reflects that LDA fits a mean per active class, whereas KNN has no analogous global class fit. Existing sensitivity results motivate, but do not isolate, this policy; no universal improvement is claimed.

        The selected landmark vector is projected to all samples with the same classifier. If $a_{ij}$ is the fraction of projected runs agreeing on graph edge $(i,j)$, KODAMA returns $d'_{ij}=(1+d_{ij})/a_{ij}^{2}$ and removes zero-agreement edges. Exact proposal, temperature, quota, and coarsening formulas are given in the supplement.

        \begin{figure}[h!]
        \hrule\vspace{2pt}
        \begin{minipage}{0.98\linewidth}
        \scriptsize
        \textbf{Algorithm 1: KODAMA ensemble}\\
        \textbf{Input:} $X,\mathcal F,M,T,L,K_0,\Pi,s$.\quad \textbf{Output:} $G,C,A_{\rm best},Z_U,Z_T$.
        \begin{tabbing}
        \quad\=\quad\=\quad\=\kill
        $X_{32}\leftarrow\textsc{Float32}(X)$; $(G_0,Z_U,Z_T)\leftarrow\textsc{KODAMA.Graph}(X_{32})$ once\\
        \textbf{for} $r=1,\ldots,M$ \textbf{do}\\
        \> $I_r\leftarrow\textsc{ExactQuotaSample}(\textsc{Partition}(X,K_0),L,s+r)$\\
        \> $y\leftarrow\textsc{InitialPartition}(X_{I_r},K_0)$; $(A,\widehat y)\leftarrow\textsc{CV}(\mathcal F,X_{I_r},y,\Pi_r)$\\
        \> \textbf{for} $t=1,\ldots,T$ \textbf{do}\\
        \>\> $y'\leftarrow\textsc{GuidedProposal}(y,\widehat y,t,T)$; if $\mathcal F={\rm PLS\mbox{-}LDA}$, apply transition coarsening\\
        \>\> $(A',\widehat y')\leftarrow\textsc{CV}(\mathcal F,X_{I_r},y',\Pi_r)$; update by $S_{\mathcal F}$ and cooling\\
        \> \textbf{end for}\\
        \> $C_{\cdot r}\leftarrow\textsc{Project}(\mathcal F,X_{I_r},y_{\rm best},X)$\\
        \textbf{end for}\\
        $G\leftarrow\textsc{AgreementCorrectionInPlace}(G_0,C)$; \textbf{return} $G,C,A_{\rm best},Z_U,Z_T$
        \end{tabbing}
        \end{minipage}
        \vspace{2pt}\hrule
        \caption{Compact pseudocode; PLS--LDA-specific steps are marked.}
        \label{alg:kodama}
        \end{figure}

        \section{Standalone heterogeneous implementation}

        The typed C++ API owns folds, proposals, classifiers, landmark projection, graph correction, and metadata; R and Python only convert host objects. CPU provides parallel HNSW over contiguous float32 storage \citep{malkov2020hnsw}. CUDA and Metal provide exact/IVF search, k-means, label-aware SIMPLS--LDA, PCA, and persistent workspaces. Large accelerator graphs use one resident IVF-Flat index whose list count is independent of the probe bound; a fixed exact pilot raises \texttt{nprobe} until recall reaches 0.99. Unsupported backends raise instead of silently falling back.

        \texttt{KODAMA.graph} prepares one graph and backend-matched PCA starts without retaining the raw matrix. \texttt{KODAMA.matrix} accepts raw features, this prepared object, or a bare paired $n\times k$ graph. PLS--LDA avoids dense one-hot responses, compacts active labels, and reuses backend-specific fold matrices and float32 workspaces. Metal retains a bounded fold working set and schedules independent $M$ runs against the device memory budget. The API also exposes both CV kernels, both core optimizers, PCA, and float32 UMAP/openTSNE on CPU, CUDA, and Metal \citep{mcinnes2018umap,vandermaaten2008tsne}. Full architecture, graph-only semantics, numerical safeguards, implementation changes, and wrapper contracts are documented in the supplement.

        \begin{figure}[h!]
        \centering
        \includegraphics[width=0.56\linewidth]{jmlr_imagenet_comparison_20260807/imagenet_classic_vs_kodama.png}
        \caption{Illustrative ImageNet layouts ($n=1{,}281{,}167$, 1,000 reference classes; colors repeat), CUDA seed 4. Classic embeddings use the raw feature graph; KODAMA panels use the KNN-corrected graph at $M=T=100$. Three-seed medians show increased sampled trustworthiness (openTSNE 0.680 to 0.694; UMAP 0.612 to 0.686) but lower truth-label silhouette (0.615 to 0.103; 0.629 to 0.178). Quantitative comparisons use paired metric samples; a frozen rerun will render the same plotting indices in all panels.}
        \label{fig:imagenet}
        \end{figure}

        \section{Validation, availability, and scope}

        \begin{table}[h!]
        \caption{Systems validation and exploratory matched classic-versus-KODAMA visualization results. Silhouette changes are KODAMA minus classic fastEmbedR medians over seeds 4, 17, and 42; $M=T=100$.}
        \label{tab:validation}
        \centering\scriptsize
        {\setlength{\tabcolsep}{3pt}\begin{tabular}{llllr}
        \toprule
        Scope & Comparison & Backend & Matched data & Result \\
        \midrule
        Kernel & KNNCV, MNIST CPU/device & CUDA & one dataset & 16.2x; .974/.973 acc. \\
        Kernel & PLSLDACV, MetRef CPU/device & CUDA & one dataset & 4.20x; .992/.993 acc. \\
        Predecessor & MetRef PLS--DA (5 comp.), $M=T=100$ & CPU1 & contextual & 965.8 s; CV 1.000/.975 \\
        Current & MetRef PLS--LDA (50 comp.), $M=T=100$ & CPU4/CUDA & contextual & 316.3/270.4 s \\
        Optimizer & shared graph CPU4/device & Metal & 15 paired cells & 1.83x [1.53,2.12] \\
        Pipeline & flow18 graph, exact/resident-IVF & CUDA & one dataset & 296.5/25.9 s; 11.5x \\
        Visualization & KNN KODAMA vs UMAP & CUDA & 10 datasets & 2/10 improved; $-0.094$ \\
        Visualization & PLS--LDA KODAMA vs UMAP & CUDA & 10 datasets & 6/10 improved; $+0.025$ \\
        Visualization & KNN KODAMA vs openTSNE & CUDA & 10 datasets & 0/10 improved; $-0.084$ \\
        Visualization & PLS--LDA KODAMA vs openTSNE & CUDA & 10 datasets & 3/10 improved; $-0.056$ \\
        \bottomrule
        \end{tabular}}
        \end{table}

        Table~\ref{tab:validation} separates acceleration from downstream structure. The historical MetRef PLS--DA route required 965.8 s at its five-component default; independently retained current 50-component PLS--LDA runs required 316.3 s on CPU4 and 270.4 s on CUDA. This is a contextual product comparison, not an implementation-only speedup: evaluator, components, search, parallelism, and retained seeds differ. On flow18 ($n=1{,}000{,}021$), resident IVF reduced full graph time from 296.5 to 25.9 s. The matched three-seed archive contains 10 datasets per classifier. PLS--LDA UMAP was the only positive cross-dataset median ($+0.025$, 95\% bootstrap interval $[-0.067,+0.084]$); KNN openTSNE was worse on all 10 (Holm $p=0.0078$). Figure~\ref{fig:imagenet} exposes the corresponding ImageNet trade-off instead of selecting only a favorable example. Labels were withheld during optimization. KODAMA R 2.4.1/2.4 is the sole predecessor; the supplement also reports matched MetRef KNN. All per-dataset timings, sensitivity analyses, and implementation evidence are in the supplement.

        Tests cover float32 numerics, backend identity, graph/PCA contracts, visualization initialization, wrappers, and direct public accelerator entry points; physical Metal and current-source CUDA execution are recorded separately from compile-only checks. The MIT repositories are \href{https://github.com/tkcaccia/kodama-cpp}{kodama-cpp} and \href{https://github.com/tkcaccia/kodama-r}{kodama-r}. Public snapshots checked on 6 August 2026 were \texttt{0b8b873} and \texttt{7090592}. A tagged confirmatory rerun remains a submission prerequisite; fragmentation, coverage, timings, and limitations are detailed in the supplement.

        \newpage
        \bibliography{kodama_cpp_refs}

        \end{document}
        """
    ).strip() + "\n"
    MAIN_TEX.write_text(text, encoding="utf-8")


def add_label_value(doc: Document, label: str, value: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    paragraph.add_run(value)


def build_cover_letter() -> None:
    doc = Document()
    supplement.apply_styles(doc)
    supplement.apply_compact_memo_styles(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    normal = doc.styles["Normal"]
    normal.font.size = Pt(9.4)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.space_after = Pt(2)
    heading = doc.styles["Heading 1"]
    heading.font.size = Pt(12.5)
    heading.paragraph_format.space_before = Pt(6)
    heading.paragraph_format.space_after = Pt(2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Cover Letter: JMLR Machine Learning Open Source Software")
    run.bold = True
    run.font.size = Pt(13.5)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("kodama-cpp: Cross-Validated Accuracy Maximization on CPU, CUDA, and Apple Metal").italic = True

    doc.add_paragraph("Dear Editors of the Journal of Machine Learning Research,")
    doc.add_paragraph(
        "We submit kodama-cpp for consideration in the Machine Learning Open Source Software track. "
        "The software is a standalone C++17 implementation of KODAMA with explicit CPU, NVIDIA CUDA, "
        "and Apple Metal backends and thin R and Python wrappers."
    )

    doc.add_heading("Submission identity", level=1)
    add_label_value(doc, "Open-source license", "MIT project code with compatible retained third-party terms; see PROVENANCE.md and THIRD_PARTY_NOTICES.md")
    add_label_value(doc, "Project URL", "https://github.com/tkcaccia/kodama-cpp")
    add_label_value(doc, "R wrapper URL", "https://github.com/tkcaccia/kodama-r")
    add_label_value(
        doc,
        "Python binding source",
        "https://github.com/tkcaccia/kodama-cpp/tree/main/split-repos/kodama-python "
        "(versioned and tested with the core source)",
    )
    add_label_value(doc, "Software version", f"{SOFTWARE_VERSION} release candidate")
    add_label_value(doc, "Reviewed baseline", REVIEWED_BASELINE)
    add_label_value(doc, "Reviewed R-wrapper baseline", R_WRAPPER_BASELINE)

    doc.add_heading("Significance and software contribution", level=1)
    doc.add_paragraph(
        "The submission preserves KODAMA's cross-validated label-predictability signal while making "
        "five contributions: a standalone C++17 core for thin R/Python wrappers; native and observable "
        "CPU/CUDA/Metal execution; label-aware SIMPLS followed by latent-space LDA as a new KODAMA "
        "evaluator; guided label evolution based on out-of-fold predictions and class transitions; "
        "and an exact-quota landmark-projection strategy for scalable, reproducible subsampling. "
        "The repository contains installation instructions, "
        "API examples, explicit pseudocode, tests, benchmark drivers, a public API snapshot, and a "
        "machine-checkable licensing/provenance audit."
    )

    doc.add_heading("Prior publications and delta", level=1)
    doc.add_paragraph(
        "The KODAMA method was published by Cacciatore, Luchinat, and Tenori in PNAS (2014), and the "
        "historical R package was described by Cacciatore et al. in Bioinformatics (2017). Prior "
        "publication of the method is disclosed. This submission concerns the new standalone software: "
        "a C++17 core independent of R, native CUDA and Metal backends, "
        "guided search mechanics, exact-quota landmark selection, float32 label-aware SIMPLS/LDA, "
        "package-owned neighbor search, strict backend metadata, and thin R/Python bindings. "
        "KODAMA R 2.4.1/2.4 is the sole previous-version comparator; later KODAMA releases are excluded from "
        "performance, quality, and methodological comparisons. The matched KNN comparison is complemented by "
        "a contextual MetRef PLS-DA benchmark. It is reported without a matched-speedup claim because the new "
        "implementation introduces label-aware SIMPLS followed by latent-space LDA. "
        "The detailed delta is recorded in docs/previous-version-delta.md. The prior peer-reviewed papers "
        "will be supplied as supplementary material."
    )

    doc.add_heading("Evidence of use and openness", level=1)
    doc.add_paragraph(
        "KODAMA has an established public predecessor: a peer-reviewed method paper, a peer-reviewed R-package "
        "paper, a maintained public R package, documentation, and public source history. The CRANlogs public API "
        "reports 48,480 KODAMA downloads from 2017-01-01 through 2026-07-26, including 4,526 during "
        "2025-07-27--2026-07-26 (https://cranlogs.r-pkg.org/downloads/total/2017-01-01:2026-07-26/KODAMA and "
        "https://cranlogs.r-pkg.org/downloads/total/2025-07-27:2026-07-26/KODAMA). These counts concern the "
        "predecessor R package, not the new implementation. kodama-cpp itself is a new repository and, on "
        "2026-07-27, had zero GitHub stars, forks, releases, and public issues; these figures are disclosed rather "
        "than presented as adoption. "
        "The repository provides an issue tracker, contribution guide, changelog, API-stability policy, reproducible "
        "tests, and release checklist."
    )

    doc.add_heading("Keywords", level=1)
    doc.add_paragraph("KODAMA; cross-validation; unsupervised learning; heterogeneous computing; manifold learning")

    doc.add_paragraph("Sincerely,")
    doc.add_paragraph(
        "Leonardo Tenori and Stefano Cacciatore\n"
        "Co-corresponding authors\n"
        "tenori@cerm.unifi.it; stefano.cacciatore@icgeb.org"
    )
    doc.save(COVER_DOCX)

    cover_text = "\n\n".join(
        paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()
    )
    COVER_TXT.write_text(cover_text + "\n", encoding="utf-8")


def build_readiness_report() -> None:
    doc = Document()
    supplement.apply_styles(doc)
    supplement.apply_compact_memo_styles(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    doc.styles["Normal"].font.size = Pt(8.8)
    supplement.add_title(
        doc,
        "JMLR MLOSS reviewer comments and response status",
        "kodama-cpp release candidate 0.1.0",
        "Internal submission-readiness report",
    )
    doc.add_paragraph(
        "This report separates completed manuscript and repository revisions from evidence absent from the reviewed archive."
    )
    rows = [
        ("KODAMA 2.4.1/2.4 comparison", "MetRef includes a classifier-matched KNN benchmark and a separately labeled historical PLS-DA benchmark. The latter contextualizes the predecessor but is not a matched speedup baseline for the new SIMPLS plus latent-space LDA route. Later KODAMA releases are excluded.", "Historical PLS-DA rerun completed on chiamaka"),
        ("Frozen public API", "Version 0.1.0 macros, SemVer policy, CMake package version, wrapper versions, and a compile-link API snapshot test are present.", "Completed locally"),
        ("Version/tag/archive", "The manuscript records public core commit 0b8b873 and R-wrapper commit 7090592. A release checklist and clean-tag archive script prevent an uncommitted or dirty artifact from being cited; a DOI is optional.", "External tag and deposit pending"),
        ("License provenance", "Pinned upstream snapshots, per-component licensing, SPDX tests, retained licenses, and the Metal Apache exception are documented.", "Engineering audit complete; coauthor contribution confirmation remains"),
        ("MLOSS page budget", "The detailed manuscript is retained as a technical supplement; a separate official-style main paper targets four description pages plus references.", "Completed; four description pages plus references verified"),
        ("CUDA release build", "Commit a32f718 passed all four configured license, numerical/core, public-API, and float32 smoke tests in the CUDA 13.0 build on chiamaka.", "Completed on chiamaka"),
        ("Continuous integration", "GitHub Actions builds and tests the CPU core on Linux/macOS, native Metal on macOS, and the R/Python wrappers on Ubuntu. Separate coverage and API-documentation workflows passed for a32f718.", "Completed at a32f718"),
        ("User documentation", "A compact C++/R/Python walkthrough and generated Doxygen API are published through GitHub Pages.", "Completed and URL verified"),
        ("Active user community", "The cover letter reports dated CRANlogs evidence for the predecessor R package and explicitly discloses that the new library has no public adoption metrics yet.", "Independent kodama-cpp usage evidence pending"),
        ("Local wrapper release checks", "R CMD check --as-cran passed with no errors or warnings and one environment-only macOS detritus NOTE; an isolated wheel built against the installed core passed eight tests including physical-Metal visualization diagnostics.", "Completed locally on current worktree"),
        ("Python packaging", "The tested Python wrapper is versioned with the public core repository and exercised by CI.", "Completed for the bundled-source release architecture"),
        ("Coverage depth", "LLVM reports 80.48% line and 69.33% branch CPU coverage. Dedicated graph/clustering, binary UMAP, matrix orchestration, and resident-IVF ownership/error contracts exercise portable behavior; physical Metal separately covers the native resident index and accelerator failures.", "Improved locally; accelerator implementation branches and CUDA hardware CI remain pending"),
        ("Large-sample SIMPLS range", "An overflow-only power-norm guard restores all 50 requested components on MNIST70k. Three-replicate CPU/Metal validation, a synthetic regression, AddressSanitizer, and physical Metal CTest pass.", "Local evidence complete; frozen CUDA confirmation pending"),
        ("CPU PLS-LDA memory traffic", "Training class sums and the latent Gram matrix are streamed without complete score matrices. Four-dataset predictions match; TabulaMuris is 1.47x faster with a 0.587 peak-memory ratio.", "Accepted locally"),
        ("M/Tcycle rationale", "Named MetRef/USPS sweeps report CV accuracy, ARI, active classes, runtime, and agreement-graph convergence; M=100 is justified by ensemble precision rather than external-label tuning.", "Completed on CUDA"),
        ("Full-cycle local backend controls", "Shared-graph MetRef, USPS, ImageSegmentation, PageBlocks, PenDigits, and SatImage runs use M=Tcycle=100, CPU4 and physical Metal. Raw COIL20 adds paired KNN timing and a censored PLS-LDA stress cell. After the Metal lifecycle correction, SatImage KNN is 1.21x faster than CPU4 and PLS-LDA is 1.93x faster than its pre-fix Metal path, while remaining 3.04x slower than CPU4 in this low-dimensional regime. MetRef independently confirms the correction.", "Completed locally; clean tagged CPU4/CUDA replication pending"),
        ("Metal lifecycle and graph correction", "A bounded resident fold cache and device-budgeted independent-M scheduler (up to 32 lanes) remove repeated uploads without changing KODAMA or classifier mathematics. MetRef M=T=100 PLS-LDA improved 631.415 to 312.029 s with exact core diagnostics.", "Accepted locally; physical Metal suite passed"),
        ("Final ablations", "The release driver fixes M=100, Tcycle=100, landmarks, k, ncomp, splitting, graph correction, backend, wrappers, and external metrics before inspection.", "Eleven-dataset CUDA archive analyzed with three seed-matched replicates; replicated KODAMA 2.4.1 predecessor runs and the frozen confirmatory matrix remain required"),
        ("Visualization claims", "Classic and KODAMA UMAP/openTSNE are paired by dataset, seed, and CUDA profile. Dataset-level seed medians, 100,000-resample bootstrap intervals, exact Wilcoxon and sign tests, and Holm adjustment are reported. The archive is explicitly exploratory and no universal improvement is claimed.", "PLS-LDA UMAP: 6/10, median +0.025, 95% CI [-0.067,+0.084], Holm p=0.557; KNN openTSNE: 0/10, Holm p=0.0078"),
        ("HPC provenance", "Uploaded run manifests preserve dataset and parameter metadata but record git_commit=NA. The manuscript does not retrospectively assign them to a source revision.", "Clean tagged confirmatory rerun with source checksum and container digest pending"),
        ("Cover letter", "A track-specific draft includes prior-publication disclosure, license, URL, version, software delta, and explicit author-only fields.", "Draft complete; author fields pending"),
        ("Release sustainability", "README, changelog, contribution guide, API policy, prior-version delta, validation protocol, and release checklist are included.", "Completed locally"),
    ]
    supplement.add_table(doc, ("Comment", "Response", "Status"), rows, [1.55, 4.15, 1.25])
    doc.save(READINESS_DOCX)


def main() -> None:
    supplement.build_bib()
    build_main_tex()
    build_cover_letter()
    build_readiness_report()
    print(MAIN_TEX)
    print(COVER_DOCX)
    print(COVER_TXT)
    print(READINESS_DOCX)


if __name__ == "__main__":
    main()
